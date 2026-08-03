import io
import os
import re
import json
import enum
import uuid
import math
import hashlib
from collections import Counter
from datetime import datetime
from typing import List, Optional, Dict, Any, Tuple

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from sqlalchemy import Column, String, Text, Integer, Float, DateTime, Enum, ForeignKey, create_engine, desc, func as sa_func
from sqlalchemy.orm import declarative_base, sessionmaker, Session, relationship
from sqlalchemy.sql import func
from pydantic import BaseModel

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None

try:
    import docx as docx_lib  # python-docx
except ImportError:  # pragma: no cover
    docx_lib = None

DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./bug_analyzer.db")
connect_args = {"check_same_thread": False} if DATABASE_URL.startswith("sqlite") else {}

engine = create_engine(DATABASE_URL, connect_args=connect_args)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

class SourceType(str, enum.Enum):
    PASTE = "paste"
    UPLOAD = "upload"


class Severity(str, enum.Enum):
    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"
    CRITICAL = "critical"


class BugSubmission(Base):
    __tablename__ = "bug_submissions"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    bug_number = Column(Integer, nullable=False)
    bug_id = Column(String(20), unique=True, index=True, nullable=False)

    title = Column(String(255), nullable=False)
    severity = Column(Enum(Severity), nullable=False, default=Severity.MEDIUM)
    source_type = Column(Enum(SourceType), nullable=False)
    original_filename = Column(String(255), nullable=True)
    file_type = Column(String(50), nullable=True)
    content = Column(Text, nullable=False)
    content_length = Column(String(20), nullable=True)
    created_at = Column(DateTime(timezone=True), server_default=func.now())

    analysis = relationship(
        "AgentAnalysis", back_populates="bug", uselist=False, cascade="all, delete-orphan"
    )


class AgentAnalysis(Base):

    __tablename__ = "agent_analyses"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    bug_id = Column(String, ForeignKey("bug_submissions.id"), unique=True, nullable=False, index=True)

    triage_severity = Column(String(20), nullable=False)
    triage_priority = Column(String(10), nullable=False)
    triage_type = Column(String(30), nullable=False)
    triage_component = Column(String(100), nullable=False)
    triage_confidence = Column(Float, nullable=False)
    triage_reasoning = Column(Text, nullable=False)

    log_exception_type = Column(String(150), nullable=True)
    log_message = Column(Text, nullable=True)
    log_language = Column(String(30), nullable=True)
    log_failure_file = Column(String(255), nullable=True)
    log_failure_line = Column(String(20), nullable=True)
    log_affected_paths = Column(Text, nullable=True)
    log_error_signature = Column(String(64), nullable=True)
    log_confidence = Column(Float, nullable=False)
    log_reasoning = Column(Text, nullable=False)

    root_cause_hypothesis = Column(Text, nullable=True)
    root_cause_confidence = Column(Float, nullable=True)
    root_cause_evidence = Column(Text, nullable=True)  

    duplicate_matches = Column(Text, nullable=True)     
    remediation_summary = Column(Text, nullable=True)
    remediation_steps = Column(Text, nullable=True)     
    remediation_confidence = Column(Float, nullable=True)
    remediation_sources = Column(Text, nullable=True)   
    combined_json = Column(Text, nullable=False)

    created_at = Column(DateTime(timezone=True), server_default=func.now())

    bug = relationship("BugSubmission", back_populates="analysis")


class HistoricalDefect(Base):
  
    __tablename__ = "historical_defects"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    defect_id = Column(String(20), unique=True, nullable=False)
    title = Column(String(255), nullable=False)
    symptom_text = Column(Text, nullable=False)     # what the report/log looked like
    root_cause = Column(Text, nullable=False)
    resolution_summary = Column(Text, nullable=False)
    component = Column(String(100), nullable=False)
    tags = Column(String(255), nullable=True)        # comma-separated

class BugSubmissionResponse(BaseModel):
    id: str
    bug_id: str
    title: str
    severity: str
    source_type: str
    original_filename: Optional[str]
    file_type: Optional[str]
    content_length: Optional[str]
    created_at: datetime

    class Config:
        from_attributes = True


class BugSubmissionDetail(BugSubmissionResponse):
    content: str


class BugSubmissionCreatedResponse(BaseModel):
    id: str
    bug_id: str
    message: str
    analysis: Optional[Dict[str, Any]] = None


TEXT_EXTENSIONS = {".txt", ".log", ".json", ".xml", ".csv", ".stacktrace", ".out"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_EXTENSIONS = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB
MAX_FILES_PER_UPLOAD = 10
UPLOAD_DIR = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

BUG_ID_PREFIX = "BUG-"
BUG_ID_PAD = 5  # BUG-00001

Base.metadata.create_all(bind=engine)

app = FastAPI(
    title="AI Smart Bug Analyzer & Fix Advisor — Full Agent Layer",
    version="3.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def _validate_severity(severity: str) -> Severity:
    try:
        return Severity(severity.lower())
    except ValueError:
        allowed = "low, medium, high, critical"
        raise HTTPException(
            status_code=400,
            detail=f"Invalid severity '{severity}'. Must be one of: {allowed}.",
        )


def _validate_extension(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed types: {allowed}",
        )
    return ext


def _extract_pdf_text(filename: str, raw: bytes) -> str:
    if PdfReader is None:
        raise HTTPException(
            status_code=500,
            detail="PDF support is not installed on the server. Run: pip install pypdf",
        )
    try:
        reader = PdfReader(io.BytesIO(raw))
        pages_text = []
        for page in reader.pages:
            pages_text.append(page.extract_text() or "")
        text = "\n".join(pages_text).strip()
    except Exception:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' could not be read as a valid PDF file.",
        )
    if not text:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' contains no extractable text (it may be a scanned/image-only PDF).",
        )
    return text


def _extract_docx_text(filename: str, raw: bytes) -> str:
    if docx_lib is None:
        raise HTTPException(
            status_code=500,
            detail="DOCX support is not installed on the server. Run: pip install python-docx",
        )
    try:
        document = docx_lib.Document(io.BytesIO(raw))
        paragraphs = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        paragraphs.append(cell.text)
        text = "\n".join(paragraphs).strip()
    except Exception:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' could not be read as a valid DOCX file.",
        )
    if not text:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' contains no extractable text.",
        )
    return text


def _extract_text_content(filename: str, ext: str, raw: bytes) -> str:
    if ext == ".pdf":
        return _extract_pdf_text(filename, raw)
    if ext == ".docx":
        return _extract_docx_text(filename, raw)
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        raise HTTPException(
            status_code=400, detail=f"'{filename}' must be UTF-8 encoded plain text."
        )


def _next_bug_number(db: Session) -> int:
    current_max = db.query(sa_func.max(BugSubmission.bug_number)).scalar()
    return (current_max or 0) + 1


def _format_bug_id(bug_number: int) -> str:
    return f"{BUG_ID_PREFIX}{bug_number:0{BUG_ID_PAD}d}"


def _serialize(bug: BugSubmission) -> dict:
    return {
        "id": bug.id,
        "bug_id": bug.bug_id,
        "title": bug.title,
        "severity": bug.severity.value if hasattr(bug.severity, "value") else bug.severity,
        "source_type": bug.source_type.value if hasattr(bug.source_type, "value") else bug.source_type,
        "original_filename": bug.original_filename,
        "file_type": bug.file_type,
        "content_length": bug.content_length,
        "created_at": bug.created_at,
    }
class TriageAgent:
    SEVERITY_KEYWORDS = {
        Severity.CRITICAL: [
            "data loss", "corrupt", "security", "exploit", "vulnerability",
            "breach", "outage", "down", "crash", "crashed", "crashing",
            "deadlock", "unresponsive", "cannot start", "won't start",
            "production is down", "database is down", "payment failure",
            "segfault", "segmentation fault", "fatal", "kernel panic",
        ],
        Severity.HIGH: [
            "exception", "error", "fail", "failed", "failure", "broken",
            "blocked", "blocker", "500", "internal server error",
            "null pointer", "nullpointerexception", "stack trace",
            "traceback", "cannot login", "can't login", "not working",
            "regression", "timeout", "connection refused",
        ],
        Severity.MEDIUM: [
            "warning", "warn", "intermittent", "sometimes", "occasionally",
            "slow", "performance", "deprecated", "inconsistent",
            "unexpected behavior", "edge case", "workaround",
        ],
        Severity.LOW: [
            "typo", "cosmetic", "misaligned", "misalignment", "ui glitch",
            "minor", "spacing", "color", "colour", "tooltip", "wording",
            "nice to have", "enhancement",
        ],
    }

    SEVERITY_WEIGHT = {
        Severity.CRITICAL: 4,
        Severity.HIGH: 3,
        Severity.MEDIUM: 2,
        Severity.LOW: 1,
    }

    STRUCTURAL_CRITICAL_PATTERNS = [
        r"Traceback \(most recent call last\)",
        r"Exception in thread",
        r"Segmentation fault",
        r"PANIC|panic:",
    ]

    COMPONENT_KEYWORDS = {
        "Authentication": ["login", "auth", "token", "session", "password", "oauth", "jwt"],
        "Payments": ["payment", "checkout", "stripe", "billing", "invoice", "charge", "refund"],
        "Database": ["database", "sql", "query", "postgres", "sqlite", "orm", "connection pool", "deadlock"],
        "API": ["api", "endpoint", "request", "response", "500", "404", "rest", "http"],
        "Frontend/UI": ["ui", "button", "css", "layout", "render", "dom", "browser", "misaligned", "tooltip"],
        "File Upload": ["upload", "file", "attachment", "pdf", "docx", "multipart"],
        "Networking": ["timeout", "connection refused", "network", "socket", "dns"],
        "Infrastructure": ["deploy", "docker", "kubernetes", "server", "outage", "down", "crash"],
    }

    TYPE_KEYWORDS = {
        "Enhancement": ["enhancement", "feature request", "would be nice", "nice to have", "improve"],
        "Others": ["question", "how do i", "documentation", "docs"],
    }

    def classify(self, content: str) -> Dict[str, Any]:
        text = content or ""
        lower = text.lower()
        reasoning_parts: List[str] = []

        scores: Dict[Severity, int] = {s: 0 for s in Severity}
        matched_keywords: Dict[Severity, List[str]] = {s: [] for s in Severity}

        for sev, keywords in self.SEVERITY_KEYWORDS.items():
            for kw in keywords:
                if kw in lower:
                    scores[sev] += 1
                    matched_keywords[sev].append(kw)

        structural_hit = False
        for pattern in self.STRUCTURAL_CRITICAL_PATTERNS:
            if re.search(pattern, text):
                scores[Severity.HIGH] += 2
                structural_hit = True
                reasoning_parts.append(f"structural signal matched pattern '{pattern}'")

        best_severity = Severity.MEDIUM
        best_weighted = -1
        best_matches = -1
        total_matches = 0

        severity_rank = {
            Severity.LOW: 0,
            Severity.MEDIUM: 1,
            Severity.HIGH: 2,
            Severity.CRITICAL: 3,
        }

        for sev, score in scores.items():
            total_matches += score
            weighted = score * self.SEVERITY_WEIGHT[sev]
            if score == 0:
                continue
            if (
                weighted > best_weighted
                or (weighted == best_weighted and score > best_matches)
                or (weighted == best_weighted and score == best_matches
                    and severity_rank[sev] > severity_rank[best_severity])
            ):
                best_weighted = weighted
                best_matches = score
                best_severity = sev

        MIN_SCORE_FOR_ELEVATED = 2
        if total_matches == 0:
            best_severity = Severity.MEDIUM
            reasoning_parts.append("no strong severity keywords found; defaulted to Medium")
        elif (
            best_severity in (Severity.HIGH, Severity.CRITICAL)
            and best_matches < MIN_SCORE_FOR_ELEVATED
            and not structural_hit
        ):
            reasoning_parts.append(
                f"only a single weak '{best_severity.value}' signal found "
                f"({matched_keywords[best_severity]}); downgraded to Medium "
                f"since evidence was insufficient"
            )
            best_severity = Severity.MEDIUM
        else:
            kws = matched_keywords[best_severity]
            if kws:
                reasoning_parts.append(
                    f"matched {len(kws)} {best_severity.value}-severity keyword(s): "
                    + ", ".join(sorted(set(kws))[:5])
                )

        if total_matches > 0:
            confidence = matched_keywords[best_severity].__len__() / max(total_matches, 1)
            confidence = 0.5 + confidence * 0.4
        else:
            confidence = 0.4
        if structural_hit and best_severity in (Severity.HIGH, Severity.CRITICAL):
            confidence = min(confidence + 0.15, 0.97)
        confidence = round(min(max(confidence, 0.35), 0.97), 2)

        urgency_words = ["urgent", "asap", "immediately", "blocking release", "blocker"]
        urgent_hit = any(w in lower for w in urgency_words)
        priority_map = {
            Severity.CRITICAL: "P0",
            Severity.HIGH: "P1",
            Severity.MEDIUM: "P2",
            Severity.LOW: "P3",
        }
        priority = priority_map[best_severity]
        if urgent_hit and priority in ("P1", "P2"):
            priority = "P0" if priority == "P1" else "P1"
            reasoning_parts.append("urgency language detected; priority bumped one level")

        component_scores: Dict[str, int] = {c: 0 for c in self.COMPONENT_KEYWORDS}
        for comp, keywords in self.COMPONENT_KEYWORDS.items():
            for kw in keywords:
                if kw in lower:
                    component_scores[comp] += 1
        best_component = max(component_scores, key=component_scores.get)
        if component_scores[best_component] == 0:
            best_component = "Unclassified"
            reasoning_parts.append("no component keywords matched; marked Unclassified")
        else:
            reasoning_parts.append(
                f"component inferred from keyword(s) associated with '{best_component}'"
            )

        bug_type = "Bug"
        for t, keywords in self.TYPE_KEYWORDS.items():
            if any(kw in lower for kw in keywords):
                bug_type = t
                break

        return {
            "severity": best_severity.value,
            "priority": priority,
            "type": bug_type,
            "component": best_component,
            "confidence": confidence,
            "reasoning": "; ".join(reasoning_parts) if reasoning_parts else "insufficient signal in report",
        }


class LogAnalysisAgent:
    CODE_FILE_PATTERN = re.compile(
        r"([\w./\\-]+\.(?:py|js|ts|jsx|tsx|java|go|rb|cpp|c|cs|php|kt|rs)):(\d+)"
    )

    def analyze(self, content: str) -> Dict[str, Any]:
        text = content or ""
        reasoning_parts: List[str] = []

        language = None
        exception_type = None
        message = None
        failure_file = None
        failure_line = None

        if re.search(r"Traceback \(most recent call last\)", text):
            language = "Python"
            reasoning_parts.append("detected Python traceback header")
            file_matches = re.findall(r'File "([^"]+)", line (\d+)', text)
            if file_matches:
                failure_file, failure_line = file_matches[-1]
                reasoning_parts.append("failure point taken as last frame in traceback")
            exc_match = re.search(r"^([A-Za-z_.]+(?:Error|Exception|Warning))\s*:\s*(.*)$", text, re.MULTILINE)
            if exc_match:
                exception_type, message = exc_match.group(1), exc_match.group(2).strip()

        elif re.search(r"Exception in thread|at [\w$.<>]+\(.*\.java:\d+\)", text):
            language = "Java"
            reasoning_parts.append("detected Java stack trace signature")
            m = re.search(r"\(([\w.$]+\.java):(\d+)\)", text)
            if m:
                failure_file, failure_line = m.group(1), m.group(2)
            exc_match = re.search(r"([\w.]+(?:Exception|Error))(?::\s*(.*))?", text)
            if exc_match:
                exception_type = exc_match.group(1)
                message = (exc_match.group(2) or "").strip() or None

        elif re.search(r"at .*\(.*:\d+:\d+\)|Uncaught|TypeError|ReferenceError", text):
            language = "JavaScript"
            reasoning_parts.append("detected JS/Node runtime error signature")
            m = re.search(r"\(([^()]+\.(?:js|ts|jsx|tsx)):(\d+):(\d+)\)", text)
            if not m:
                m = re.search(r"([\w./\\-]+\.(?:js|ts|jsx|tsx)):(\d+):(\d+)", text)
            if m:
                failure_file, failure_line = m.group(1), m.group(2)
            exc_match = re.search(r"(?:Uncaught\s+)?([A-Za-z]+(?:Error|Exception))\s*:\s*(.*)", text)
            if exc_match:
                exception_type = exc_match.group(1)
                message = exc_match.group(2).strip().split("\n")[0] or None

        elif re.search(r"^\d{4}-\d{2}-\d{2}[T ]\d{2}:\d{2}:\d{2}", text, re.MULTILINE):
            language = "log"
            reasoning_parts.append("detected structured timestamped log format")
            lvl = re.search(r"\b(ERROR|WARN|FATAL|CRITICAL)\b", text)
            if lvl:
                exception_type = f"{lvl.group(1)} log entry"

        else:
            reasoning_parts.append("no recognized stack trace signature; treated as unstructured report")

        if not failure_file:
            m = self.CODE_FILE_PATTERN.search(text)
            if m:
                failure_file, failure_line = m.group(1), m.group(2)
                reasoning_parts.append("failure point recovered via generic file:line pattern")

        affected_paths = []
        for m in self.CODE_FILE_PATTERN.finditer(text):
            path = m.group(1)
            if path not in affected_paths:
                affected_paths.append(path)
        for m in re.finditer(r'File "([^"]+)"', text):
            path = m.group(1)
            if path not in affected_paths:
                affected_paths.append(path)
        if failure_file and failure_file not in affected_paths:
            affected_paths.insert(0, failure_file)
        affected_paths = affected_paths[:10]

        sig_source = f"{exception_type or 'unknown'}::{(message or '').strip().lower()[:200]}"
        error_signature = hashlib.sha256(sig_source.encode("utf-8")).hexdigest()[:16]

        confidence = 0.4
        if language and language != "log":
            confidence += 0.25
        if exception_type:
            confidence += 0.2
        if failure_file:
            confidence += 0.1
        confidence = round(min(confidence, 0.97), 2)

        if not text.strip():
            confidence = 0.0
            reasoning_parts = ["empty content — nothing to analyze"]

        return {
            "exception_type": exception_type,
            "message": message,
            "language": language,
            "failure_file": failure_file,
            "failure_line": failure_line,
            "affected_paths": affected_paths,
            "error_signature": error_signature,
            "confidence": confidence,
            "reasoning": "; ".join(reasoning_parts),
        }

_TOKEN_RE = re.compile(r"[A-Za-z][A-Za-z0-9_]{1,}")

_STOPWORDS = {
    "the", "a", "an", "is", "are", "was", "were", "be", "been", "being",
    "to", "of", "in", "on", "at", "for", "and", "or", "but", "with",
    "this", "that", "it", "its", "as", "by", "from", "into", "not",
    "has", "have", "had", "do", "does", "did", "will", "would", "can",
    "could", "should", "if", "then", "than", "so", "we", "you", "i",
    "my", "your", "our", "their", "his", "her", "them", "us", "line",
    "file", "most", "recent", "call", "last",
}


def _tokenize(text: str) -> List[str]:
    return [t.lower() for t in _TOKEN_RE.findall(text or "") if t.lower() not in _STOPWORDS]


class TfidfIndex:
    def __init__(self, documents: List[Tuple[str, str]]):
        # documents: list of (doc_id, raw_text)
        self.doc_ids = [d[0] for d in documents]
        self.doc_tokens = [_tokenize(d[1]) for d in documents]
        self.n_docs = len(documents)

        df = Counter()
        for tokens in self.doc_tokens:
            for term in set(tokens):
                df[term] += 1
        self.idf = {
            term: math.log((1 + self.n_docs) / (1 + count)) + 1.0
            for term, count in df.items()
        }

        self.doc_vectors = [self._vectorize(tokens) for tokens in self.doc_tokens]

    def _vectorize(self, tokens: List[str]) -> Dict[str, float]:
        if not tokens:
            return {}
        tf = Counter(tokens)
        max_tf = max(tf.values())
        vec = {}
        for term, count in tf.items():
            idf = self.idf.get(term, math.log((1 + self.n_docs) / 1) + 1.0)
            vec[term] = (count / max_tf) * idf
        norm = math.sqrt(sum(v * v for v in vec.values())) or 1.0
        return {k: v / norm for k, v in vec.items()}

    def _cosine(self, vec_a: Dict[str, float], vec_b: Dict[str, float]) -> float:
        if not vec_a or not vec_b:
            return 0.0
        shared = set(vec_a.keys()) & set(vec_b.keys())
        return sum(vec_a[t] * vec_b[t] for t in shared)

    def top_k_matches(self, query_text: str, k: int = 3, exclude_id: Optional[str] = None) -> List[Tuple[str, float]]:
        query_tokens = _tokenize(query_text)
        query_vec = self._vectorize(query_tokens)
        scored = []
        for doc_id, doc_vec in zip(self.doc_ids, self.doc_vectors):
            if doc_id == exclude_id:
                continue
            sim = self._cosine(query_vec, doc_vec)
            if sim > 0:
                scored.append((doc_id, round(sim, 4)))
        scored.sort(key=lambda x: x[1], reverse=True)
        return scored[:k]

SEED_HISTORICAL_DEFECTS = [
    {
        "defect_id": "HDB-001",
        "title": "Checkout fails with CardError on Stripe charge",
        "symptom_text": (
            "Traceback (most recent call last): File \"app/views.py\", line 88, in checkout "
            "charge = stripe.Charge.create(...) stripe.error.CardError: Your card was declined. "
            "500 internal server error during checkout payment"
        ),
        "root_cause": (
            "Stripe API key for the live environment was pointed at a test-mode key after a "
            "config deploy, causing all real cards to be rejected as declined."
        ),
        "resolution_summary": (
            "Corrected STRIPE_SECRET_KEY env var to the live key in the deploy pipeline and "
            "added a startup check that fails fast if a test key is used in production."
        ),
        "component": "Payments",
        "tags": "stripe,charge,carderror,payment,checkout,500",
    },
    {
        "defect_id": "HDB-002",
        "title": "Login endpoint throws NullPointerException on session lookup",
        "symptom_text": (
            "Exception in thread \"main\" java.lang.NullPointerException at "
            "com.app.auth.SessionService.validate(SessionService.java:112) cannot login token session"
        ),
        "root_cause": (
            "Session cache entry expired but the code did not null-check before dereferencing "
            "the session object, causing a NullPointerException instead of a clean re-login prompt."
        ),
        "resolution_summary": (
            "Added a null check with graceful redirect-to-login fallback in SessionService.validate; "
            "added unit test for expired-session path."
        ),
        "component": "Authentication",
        "tags": "auth,login,session,nullpointerexception,java,token",
    },
    {
        "defect_id": "HDB-003",
        "title": "Database connection pool exhausted under load, requests time out",
        "symptom_text": (
            "connection refused timeout database postgres query connection pool deadlock "
            "requests hanging under load, 500s spike"
        ),
        "root_cause": (
            "Connection pool size was left at the default (5) while concurrent request volume "
            "grew, so the pool exhausted and new requests queued until timeout."
        ),
        "resolution_summary": (
            "Increased SQLAlchemy pool_size/max_overflow and added pool usage metrics + alerting "
            "so exhaustion is caught before it causes timeouts."
        ),
        "component": "Database",
        "tags": "database,postgres,connection pool,timeout,deadlock,sql",
    },
    {
        "defect_id": "HDB-004",
        "title": "Uncaught TypeError on checkout button click (undefined cart total)",
        "symptom_text": (
            "Uncaught TypeError: Cannot read properties of undefined (reading 'total') "
            "at CheckoutButton.js:45:12 button ui render dom browser"
        ),
        "root_cause": (
            "Cart total was read from state before the async cart-fetch resolved, so `cart` was "
            "still undefined on first render, causing a TypeError on the .total access."
        ),
        "resolution_summary": (
            "Added a loading guard so CheckoutButton renders a disabled state until cart data "
            "resolves; added default empty-cart object to prevent undefined access."
        ),
        "component": "Frontend/UI",
        "tags": "typeerror,javascript,ui,checkout,button,undefined,cart",
    },
    {
        "defect_id": "HDB-005",
        "title": "File upload endpoint returns 500 for large PDF attachments",
        "symptom_text": (
            "500 internal server error api endpoint upload file attachment pdf multipart "
            "request failed for files over a few MB"
        ),
        "root_cause": (
            "Upload endpoint had no explicit request body size limit configured, so large "
            "multipart PDFs exceeded the ASGI server's default buffer and the request was "
            "dropped mid-stream, surfacing as an unhandled 500."
        ),
        "resolution_summary": (
            "Configured explicit max upload size with a clear 413 response before the limit, "
            "and streamed large multipart bodies to disk instead of buffering in memory."
        ),
        "component": "File Upload",
        "tags": "upload,file,pdf,multipart,api,500,attachment",
    },
    {
        "defect_id": "HDB-006",
        "title": "Intermittent slow response on API endpoint during peak hours",
        "symptom_text": (
            "warning intermittent slow performance sometimes occasionally endpoint response "
            "degraded during peak hours, no errors just latency"
        ),
        "root_cause": (
            "N+1 query pattern in the API serializer was fetching related rows one at a time "
            "instead of using a join/prefetch, so latency scaled with result set size and only "
            "became noticeable under peak load."
        ),
        "resolution_summary": (
            "Refactored serializer to use select_related/prefetch to eliminate the N+1 pattern; "
            "added query count assertions in tests to prevent regression."
        ),
        "component": "API",
        "tags": "api,performance,slow,n+1,query,endpoint",
    },
    {
        "defect_id": "HDB-007",
        "title": "Production outage: service unresponsive, health checks failing",
        "symptom_text": (
            "production is down outage unresponsive crash server infrastructure deploy docker "
            "kubernetes health check failing fatal"
        ),
        "root_cause": (
            "A bad deploy shipped a container with a missing environment variable required at "
            "startup, causing the process to crash-loop and fail all health checks."
        ),
        "resolution_summary": (
            "Rolled back to the previous image, added the missing env var to the deploy config, "
            "and added a pre-deploy config validation step to the CI pipeline."
        ),
        "component": "Infrastructure",
        "tags": "outage,down,crash,deploy,docker,kubernetes,infrastructure",
    },
    {
        "defect_id": "HDB-008",
        "title": "Minor UI misalignment on settings page tooltip",
        "symptom_text": (
            "cosmetic minor misaligned tooltip spacing color ui glitch settings page, nice to have fix"
        ),
        "root_cause": (
            "A CSS flexbox container was missing an explicit align-items rule, causing the "
            "tooltip to inherit default stretch alignment instead of centering."
        ),
        "resolution_summary": (
            "Added align-items: center to the tooltip's flex container; verified across "
            "breakpoints in the design review."
        ),
        "component": "Frontend/UI",
        "tags": "ui,cosmetic,tooltip,css,misaligned,minor",
    },
    {
        "defect_id": "HDB-009",
        "title": "Socket timeout connecting to third-party pricing API",
        "symptom_text": (
            "connection refused network socket dns timeout third-party api request failed "
            "networking error"
        ),
        "root_cause": (
            "Outbound requests to the third-party pricing API had no configured timeout, so a "
            "slow upstream response held the connection open until the OS-level socket timeout, "
            "surfacing as an unhandled hang rather than a fast, retryable failure."
        ),
        "resolution_summary": (
            "Set an explicit client-side timeout with exponential backoff retry and a circuit "
            "breaker around the third-party pricing client."
        ),
        "component": "Networking",
        "tags": "networking,timeout,socket,dns,connection refused,api",
    },
    {
        "defect_id": "HDB-010",
        "title": "Data corruption after concurrent writes to same order record",
        "symptom_text": (
            "data loss corrupt database deadlock concurrent writes order record inconsistent state"
        ),
        "root_cause": (
            "Two request handlers updated the same order row without row-level locking or "
            "optimistic concurrency control, causing a lost-update race condition under "
            "concurrent traffic."
        ),
        "resolution_summary": (
            "Added optimistic locking (version column) on the order model and retry-on-conflict "
            "logic in the update path."
        ),
        "component": "Database",
        "tags": "data loss,corrupt,database,deadlock,concurrency,race condition",
    },
]


def _seed_historical_defects(db: Session):
    existing = db.query(sa_func.count(HistoricalDefect.id)).scalar()
    if existing and existing > 0:
        return
    for d in SEED_HISTORICAL_DEFECTS:
        db.add(HistoricalDefect(
            id=str(uuid.uuid4()),
            defect_id=d["defect_id"],
            title=d["title"],
            symptom_text=d["symptom_text"],
            root_cause=d["root_cause"],
            resolution_summary=d["resolution_summary"],
            component=d["component"],
            tags=d["tags"],
        ))
    db.commit()
class RootCauseAgent:
    def __init__(self):
        pass

    def _build_query_text(self, content: str, triage: Dict[str, Any], log_result: Dict[str, Any]) -> str:
        """'Build the Root Cause Prompt' step — assembles the retrieval query
        from the raw report plus structured Triage/Log Analysis signals, so
        retrieval isn't relying on raw text alone."""
        parts = [content or ""]
        if triage.get("component"):
            parts.append(triage["component"])
        if log_result.get("exception_type"):
            parts.append(log_result["exception_type"])
        if log_result.get("message"):
            parts.append(log_result["message"])
        return " ".join(parts)

    def analyze(
        self,
        content: str,
        triage: Dict[str, Any],
        log_result: Dict[str, Any],
        historical_defects: List[HistoricalDefect],
        exclude_defect_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        if not historical_defects:
            return {
                "hypothesis": "No historical defect knowledge base available to ground a root cause hypothesis.",
                "confidence": 0.0,
                "evidence": [],
                "reasoning": "Historical defect KB is empty.",
            }
        corpus = [(d.defect_id, f"{d.title} {d.symptom_text} {d.tags or ''}") for d in historical_defects]
        index = TfidfIndex(corpus)
        query_text = self._build_query_text(content, triage, log_result)
        matches = index.top_k_matches(query_text, k=3, exclude_id=exclude_defect_id)

        by_id = {d.defect_id: d for d in historical_defects}
        evidence = []
        for defect_id, sim in matches:
            d = by_id[defect_id]
            evidence.append({
                "defect_id": d.defect_id,
                "title": d.title,
                "similarity": sim,
                "root_cause": d.root_cause,
                "component": d.component,
            })

        reasoning_parts = [f"retrieved {len(evidence)} candidate historical defect(s) via TF-IDF cosine similarity"]
        SIM_STRONG = 0.35
        SIM_WEAK = 0.15

        if evidence and evidence[0]["similarity"] >= SIM_STRONG:
            top = evidence[0]
            hypothesis = (
                f"Likely same root cause pattern as historical defect {top['defect_id']} "
                f"(\"{top['title']}\"): {top['root_cause']}"
            )
            confidence = round(min(0.55 + top["similarity"] * 0.5, 0.95), 2)
            reasoning_parts.append(
                f"top match {top['defect_id']} scored {top['similarity']} (>= {SIM_STRONG} strong-match threshold)"
            )
        elif evidence and evidence[0]["similarity"] >= SIM_WEAK:
            top = evidence[0]
            hypothesis = (
                f"Possible partial match to historical defect {top['defect_id']} "
                f"(\"{top['title']}\"), with root cause: {top['root_cause']} — "
                f"treat as a weak hypothesis pending engineer review, since similarity is moderate."
            )
            confidence = round(0.3 + top["similarity"] * 0.4, 2)
            reasoning_parts.append(
                f"top match {top['defect_id']} scored {top['similarity']} (weak-match band, {SIM_WEAK}-{SIM_STRONG})"
            )
        else:
            # Heuristic fallback grounded in Log Analysis + Triage, no strong
            # historical precedent found.
            exc = log_result.get("exception_type")
            comp = triage.get("component")
            if exc and comp and comp != "Unclassified":
                hypothesis = (
                    f"No close historical match found. Based on the '{exc}' exception signature "
                    f"and '{comp}' component classification, the likely root cause is a fault in "
                    f"the {comp.lower()} code path around the failure point"
                    + (f" ({log_result.get('failure_file')}:{log_result.get('failure_line')})"
                       if log_result.get("failure_file") else "")
                    + " — recommend engineer investigation to confirm, as this is a heuristic estimate."
                )
                confidence = 0.35
            elif exc:
                hypothesis = (
                    f"No close historical match found. The '{exc}' exception signature suggests "
                    f"an unhandled error condition at the failure point"
                    + (f" ({log_result.get('failure_file')}:{log_result.get('failure_line')})"
                       if log_result.get("failure_file") else "")
                    + " — root cause could not be narrowed further without deeper investigation."
                )
                confidence = 0.3
            else:
                hypothesis = (
                    "No close historical match and no structured exception signature found; "
                    "insufficient evidence to propose a specific root cause. Recommend manual triage."
                )
                confidence = 0.2
            reasoning_parts.append(
                f"best available similarity {evidence[0]['similarity'] if evidence else 0.0} "
                f"below weak-match threshold ({SIM_WEAK}); fell back to Log Analysis/Triage heuristic"
            )

        return {
            "hypothesis": hypothesis,
            "confidence": confidence,
            "evidence": evidence,
            "reasoning": "; ".join(reasoning_parts),
        }

class DuplicateDetectionAgent:
    SIMILARITY_THRESHOLD = 0.20  # below this, not considered a duplicate candidate

    def _build_query_text(self, content: str, triage: Dict[str, Any], log_result: Dict[str, Any]) -> str:
        parts = [content or ""]
        if triage.get("component"):
            parts.append(triage["component"])
        if log_result.get("exception_type"):
            parts.append(log_result["exception_type"])
        return " ".join(parts)

    def search(
        self,
        content: str,
        triage: Dict[str, Any],
        log_result: Dict[str, Any],
        historical_defects: List[HistoricalDefect],
        prior_submissions: List[Tuple[BugSubmission, Optional[AgentAnalysis]]],
        exclude_bug_db_id: Optional[str] = None,
        top_k: int = 5,
    ) -> Dict[str, Any]:
        query_text = self._build_query_text(content, triage, log_result)

        corpus: List[Tuple[str, str]] = []
        lookup: Dict[str, Dict[str, Any]] = {}

        for d in historical_defects:
            doc_id = f"kb:{d.defect_id}"
            corpus.append((doc_id, f"{d.title} {d.symptom_text} {d.tags or ''}"))
            lookup[doc_id] = {
                "source": "historical_kb",
                "ref_id": d.defect_id,
                "title": d.title,
                "resolution_summary": d.resolution_summary,
                "component": d.component,
            }

        for bug, analysis in prior_submissions:
            if bug.id == exclude_bug_db_id:
                continue
            doc_id = f"sub:{bug.id}"
            text = bug.title + " " + (bug.content or "")
            corpus.append((doc_id, text))
            resolution_summary = None
            if analysis and analysis.root_cause_hypothesis:
                resolution_summary = analysis.root_cause_hypothesis
            lookup[doc_id] = {
                "source": "prior_submission",
                "ref_id": bug.bug_id,
                "title": bug.title,
                "resolution_summary": resolution_summary or "No resolution recorded yet for this prior submission.",
                "component": analysis.triage_component if analysis else None,
            }

        if not corpus:
            return {"matches": [], "reasoning": "No historical KB entries or prior submissions to compare against."}

        index = TfidfIndex(corpus)
        raw_matches = index.top_k_matches(query_text, k=top_k)

        matches = []
        for doc_id, sim in raw_matches:
            if sim < self.SIMILARITY_THRESHOLD:
                continue
            info = lookup[doc_id]
            matches.append({
                "source": info["source"],
                "id": info["ref_id"],
                "title": info["title"],
                "similarity": sim,
                "similarity_pct": round(sim * 100, 1),
                "resolution_summary": info["resolution_summary"],
                "component": info["component"],
            })

        reasoning = (
            f"searched {len(corpus)} candidate(s) ({len(historical_defects)} historical KB, "
            f"{len(corpus) - len(historical_defects)} prior submissions) via TF-IDF cosine similarity; "
            f"{len(matches)} passed the {self.SIMILARITY_THRESHOLD} similarity threshold"
        )

        return {"matches": matches, "reasoning": reasoning}
class RemediationAgent:
   
    BEST_PRACTICE_PLAYBOOKS = {
        "Authentication": [
            "Add explicit null/expiry checks around session and token lookups before use.",
            "Add a regression test covering the expired-session / invalid-token path.",
            "Ensure failed auth attempts degrade gracefully (redirect/login prompt) rather than throwing unhandled exceptions.",
        ],
        "Payments": [
            "Verify payment provider API keys/environment config match the target environment (test vs live).",
            "Add a startup/config validation check that fails fast on misconfigured payment credentials.",
            "Add idempotency handling for retried payment requests to avoid double-charging during fix rollout.",
        ],
        "Database": [
            "Review connection pool sizing (pool_size/max_overflow) against current concurrent load.",
            "Add row-level locking or optimistic concurrency control for any concurrent-write paths implicated.",
            "Add query performance monitoring/alerting around the affected query path.",
        ],
        "API": [
            "Check for N+1 query patterns or unbounded loops in the affected endpoint's serialization path.",
            "Add explicit request size/timeout limits to the endpoint if not already present.",
            "Add an integration test reproducing the reported failure to prevent regression.",
        ],
        "Frontend/UI": [
            "Add a loading/guard state so components don't render against not-yet-resolved async data.",
            "Add a default/fallback value for any state accessed before initial load completes.",
            "Add a visual regression or unit test covering this component's failure state.",
        ],
        "File Upload": [
            "Set explicit max upload size limits and stream large files to disk instead of buffering in memory.",
            "Return a clear 413/validation error before hitting server-side limits, instead of an unhandled 500.",
            "Add a test with a large/edge-case file matching the reported file type and size.",
        ],
        "Networking": [
            "Set explicit client-side timeouts with exponential backoff retry for the affected outbound call.",
            "Add a circuit breaker around the failing dependency to fail fast under sustained outages.",
            "Add monitoring/alerting on timeout rate for this dependency.",
        ],
        "Infrastructure": [
            "Confirm the last deploy's config/environment variables against the previous known-good release.",
            "Add a pre-deploy config validation step to CI to catch missing/invalid env vars before rollout.",
            "Add health-check-based automatic rollback if the new deploy fails startup checks.",
        ],
        "Unclassified": [
            "Reproduce the issue locally with the reported steps/log to confirm scope before fixing.",
            "Add logging around the suspected failure point to narrow down root cause with more confidence.",
            "Once root cause is confirmed, add a regression test before closing the fix.",
        ],
    }

    def _collect_context(
        self, triage: Dict[str, Any], log_result: Dict[str, Any],
        root_cause: Dict[str, Any], duplicates: Dict[str, Any],
    ) -> Dict[str, Any]:
        """'Collect Context' step — just a light structuring pass; kept as
        its own method so the seam is explicit and testable."""
        return {
            "component": triage.get("component", "Unclassified"),
            "severity": triage.get("severity", "medium"),
            "exception_type": log_result.get("exception_type"),
            "root_cause_hypothesis": root_cause.get("hypothesis"),
            "root_cause_confidence": root_cause.get("confidence", 0.0),
            "top_evidence": root_cause.get("evidence", [])[:1],
            "top_duplicate": duplicates.get("matches", [])[:1],
        }

    def generate(
        self, triage: Dict[str, Any], log_result: Dict[str, Any],
        root_cause: Dict[str, Any], duplicates: Dict[str, Any],
    ) -> Dict[str, Any]:
        ctx = self._collect_context(triage, log_result, root_cause, duplicates)
        sources: List[str] = []
        steps: List[str] = []

        top_dup = ctx["top_duplicate"][0] if ctx["top_duplicate"] else None
        strong_duplicate = (
            top_dup and top_dup["similarity"] >= 0.35
            and top_dup["resolution_summary"]
            and "No resolution recorded" not in top_dup["resolution_summary"]
        )

        if strong_duplicate:
            steps.append(
                f"Apply the same fix pattern used for {top_dup['source'].replace('_', ' ')} "
                f"'{top_dup['id']}': {top_dup['resolution_summary']}"
            )
            sources.append(f"{top_dup['source']}:{top_dup['id']} (similarity {top_dup['similarity_pct']}%)")

        # Priority 2: root cause hypothesis (if confident).
        if ctx["root_cause_confidence"] and ctx["root_cause_confidence"] >= 0.5:
            steps.append(
                f"Address the identified root cause directly: {ctx['root_cause_hypothesis']}"
            )
            if ctx["top_evidence"]:
                sources.append(f"root_cause_evidence:{ctx['top_evidence'][0]['defect_id']}")
            else:
                sources.append("root_cause_agent_heuristic")
        elif not strong_duplicate:
            steps.append(
                "Root cause confidence is low — first reproduce the issue and add targeted logging "
                "around the failure point before attempting a fix, to avoid treating a symptom as the cause."
            )
            sources.append("remediation_agent_low_confidence_fallback")

        component = ctx["component"] if ctx["component"] in self.BEST_PRACTICE_PLAYBOOKS else "Unclassified"
        playbook = self.BEST_PRACTICE_PLAYBOOKS[component]
        steps.extend(playbook)
        sources.append(f"best_practice_playbook:{component}")
        seen = set()
        deduped_steps = []
        for s in steps:
            if s not in seen:
                deduped_steps.append(s)
                seen.add(s)
        dup_conf = top_dup["similarity"] if strong_duplicate else 0.0
        rc_conf = ctx["root_cause_confidence"] or 0.0
        confidence = round(min(0.9, max(dup_conf, rc_conf) * 0.9 + 0.1), 2)

        if strong_duplicate:
            summary = (
                f"Recommended fix is grounded primarily in a closely matching resolved issue "
                f"({top_dup['id']}, {top_dup['similarity_pct']}% similarity), supplemented with "
                f"{component} best-practice hardening steps."
            )
        elif ctx["root_cause_confidence"] and ctx["root_cause_confidence"] >= 0.5:
            summary = (
                f"Recommended fix targets the hypothesized root cause directly, supplemented with "
                f"{component} best-practice hardening steps."
            )
        else:
            summary = (
                f"No strong historical match or high-confidence root cause was found — recommendation "
                f"leads with investigation steps, followed by standard {component} best practices."
            )

        return {
            "summary": summary,
            "steps": deduped_steps,
            "confidence": confidence,
            "sources": sources,
        }


triage_agent = TriageAgent()
log_analysis_agent = LogAnalysisAgent()
root_cause_agent = RootCauseAgent()
duplicate_detection_agent = DuplicateDetectionAgent()
remediation_agent = RemediationAgent()


def run_orchestration(
    content: str,
    db: Session,
    exclude_bug_db_id: Optional[str] = None,
) -> Dict[str, Any]:
   
    triage_result = triage_agent.classify(content)
    log_result = log_analysis_agent.analyze(content)

    historical_defects = db.query(HistoricalDefect).all()

    root_cause_result = root_cause_agent.analyze(
        content, triage_result, log_result, historical_defects
    )

    prior_bugs = (
        db.query(BugSubmission)
        .filter(BugSubmission.id != exclude_bug_db_id if exclude_bug_db_id else True)
        .order_by(desc(BugSubmission.created_at))
        .limit(200)  # cap corpus size for a live-demo dataset
        .all()
    )
    prior_submissions = [(b, b.analysis) for b in prior_bugs]

    duplicate_result = duplicate_detection_agent.search(
        content, triage_result, log_result, historical_defects, prior_submissions,
        exclude_bug_db_id=exclude_bug_db_id,
    )

    remediation_result = remediation_agent.generate(
        triage_result, log_result, root_cause_result, duplicate_result
    )

    return {
        "triage": triage_result,
        "log_analysis": log_result,
        "root_cause": root_cause_result,
        "duplicates": duplicate_result,
        "remediation": remediation_result,
        "analyzed_at": datetime.utcnow().isoformat(),
    }


def _persist_analysis(db: Session, bug: BugSubmission, combined: Dict[str, Any]) -> AgentAnalysis:
    triage = combined["triage"]
    log_res = combined["log_analysis"]
    root_cause = combined["root_cause"]
    duplicates = combined["duplicates"]
    remediation = combined["remediation"]

    record = AgentAnalysis(
        bug_id=bug.id,
        triage_severity=triage["severity"],
        triage_priority=triage["priority"],
        triage_type=triage["type"],
        triage_component=triage["component"],
        triage_confidence=triage["confidence"],
        triage_reasoning=triage["reasoning"],
        log_exception_type=log_res["exception_type"],
        log_message=log_res["message"],
        log_language=log_res["language"],
        log_failure_file=log_res["failure_file"],
        log_failure_line=log_res["failure_line"],
        log_affected_paths=json.dumps(log_res["affected_paths"]),
        log_error_signature=log_res["error_signature"],
        log_confidence=log_res["confidence"],
        log_reasoning=log_res["reasoning"],
        root_cause_hypothesis=root_cause["hypothesis"],
        root_cause_confidence=root_cause["confidence"],
        root_cause_evidence=json.dumps(root_cause["evidence"]),
        duplicate_matches=json.dumps(duplicates["matches"]),
        remediation_summary=remediation["summary"],
        remediation_steps=json.dumps(remediation["steps"]),
        remediation_confidence=remediation["confidence"],
        remediation_sources=json.dumps(remediation["sources"]),
        combined_json=json.dumps({"bug_id": bug.bug_id, **combined}),
    )
    db.add(record)
    db.commit()
    db.refresh(record)
    return record


def _serialize_analysis(record: AgentAnalysis, bug_id: str) -> Dict[str, Any]:
    return {
        "bug_id": bug_id,
        "triage": {
            "severity": record.triage_severity,
            "priority": record.triage_priority,
            "type": record.triage_type,
            "component": record.triage_component,
            "confidence": record.triage_confidence,
            "reasoning": record.triage_reasoning,
        },
        "log_analysis": {
            "exception_type": record.log_exception_type,
            "message": record.log_message,
            "language": record.log_language,
            "failure_file": record.log_failure_file,
            "failure_line": record.log_failure_line,
            "affected_paths": json.loads(record.log_affected_paths or "[]"),
            "error_signature": record.log_error_signature,
            "confidence": record.log_confidence,
            "reasoning": record.log_reasoning,
        },
        "root_cause": {
            "hypothesis": record.root_cause_hypothesis,
            "confidence": record.root_cause_confidence,
            "evidence": json.loads(record.root_cause_evidence or "[]"),
        },
        "duplicates": {
            "matches": json.loads(record.duplicate_matches or "[]"),
        },
        "remediation": {
            "summary": record.remediation_summary,
            "steps": json.loads(record.remediation_steps or "[]"),
            "confidence": record.remediation_confidence,
            "sources": json.loads(record.remediation_sources or "[]"),
        },
        "analyzed_at": record.created_at,
    }


FRONTEND_HTML = '''
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8" />
<meta name="viewport" content="width=device-width, initial-scale=1.0" />
<title>Bug Submission — AI Smart Bug Analyzer &amp; Fix Advisor</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;500;600&family=IBM+Plex+Sans:wght@400;500;600&display=swap');

  :root{
    --ink:#000000;
    --panel:#181C25;
    --panel-2:#1F2430;
    --panel-3:#242A38;
    --line:#2B303C;
    --line-soft:#232833;
    --text:#E7EAF0;
    --muted:#8992A6;
    --muted-dim:#5C6478;
    --accent:#57C7D4;
    --accent-dim:rgba(87,199,212,0.12);
    --accent-line:rgba(87,199,212,0.35);
    --danger:#E2574C;
    --danger-dim:rgba(226,87,76,0.12);
    --warning:#E0A438;
    --warning-dim:rgba(224,164,56,0.12);
    --success:#48B884;
    --success-dim:rgba(72,184,132,0.12);
    --purple:#9C8CF0;
    --purple-dim:rgba(156,140,240,0.12);
  }

  *{box-sizing:border-box;}
  html,body{margin:0;padding:0;}
  body{
    background:var(--ink);
    color:var(--text);
    font-family:'IBM Plex Sans', sans-serif;
    font-size:15px;
    line-height:1.6;
    -webkit-font-smoothing:antialiased;
  }
  .mono{font-family:'IBM Plex Mono', monospace;}
  .wrap{ max-width:1100px; margin:0 auto; padding:40px 24px 64px; }

  .eyebrow{
    font-family:'IBM Plex Mono', monospace; font-size:11px; letter-spacing:0.14em;
    text-transform:uppercase; color:var(--muted-dim); margin:0 0 10px;
  }
  .header-row{ display:flex; align-items:flex-end; justify-content:space-between; gap:16px; margin-bottom:6px; flex-wrap:wrap; }
  h1{ font-family:'IBM Plex Mono', monospace; font-size:26px; font-weight:600; letter-spacing:-0.01em; margin:0; }
  .subtitle{ color:var(--muted); font-size:14.5px; margin:10px 0 32px; max-width:640px; }
  .btn-clear{
    background:transparent; border:1px solid var(--line); color:var(--muted);
    font-family:'IBM Plex Mono', monospace; font-size:12px; letter-spacing:0.04em;
    padding:8px 14px; border-radius:6px; cursor:pointer; transition:border-color .15s ease, color .15s ease;
  }
  .btn-clear:hover{ border-color:var(--muted-dim); color:var(--text); }

  .grid{ display:grid; grid-template-columns:1.55fr 1fr; gap:20px; align-items:start; }
  @media (max-width:820px){ .grid{ grid-template-columns:1fr; } }

  .panel{ background:var(--panel); border:1px solid var(--line); border-radius:10px; padding:22px; }
  .label{
    font-family:'IBM Plex Mono', monospace; font-size:11px; letter-spacing:0.1em;
    text-transform:uppercase; color:var(--muted-dim); display:block; margin-bottom:8px;
  }

  input[type="text"], select, textarea{
    width:100%; background:var(--panel-2); border:1px solid var(--line); color:var(--text);
    border-radius:7px; padding:11px 13px; font-family:'IBM Plex Sans', sans-serif; font-size:14px;
    outline:none; transition:border-color .15s ease, box-shadow .15s ease;
  }
  input[type="text"]:focus, select:focus, textarea:focus{
    border-color:var(--accent-line); box-shadow:0 0 0 3px var(--accent-dim);
  }
  textarea{
    font-family:'IBM Plex Mono', monospace; font-size:13px; line-height:1.65;
    min-height:190px; resize:vertical; white-space:pre;
  }
  textarea::placeholder, input::placeholder{ color:var(--muted-dim); }

  .field{ margin-bottom:20px; }
  .field-row{ display:flex; gap:14px; }
  .field-row .field{ flex:1; margin-bottom:0; }

  hr.divider{ border:none; border-top:1px solid var(--line-soft); margin:22px 0; }

  .dropzone{
    border:1.5px dashed var(--line); border-radius:9px; padding:26px 20px; text-align:center;
    cursor:pointer; transition:border-color .15s ease, background .15s ease;
  }
  .dropzone.drag{ border-color:var(--accent); background:var(--accent-dim); }
  .dropzone .dz-icon{ font-family:'IBM Plex Mono', monospace; font-size:20px; color:var(--muted-dim); margin-bottom:8px; }
  .dropzone .dz-title{ font-size:14px; color:var(--text); margin-bottom:4px; }
  .dropzone .dz-sub{ font-size:12.5px; color:var(--muted-dim); font-family:'IBM Plex Mono', monospace; }
  input[type="file"]{ display:none; }

  .file-list{ margin-top:14px; display:flex; flex-direction:column; gap:8px; }
  .file-chip{
    display:flex; align-items:center; gap:10px; background:var(--panel-2); border:1px solid var(--line);
    border-radius:7px; padding:9px 11px; font-size:13px;
  }
  .file-chip.file-error{ border-color:rgba(226,87,76,0.4); }
  .file-chip .fc-ext{
    font-family:'IBM Plex Mono', monospace; font-size:10px; letter-spacing:0.04em;
    background:var(--panel-3); color:var(--muted); border-radius:4px; padding:3px 6px; flex-shrink:0;
  }
  .file-chip .fc-name{ flex:1; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }
  .file-chip .fc-meta{ color:var(--muted-dim); font-size:11.5px; font-family:'IBM Plex Mono', monospace; flex-shrink:0; }
  .file-chip .fc-error-text{ color:var(--danger); font-size:11.5px; font-family:'IBM Plex Mono', monospace; flex-shrink:0; }
  .file-chip .fc-remove{
    background:none; border:none; color:var(--muted-dim); cursor:pointer;
    font-size:15px; line-height:1; padding:2px 4px; border-radius:4px;
  }
  .file-chip .fc-remove:hover{ color:var(--danger); background:var(--danger-dim); }

  .submit-btn{
    width:100%; margin-top:22px; background:var(--accent); color:#0A2C31; border:none; border-radius:8px;
    padding:13px; font-family:'IBM Plex Mono', monospace; font-size:13.5px; font-weight:600;
    letter-spacing:0.02em; cursor:pointer; transition:filter .15s ease, opacity .15s ease;
  }
  .submit-btn:hover{ filter:brightness(1.08); }
  .submit-btn:disabled{ opacity:0.45; cursor:not-allowed; filter:none; }

  .hint{ font-size:12px; color:var(--muted-dim); margin-top:9px; text-align:center; }
  .hint.error-hint{ color:var(--danger); }

  .preview-panel{ position:sticky; top:20px; }
  .preview-head{ display:flex; align-items:center; justify-content:space-between; margin-bottom:16px; }
  .status-pill{
    display:inline-flex; align-items:center; gap:7px; font-family:'IBM Plex Mono', monospace;
    font-size:11px; letter-spacing:0.05em; text-transform:uppercase; color:var(--muted-dim);
  }
  .status-dot{ width:7px; height:7px; border-radius:50%; background:var(--muted-dim); }
  .status-dot.live{ background:var(--accent); animation:pulse 1.6s ease-in-out infinite; }
  @keyframes pulse{ 0%,100%{ box-shadow:0 0 0 0 var(--accent-dim); } 50%{ box-shadow:0 0 0 5px transparent; } }

  .empty-state{ padding:30px 6px; text-align:center; color:var(--muted-dim); font-size:13px; }
  .empty-state .ee-icon{ font-family:'IBM Plex Mono', monospace; font-size:22px; margin-bottom:10px; opacity:0.6; }

  .badge{
    display:inline-block; font-family:'IBM Plex Mono', monospace; font-size:11.5px; letter-spacing:0.03em;
    padding:5px 10px; border-radius:5px; background:var(--accent-dim); color:var(--accent);
    border:1px solid var(--accent-line); margin-bottom:16px;
  }

  .kv{ margin-bottom:14px; }
  .kv-label{
    font-family:'IBM Plex Mono', monospace; font-size:10.5px; text-transform:uppercase;
    letter-spacing:0.08em; color:var(--muted-dim); margin-bottom:4px;
  }
  .kv-value{ font-family:'IBM Plex Mono', monospace; font-size:13px; color:var(--text); word-break:break-word; line-height:1.5; }
  .kv-value.danger-text{ color:var(--danger); }

  .preview-fade{ animation:fadeIn .25s ease; }
  @keyframes fadeIn{ from{ opacity:0; transform:translateY(3px); } to{ opacity:1; transform:translateY(0); } }

  .attached-mini{ margin-top:6px; display:flex; flex-direction:column; gap:6px; }
  .attached-mini-item{
    display:flex; justify-content:space-between; font-family:'IBM Plex Mono', monospace; font-size:11.5px;
    color:var(--muted); border-top:1px solid var(--line-soft); padding-top:6px;
  }

  .confirm-banner{
    margin-top:20px; background:var(--success-dim); border:1px solid rgba(72,184,132,0.35);
    border-radius:9px; padding:15px 16px; display:none;
  }
  .confirm-banner.show{ display:block; animation:fadeIn .25s ease; }
  .confirm-banner.error-banner{ background:var(--danger-dim); border-color:rgba(226,87,76,0.35); }
  .confirm-title{ font-family:'IBM Plex Mono', monospace; font-size:13px; color:var(--success); margin-bottom:4px; }
  .confirm-banner.error-banner .confirm-title{ color:var(--danger); }
  .confirm-sub{ font-size:12.5px; color:var(--muted); }

  .history{ margin-top:36px; }
  .history-head{ display:flex; align-items:center; justify-content:space-between; margin-bottom:14px; }
  .history-title{
    font-family:'IBM Plex Mono', monospace; font-size:11px; letter-spacing:0.1em;
    text-transform:uppercase; color:var(--muted-dim);
  }
  .btn-refresh{
    background:transparent; border:1px solid var(--line); color:var(--muted);
    font-family:'IBM Plex Mono', monospace; font-size:11px; letter-spacing:0.04em;
    padding:5px 10px; border-radius:5px; cursor:pointer;
  }
  .btn-refresh:hover{ border-color:var(--muted-dim); color:var(--text); }
  .history-empty{ color:var(--muted-dim); font-size:13px; padding:16px 0; border-top:1px solid var(--line-soft); }
  .history-item{
    display:flex; align-items:center; gap:14px; padding:12px 0; border-top:1px solid var(--line-soft); cursor:pointer;
  }
  .history-item:hover .history-title-txt{ color:var(--accent); }
  .history-id{
    font-family:'IBM Plex Mono', monospace; font-size:12px; color:var(--accent); flex-shrink:0;
    width:150px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap;
  }
  .history-title-txt{ flex:1; font-size:13.5px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; transition:color .15s ease; }
  .history-sev{
    font-family:'IBM Plex Mono', monospace; font-size:10.5px; text-transform:uppercase;
    letter-spacing:0.04em; padding:3px 8px; border-radius:5px; flex-shrink:0;
  }
  .history-time{ font-family:'IBM Plex Mono', monospace; font-size:11.5px; color:var(--muted-dim); flex-shrink:0; }

  .modal{ position:fixed; inset:0; background:rgba(0,0,0,0.6); display:flex; align-items:center; justify-content:center; padding:20px; z-index:50; }
  .modal.hidden{ display:none; }
  .modal-content{
    background:var(--panel); border:1px solid var(--line); border-radius:10px; max-width:820px; width:100%;
    max-height:86vh; overflow-y:auto; padding:24px; position:relative;
  }
  .modal-close{ position:absolute; top:12px; right:16px; background:none; border:none; color:var(--muted-dim); font-size:1.4rem; cursor:pointer; }
  .modal-meta{ color:var(--muted); font-size:12.5px; font-family:'IBM Plex Mono', monospace; }
  .modal-body{
    white-space:pre-wrap; word-break:break-word; background:var(--panel-2); border:1px solid var(--line);
    border-radius:8px; padding:14px; font-family:'IBM Plex Mono', monospace; font-size:12.5px; margin-top:12px;
    max-height:220px; overflow-y:auto;
  }

  /* --- Structured Findings Dashboard --- */
  .dash-section{ margin-top:20px; border-top:1px solid var(--line-soft); padding-top:18px; }
  .dash-section-head{ display:flex; align-items:center; gap:10px; margin-bottom:12px; }
  .dash-icon{
    width:22px; height:22px; border-radius:5px; display:flex; align-items:center; justify-content:center;
    font-family:'IBM Plex Mono', monospace; font-size:11px; flex-shrink:0;
  }
  .dash-section-title{
    font-family:'IBM Plex Mono', monospace; font-size:11.5px; letter-spacing:0.09em; text-transform:uppercase;
  }
  .dash-icon.triage{ background:var(--accent-dim); color:var(--accent); }
  .dash-section-title.triage{ color:var(--accent); }
  .dash-icon.log{ background:var(--warning-dim); color:var(--warning); }
  .dash-section-title.log{ color:var(--warning); }
  .dash-icon.rootcause{ background:var(--purple-dim); color:var(--purple); }
  .dash-section-title.rootcause{ color:var(--purple); }
  .dash-icon.dup{ background:rgba(232,135,63,0.14); color:#E8873F; }
  .dash-section-title.dup{ color:#E8873F; }
  .dash-icon.remediation{ background:var(--success-dim); color:var(--success); }
  .dash-section-title.remediation{ color:var(--success); }

  .dash-grid{ display:grid; grid-template-columns:1fr 1fr; gap:10px 18px; margin-bottom:10px; }
  .dash-kv-label{ font-family:'IBM Plex Mono', monospace; font-size:10.5px; text-transform:uppercase; letter-spacing:0.06em; color:var(--muted-dim); }
  .dash-kv-value{ font-family:'IBM Plex Mono', monospace; font-size:12.5px; color:var(--text); word-break:break-word; }
  .dash-text{ font-size:13px; color:var(--text); line-height:1.55; margin-bottom:8px; }
  .dash-reasoning{ font-size:12px; color:var(--muted); line-height:1.5; margin-top:6px; font-style:italic; }

  .confidence-bar{ height:5px; border-radius:3px; background:var(--panel-3); margin-top:4px; overflow:hidden; }
  .confidence-fill{ height:100%; background:var(--accent); }
  .confidence-fill.purple{ background:var(--purple); }
  .confidence-fill.success{ background:var(--success); }
  .confidence-fill.warning{ background:var(--warning); }

  .evidence-card, .dup-card{
    background:var(--panel-2); border:1px solid var(--line); border-radius:8px; padding:11px 13px; margin-top:8px;
  }
  .evidence-card-head, .dup-card-head{ display:flex; justify-content:space-between; align-items:center; gap:10px; margin-bottom:5px; }
  .evidence-id, .dup-id{ font-family:'IBM Plex Mono', monospace; font-size:12px; color:var(--purple); }
  .dup-card .dup-id{ color:#E8873F; }
  .sim-pill{
    font-family:'IBM Plex Mono', monospace; font-size:10.5px; padding:2px 8px; border-radius:10px;
    background:var(--panel-3); color:var(--muted);
  }
  .sim-pill.high{ background:var(--success-dim); color:var(--success); }
  .sim-pill.med{ background:var(--warning-dim); color:var(--warning); }
  .evidence-title, .dup-title{ font-size:12.5px; color:var(--text); margin-bottom:4px; }
  .evidence-cause, .dup-resolution{ font-size:12px; color:var(--muted); line-height:1.5; }
  .dup-source-tag{
    font-family:'IBM Plex Mono', monospace; font-size:9.5px; text-transform:uppercase; letter-spacing:0.05em;
    color:var(--muted-dim); border:1px solid var(--line); border-radius:4px; padding:1px 6px; margin-right:6px;
  }

  .remediation-summary{
    background:var(--success-dim); border:1px solid rgba(72,184,132,0.3); border-radius:8px;
    padding:11px 13px; font-size:12.5px; color:var(--text); line-height:1.5; margin-bottom:10px;
  }
  .remediation-steps{ margin:0; padding-left:18px; }
  .remediation-steps li{ font-size:13px; color:var(--text); line-height:1.6; margin-bottom:6px; }
  .remediation-sources{ margin-top:10px; display:flex; flex-wrap:wrap; gap:6px; }
  .source-chip{
    font-family:'IBM Plex Mono', monospace; font-size:10px; color:var(--muted); background:var(--panel-3);
    border:1px solid var(--line); border-radius:4px; padding:3px 7px;
  }
  .dash-empty{ font-size:12.5px; color:var(--muted-dim); font-style:italic; }
</style>
</head>
<body>
<div class="wrap">

  <p class="eyebrow">AI Smart Bug Analyzer &amp; Fix Advisor</p>
  <div class="header-row">
    <h1>Bug submission</h1>
    <button class="btn-clear" id="clearBtn" type="button">Clear form</button>
  </div>
  <p class="subtitle">Paste an error, or drop a log file. Triage, Log Analysis, Root Cause, Duplicate Detection, and Remediation agents run automatically on submit and are shown together in the Structured Findings dashboard.</p>

  <div class="grid">

    <div>
      <div class="panel">

        <div class="field-row">
          <div class="field">
            <label class="label" for="title">Title</label>
            <input type="text" id="title" placeholder="Checkout fails with 500 on submit" />
          </div>
        </div>

        <hr class="divider" />

        <div class="field">
          <label class="label" for="pasteArea">Paste report, stack trace, or log</label>
          <textarea id="pasteArea" placeholder="Traceback (most recent call last):&#10;  File &quot;app/views.py&quot;, line 88, in checkout&#10;    charge = stripe.Charge.create(...)&#10;stripe.error.CardError: Your card was declined."></textarea>
        </div>

        <div class="field">
          <label class="label">Attach files</label>
          <div class="dropzone" id="dropzone">
            <div class="dz-icon">&#8595;</div>
            <div class="dz-title">Drop files here, or click to browse</div>
            <div class="dz-sub">.log .txt .json .xml .csv .stacktrace .out .pdf .docx — up to 10 files, 5MB each</div>
          </div>
          <input type="file" id="fileInput" multiple accept=".txt,.log,.json,.xml,.csv,.stacktrace,.out,.pdf,.docx" />
          <div class="file-list" id="fileList"></div>
        </div>

        <button class="submit-btn" id="submitBtn" disabled>Submit for analysis</button>
        <div class="hint" id="submitHint">Add pasted text or at least one file to enable submission.</div>

        <div class="confirm-banner" id="confirmBanner">
          <div class="confirm-title" id="confirmTitle"></div>
          <div class="confirm-sub" id="confirmSub"></div>
        </div>

      </div>

      <div class="history">
        <div class="history-head">
          <div class="history-title">Submission history</div>
          <button class="btn-refresh" id="refreshHistoryBtn" type="button">Refresh</button>
        </div>
        <div id="historyList">
          <div class="history-empty">Loading submissions...</div>
        </div>
      </div>
    </div>

    <div class="preview-panel">
      <div class="panel">
        <div class="preview-head">
          <span class="label" style="margin-bottom:0;">Structured findings</span>
          <span class="status-pill"><span class="status-dot" id="statusDot"></span><span id="statusText">Idle</span></span>
        </div>
        <div id="previewBody">
          <div class="empty-state">
            <div class="ee-icon">&#10022;</div>
            Paste an error or drop a file to begin analysis.
          </div>
        </div>
      </div>
    </div>

  </div>
</div>

<div id="detailModal" class="modal hidden">
  <div class="modal-content">
    <button id="closeModal" class="modal-close" type="button">&times;</button>
    <h3 id="modalTitle" class="mono" style="margin:0 0 6px;"></h3>
    <p id="modalMeta" class="modal-meta"></p>
    <pre id="modalContent" class="modal-body"></pre>
    <div id="modalAgentSection"></div>
  </div>
</div>

<script>
(function(){
  const API_BASE = "";

  const pasteArea = document.getElementById('pasteArea');
  const dropzone = document.getElementById('dropzone');
  const fileInput = document.getElementById('fileInput');
  const fileList = document.getElementById('fileList');
  const submitBtn = document.getElementById('submitBtn');
  const submitHint = document.getElementById('submitHint');
  const previewBody = document.getElementById('previewBody');
  const statusDot = document.getElementById('statusDot');
  const statusText = document.getElementById('statusText');
  const confirmBanner = document.getElementById('confirmBanner');
  const confirmTitle = document.getElementById('confirmTitle');
  const confirmSub = document.getElementById('confirmSub');
  const historyList = document.getElementById('historyList');
  const titleInput = document.getElementById('title');
  const clearBtn = document.getElementById('clearBtn');
  const refreshHistoryBtn = document.getElementById('refreshHistoryBtn');

  const MAX_FILES = 10;
  const MAX_FILE_SIZE = 5 * 1024 * 1024;
  const ALLOWED_EXT = ['txt','log','json','xml','csv','stacktrace','out','pdf','docx'];
  const BINARY_EXT = ['pdf','docx'];

  let attachedFiles = [];
  let isSubmitting = false;

  function extForFile(name){
    const parts = name.split('.');
    return parts.length > 1 ? parts[parts.length-1].toUpperCase() : 'FILE';
  }
  function extLower(name){
    const parts = name.split('.');
    return parts.length > 1 ? parts[parts.length-1].toLowerCase() : '';
  }
  function formatSize(bytes){
    if (bytes < 1024) return bytes + ' B';
    if (bytes < 1024*1024) return (bytes/1024).toFixed(1) + ' KB';
    return (bytes/(1024*1024)).toFixed(1) + ' MB';
  }
  function escapeHtml(str){
    return String(str).replace(/[&<>"']/g, function(c){
      return { '&':'&amp;', '<':'&lt;', '>':'&gt;', '"':'&quot;', "'":'&#39;' }[c];
    });
  }

  function renderFileList(){
    fileList.innerHTML = '';
    attachedFiles.forEach(function(f, idx){
      const chip = document.createElement('div');
      chip.className = 'file-chip' + (f.error ? ' file-error' : '');
      chip.innerHTML =
        '<span class="fc-ext">' + extForFile(f.name) + '</span>' +
        '<span class="fc-name">' + escapeHtml(f.name) + '</span>' +
        (f.error
          ? '<span class="fc-error-text">' + escapeHtml(f.error) + '</span>'
          : '<span class="fc-meta">' + formatSize(f.size) + '</span>') +
        '<button class="fc-remove" type="button" aria-label="Remove file">&#10005;</button>';
      chip.querySelector('.fc-remove').addEventListener('click', function(){
        attachedFiles.splice(idx, 1);
        renderFileList();
        updatePreview();
      });
      fileList.appendChild(chip);
    });
  }

  function updatePreview(){
    const text = pasteArea.value;
    const hasFiles = attachedFiles.length > 0;
    const hasValidFiles = attachedFiles.some(function(f){ return !f.error; });
    const hasContent = (text && text.trim().length > 0) || hasValidFiles;

    submitBtn.disabled = !hasContent || isSubmitting;
    if (!isSubmitting) {
      submitHint.textContent = hasContent ? 'Ready to submit.' : 'Add pasted text or at least one file to enable submission.';
      submitHint.classList.remove('error-hint');
    }

    if (!hasContent) {
      statusDot.classList.remove('live');
      statusText.textContent = 'Idle';
      previewBody.innerHTML = '<div class="empty-state"><div class="ee-icon">&#10022;</div>Paste an error or drop a file to begin analysis.</div>';
      return;
    }

    statusDot.classList.add('live');
    statusText.textContent = 'Ready';
    previewBody.innerHTML = '<div class="empty-state"><div class="ee-icon">&#10022;</div>Submit to run the full agent pipeline (Triage → Log Analysis → Root Cause → Duplicate Detection → Remediation).</div>';
  }

  pasteArea.addEventListener('input', updatePreview);

  dropzone.addEventListener('click', function(){ fileInput.click(); });
  dropzone.addEventListener('dragover', function(e){ e.preventDefault(); dropzone.classList.add('drag'); });
  dropzone.addEventListener('dragleave', function(){ dropzone.classList.remove('drag'); });
  dropzone.addEventListener('drop', function(e){
    e.preventDefault();
    dropzone.classList.remove('drag');
    handleFiles(e.dataTransfer.files);
  });
  fileInput.addEventListener('change', function(e){
    handleFiles(e.target.files);
    fileInput.value = '';
  });

  function handleFiles(fileListObj){
    const incoming = Array.from(fileListObj);

    if (attachedFiles.length + incoming.length > MAX_FILES) {
      submitHint.textContent = 'Max ' + MAX_FILES + ' files per submission. Remove some files first.';
      submitHint.classList.add('error-hint');
      return;
    }

    incoming.forEach(function(file){
      const ext = extLower(file.name);
      const isBinary = BINARY_EXT.includes(ext);
      const entry = { file: file, name: file.name, size: file.size, content: null, error: null, binary: isBinary };

      if (!ALLOWED_EXT.includes(ext)) {
        entry.error = 'Unsupported type (.' + ext + ')';
      } else if (file.size === 0) {
        entry.error = 'File is empty';
      } else if (file.size > MAX_FILE_SIZE) {
        entry.error = 'Too large (max 5MB)';
      }

      attachedFiles.push(entry);
    });
    renderFileList();
    updatePreview();
  }

  submitBtn.addEventListener('click', async function(){
    if (submitBtn.disabled || isSubmitting) return;

    const title = titleInput.value.trim();
    const text = pasteArea.value.trim();
    const validFiles = attachedFiles.filter(function(f){ return !f.error; });

    if (!title) {
      submitHint.textContent = 'Title is required.';
      submitHint.classList.add('error-hint');
      titleInput.focus();
      return;
    }
    if (attachedFiles.some(function(f){ return f.error; })) {
      submitHint.textContent = 'Remove or fix invalid files before submitting.';
      submitHint.classList.add('error-hint');
      return;
    }
    if (!text && validFiles.length === 0) {
      submitHint.textContent = 'Add pasted text or at least one file to enable submission.';
      submitHint.classList.add('error-hint');
      return;
    }

    isSubmitting = true;
    submitBtn.disabled = true;
    submitBtn.textContent = 'Submitting...';
    statusText.textContent = 'Analyzing';
    hideConfirm();

    try {
      const results = [];

      if (text) {
        const fd = new FormData();
        fd.append('title', title);
        fd.append('content', text);
        const res = await fetch(API_BASE + '/api/bugs/paste', { method: 'POST', body: fd });
        const data = await res.json();
        if (!res.ok) throw new Error(data.detail || 'Paste submission failed.');
        results.push(data);
      }

      if (validFiles.length > 0) {
        const fd = new FormData();
        fd.append('title', title);
        validFiles.forEach(function(f){ fd.append('files', f.file); });
        const res = await fetch(API_BASE + '/api/bugs/upload', { method: 'POST', body: fd });
        const data = await res.json();
        if (!res.ok) throw new Error(data.detail || 'File upload failed.');
        results.push(data);
      }

      const firstAnalysis = results.find(function(r){ return r.analysis; });
      const bugIds = results.map(function(r){ return r.bug_id; }).join(', ');

      if (firstAnalysis && firstAnalysis.analysis) {
        showConfirm(
          false,
          bugIds + ' analyzed',
          'Severity: ' + firstAnalysis.analysis.triage.severity + ' · ' + firstAnalysis.analysis.triage.component +
          ' · Remediation confidence: ' + Math.round(firstAnalysis.analysis.remediation.confidence*100) + '%'
        );
        statusDot.classList.add('live');
        statusText.textContent = 'Analyzed';
        previewBody.innerHTML =
          '<div class="badge preview-fade">' + bugIds + ' — full analysis complete</div>' +
          renderDashboard(firstAnalysis.analysis);
        resetForm(true);
      } else {
        showConfirm(false, results.length + ' submission(s) analyzed', bugIds);
        resetForm(false);
      }

      await loadHistory();

    } catch (err) {
      showConfirm(true, 'Submission failed', err.message);
    } finally {
      isSubmitting = false;
      submitBtn.textContent = 'Submit for analysis';
      updatePreview();
    }
  });

  function showConfirm(isError, title, sub){
    confirmBanner.classList.toggle('error-banner', isError);
    confirmTitle.textContent = title;
    confirmSub.textContent = sub;
    confirmBanner.classList.add('show');
    setTimeout(hideConfirm, 5000);
  }
  function hideConfirm(){
    confirmBanner.classList.remove('show');
  }

  function resetForm(keepPreview){
    pasteArea.value = '';
    attachedFiles = [];
    titleInput.value = '';
    renderFileList();
    if (!keepPreview) {
      updatePreview();
    } else {
      submitBtn.disabled = true;
      submitHint.textContent = 'Add pasted text or at least one file to enable submission.';
      submitHint.classList.remove('error-hint');
    }
  }

  clearBtn.addEventListener('click', function(){
    resetForm();
    submitHint.textContent = 'Add pasted text or at least one file to enable submission.';
    submitHint.classList.remove('error-hint');
  });

  const sevColors = {
    low: { bg: 'var(--success-dim)', text: 'var(--success)' },
    medium: { bg: 'var(--warning-dim)', text: 'var(--warning)' },
    high: { bg: 'rgba(232,135,63,0.12)', text: '#E8873F' },
    critical: { bg: 'var(--danger-dim)', text: 'var(--danger)' }
  };

  async function loadHistory(){
    try {
      const res = await fetch(API_BASE + '/api/bugs');
      const bugs = await res.json();
      if (!res.ok) throw new Error(bugs.detail || 'Failed to load history.');

      if (!bugs.length) {
        historyList.innerHTML = '<div class="history-empty">Nothing submitted yet. Your queued analyses will appear here.</div>';
        return;
      }

      historyList.innerHTML = '';
      bugs.forEach(function(bug){
        const c = sevColors[bug.severity] || sevColors.medium;
        const time = new Date(bug.created_at).toLocaleString();
        const item = document.createElement('div');
        item.className = 'history-item';
        item.innerHTML =
          '<span class="history-id">' + escapeHtml(bug.bug_id) + '</span>' +
          '<span class="history-title-txt">' + escapeHtml(bug.title) + '</span>' +
          '<span class="history-sev" style="background:' + c.bg + ';color:' + c.text + '">' + bug.severity + '</span>' +
          '<span class="history-time">' + time + '</span>';
        item.addEventListener('click', function(){ openDetail(bug.id); });
        historyList.appendChild(item);
      });
    } catch (err) {
      historyList.innerHTML = '<div class="history-empty">Failed to load history: ' + escapeHtml(err.message) + '</div>';
    }
  }

  refreshHistoryBtn.addEventListener('click', loadHistory);

  const modal = document.getElementById('detailModal');
  document.getElementById('closeModal').addEventListener('click', function(){ modal.classList.add('hidden'); });
  modal.addEventListener('click', function(e){ if (e.target === modal) modal.classList.add('hidden'); });

  function simClass(sim){
    if (sim >= 0.5) return 'high';
    if (sim >= 0.3) return 'med';
    return '';
  }

  function renderDashboard(analysis){
    if (!analysis) {
      return '<div class="dash-section"><div class="dash-empty">No agent analysis available for this submission.</div></div>';
    }
    const t = analysis.triage;
    const l = analysis.log_analysis;
    const rc = analysis.root_cause;
    const dup = analysis.duplicates;
    const rem = analysis.remediation;

    let html = '';

    // --- Triage ---
    html += '<div class="dash-section"><div class="dash-section-head"><div class="dash-icon triage">T</div><div class="dash-section-title triage">Triage</div></div>';
    html += '<div class="dash-grid">';
    html += '<div><div class="dash-kv-label">Severity</div><div class="dash-kv-value">' + escapeHtml(t.severity) + '</div></div>';
    html += '<div><div class="dash-kv-label">Priority</div><div class="dash-kv-value">' + escapeHtml(t.priority) + '</div></div>';
    html += '<div><div class="dash-kv-label">Component</div><div class="dash-kv-value">' + escapeHtml(t.component) + '</div></div>';
    html += '<div><div class="dash-kv-label">Type</div><div class="dash-kv-value">' + escapeHtml(t.type) + '</div></div>';
    html += '</div>';
    html += '<div class="dash-kv-label">Confidence — ' + Math.round(t.confidence*100) + '%</div>';
    html += '<div class="confidence-bar"><div class="confidence-fill" style="width:' + Math.round(t.confidence*100) + '%"></div></div>';
    html += '<div class="dash-reasoning">' + escapeHtml(t.reasoning) + '</div>';
    html += '</div>';

    // --- Log Analysis ---
    html += '<div class="dash-section"><div class="dash-section-head"><div class="dash-icon log">L</div><div class="dash-section-title log">Log Analysis</div></div>';
    html += '<div class="dash-grid">';
    html += '<div><div class="dash-kv-label">Exception</div><div class="dash-kv-value">' + escapeHtml(l.exception_type || '—') + '</div></div>';
    html += '<div><div class="dash-kv-label">Language</div><div class="dash-kv-value">' + escapeHtml(l.language || '—') + '</div></div>';
    html += '<div><div class="dash-kv-label">Failure point</div><div class="dash-kv-value">' + escapeHtml((l.failure_file || '—') + (l.failure_line ? (':' + l.failure_line) : '')) + '</div></div>';
    html += '<div><div class="dash-kv-label">Error signature</div><div class="dash-kv-value">' + escapeHtml(l.error_signature || '—') + '</div></div>';
    html += '</div>';
    if (l.message) {
      html += '<div class="dash-kv-label">Message</div><div class="dash-kv-value" style="margin-bottom:8px;">' + escapeHtml(l.message) + '</div>';
    }
    if (l.affected_paths && l.affected_paths.length) {
      html += '<div class="dash-kv-label">Affected code path(s)</div><div class="dash-kv-value" style="margin-bottom:8px;">' + escapeHtml(l.affected_paths.join(', ')) + '</div>';
    }
    html += '<div class="dash-kv-label">Confidence — ' + Math.round(l.confidence*100) + '%</div>';
    html += '<div class="confidence-bar"><div class="confidence-fill" style="width:' + Math.round(l.confidence*100) + '%"></div></div>';
    html += '<div class="dash-reasoning">' + escapeHtml(l.reasoning) + '</div>';
    html += '</div>';

    // --- Root Cause ---
    html += '<div class="dash-section"><div class="dash-section-head"><div class="dash-icon rootcause">R</div><div class="dash-section-title rootcause">Root Cause Hypothesis</div></div>';
    html += '<div class="dash-text">' + escapeHtml(rc.hypothesis || 'No hypothesis generated.') + '</div>';
    html += '<div class="dash-kv-label">Confidence — ' + Math.round((rc.confidence||0)*100) + '%</div>';
    html += '<div class="confidence-bar"><div class="confidence-fill purple" style="width:' + Math.round((rc.confidence||0)*100) + '%"></div></div>';
    if (rc.evidence && rc.evidence.length) {
      html += '<div class="dash-kv-label" style="margin-top:12px;">Supporting historical evidence</div>';
      rc.evidence.forEach(function(e){
        html += '<div class="evidence-card">' +
          '<div class="evidence-card-head"><span class="evidence-id">' + escapeHtml(e.defect_id) + '</span>' +
          '<span class="sim-pill ' + simClass(e.similarity) + '">' + Math.round(e.similarity*100) + '% match</span></div>' +
          '<div class="evidence-title">' + escapeHtml(e.title) + '</div>' +
          '<div class="evidence-cause">' + escapeHtml(e.root_cause) + '</div>' +
        '</div>';
      });
    } else {
      html += '<div class="dash-empty" style="margin-top:8px;">No supporting historical evidence found.</div>';
    }
    html += '</div>';

    // --- Duplicate Detection ---
    html += '<div class="dash-section"><div class="dash-section-head"><div class="dash-icon dup">D</div><div class="dash-section-title dup">Duplicate / Similar Bugs</div></div>';
    if (dup.matches && dup.matches.length) {
      dup.matches.forEach(function(m){
        html += '<div class="dup-card">' +
          '<div class="dup-card-head"><span><span class="dup-source-tag">' + escapeHtml(m.source === 'historical_kb' ? 'KB' : 'Prior bug') + '</span><span class="dup-id">' + escapeHtml(m.id) + '</span></span>' +
          '<span class="sim-pill ' + simClass(m.similarity) + '">' + m.similarity_pct + '% similar</span></div>' +
          '<div class="dup-title">' + escapeHtml(m.title) + '</div>' +
          '<div class="dup-resolution">' + escapeHtml(m.resolution_summary || 'No resolution summary available.') + '</div>' +
        '</div>';
      });
    } else {
      html += '<div class="dash-empty">No duplicate or closely similar issues found.</div>';
    }
    html += '</div>';

    // --- Remediation ---
    html += '<div class="dash-section"><div class="dash-section-head"><div class="dash-icon remediation">F</div><div class="dash-section-title remediation">Recommended Fix</div></div>';
    html += '<div class="remediation-summary">' + escapeHtml(rem.summary || '') + '</div>';
    if (rem.steps && rem.steps.length) {
      html += '<ol class="remediation-steps">';
      rem.steps.forEach(function(s){ html += '<li>' + escapeHtml(s) + '</li>'; });
      html += '</ol>';
    }
    html += '<div class="dash-kv-label" style="margin-top:10px;">Confidence — ' + Math.round((rem.confidence||0)*100) + '%</div>';
    html += '<div class="confidence-bar"><div class="confidence-fill success" style="width:' + Math.round((rem.confidence||0)*100) + '%"></div></div>';
    if (rem.sources && rem.sources.length) {
      html += '<div class="remediation-sources">';
      rem.sources.forEach(function(s){ html += '<span class="source-chip">' + escapeHtml(s) + '</span>'; });
      html += '</div>';
    }
    html += '</div>';

    return html;
  }

  async function openDetail(id){
    try {
      const res = await fetch(API_BASE + '/api/bugs/' + id);
      const bug = await res.json();
      if (!res.ok) throw new Error(bug.detail || 'Failed to load detail.');
      document.getElementById('modalTitle').textContent = bug.bug_id + ' — ' + bug.title;
      document.getElementById('modalMeta').textContent =
        (bug.source_type === 'upload' ? 'File: ' + bug.original_filename : 'Source: Pasted text') +
        ' · ' + bug.severity + ' · ' + new Date(bug.created_at).toLocaleString() + ' · ' + bug.content_length + ' chars';
      document.getElementById('modalContent').textContent = bug.content;

      let analysis = null;
      try {
        const aRes = await fetch(API_BASE + '/api/bugs/' + id + '/analysis');
        if (aRes.ok) analysis = await aRes.json();
      } catch (e) { /* analysis not available */ }
      document.getElementById('modalAgentSection').innerHTML = renderDashboard(analysis);

      modal.classList.remove('hidden');
    } catch (err) {
      showConfirm(true, 'Could not load submission', err.message);
    }
  }

  updatePreview();
  loadHistory();
})();
</script>
</body>
</html>
'''


@app.get("/", response_class=HTMLResponse)
def serve_frontend():
    return FRONTEND_HTML


@app.on_event("startup")
def _on_startup():
    db = SessionLocal()
    try:
        _seed_historical_defects(db)
    finally:
        db.close()


@app.get("/api/health")
def health_check():
    return {"status": "ok", "time": datetime.utcnow().isoformat()}


@app.post("/api/bugs/paste", response_model=BugSubmissionCreatedResponse)
def submit_pasted_bug(
    title: str = Form(...),
    content: str = Form(...),
    db: Session = Depends(get_db),
):
    content = content.strip()
    if not content:
        raise HTTPException(status_code=400, detail="Pasted content cannot be empty.")
    if not title.strip():
        raise HTTPException(status_code=400, detail="Title cannot be empty.")

    combined = run_orchestration(content, db)
    sev = _validate_severity(combined["triage"]["severity"])

    bug_number = _next_bug_number(db)
    bug = BugSubmission(
        bug_number=bug_number,
        bug_id=_format_bug_id(bug_number),
        title=title.strip(),
        severity=sev,
        source_type=SourceType.PASTE,
        original_filename=None,
        file_type="text/plain",
        content=content,
        content_length=str(len(content)),
    )
    db.add(bug)
    db.commit()
    db.refresh(bug)

    record = _persist_analysis(db, bug, combined)

    return {
        "id": bug.id,
        "bug_id": bug.bug_id,
        "message": "Bug report submitted and analyzed successfully.",
        "analysis": _serialize_analysis(record, bug.bug_id),
    }


@app.post("/api/bugs/upload", response_model=BugSubmissionCreatedResponse)
async def submit_uploaded_bug(
    title: str = Form(...),
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
):
    if not title.strip():
        raise HTTPException(status_code=400, detail="Title cannot be empty.")
    if not files:
        raise HTTPException(status_code=400, detail="At least one file is required.")
    if len(files) > MAX_FILES_PER_UPLOAD:
        raise HTTPException(
            status_code=400,
            detail=f"Too many files. Max {MAX_FILES_PER_UPLOAD} files per submission.",
        )

    created_ids = []
    first_bug_id = None
    first_combined = None
    for file in files:
        ext = _validate_extension(file.filename)

        raw = await file.read()
        if len(raw) == 0:
            raise HTTPException(status_code=400, detail=f"'{file.filename}' is empty.")
        if len(raw) > MAX_FILE_SIZE_BYTES:
            max_mb = MAX_FILE_SIZE_BYTES // (1024 * 1024)
            raise HTTPException(
                status_code=400,
                detail=f"'{file.filename}' is too large. Max size is {max_mb} MB.",
            )

        text_content = _extract_text_content(file.filename, ext, raw)

        combined = run_orchestration(text_content, db)
        sev = _validate_severity(combined["triage"]["severity"])

        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S%f")
        safe_name = f"{timestamp}_{os.path.basename(file.filename)}"
        disk_path = os.path.join(UPLOAD_DIR, safe_name)
        with open(disk_path, "wb") as f:
            f.write(raw)

        bug_title = title.strip() if len(files) == 1 else f"{title.strip()} — {file.filename}"
        bug_number = _next_bug_number(db)
        bug = BugSubmission(
            bug_number=bug_number,
            bug_id=_format_bug_id(bug_number),
            title=bug_title,
            severity=sev,
            source_type=SourceType.UPLOAD,
            original_filename=file.filename,
            file_type=ext,
            content=text_content,
            content_length=str(len(text_content)),
        )
        db.add(bug)
        db.commit()
        db.refresh(bug)

        record = _persist_analysis(db, bug, combined)

        created_ids.append(bug.id)
        if first_bug_id is None:
            first_bug_id = bug.bug_id
            first_combined = _serialize_analysis(record, bug.bug_id)

    return {
        "id": created_ids[0],
        "bug_id": first_bug_id,
        "message": f"{len(created_ids)} file(s) uploaded, submitted, and analyzed successfully.",
        "analysis": first_combined,
    }


@app.get("/api/bugs", response_model=List[BugSubmissionResponse])
def list_bugs(db: Session = Depends(get_db)):
    bugs = db.query(BugSubmission).order_by(desc(BugSubmission.created_at)).all()
    return [_serialize(b) for b in bugs]


@app.get("/api/bugs/{bug_id}", response_model=BugSubmissionDetail)
def get_bug(bug_id: str, db: Session = Depends(get_db)):
    bug = (
        db.query(BugSubmission)
        .filter((BugSubmission.id == bug_id) | (BugSubmission.bug_id == bug_id))
        .first()
    )
    if not bug:
        raise HTTPException(status_code=404, detail="Bug submission not found.")
    data = _serialize(bug)
    data["content"] = bug.content
    return data


@app.get("/api/bugs/{bug_id}/analysis")
def get_bug_analysis(bug_id: str, db: Session = Depends(get_db)):
    """
    Returns the FULL combined agent output for a bug — Triage, Log Analysis,
    Root Cause, Duplicate Detection, and Remediation — in the exact
    structure the Structured Findings Dashboard renders.
    """
    bug = (
        db.query(BugSubmission)
        .filter((BugSubmission.id == bug_id) | (BugSubmission.bug_id == bug_id))
        .first()
    )
    if not bug:
        raise HTTPException(status_code=404, detail="Bug submission not found.")
    if not bug.analysis:
        raise HTTPException(status_code=404, detail="No agent analysis found for this bug.")
    return _serialize_analysis(bug.analysis, bug.bug_id)


@app.get("/api/historical-defects")
def list_historical_defects(db: Session = Depends(get_db)):
    """Utility endpoint to inspect the seeded knowledge base Root Cause /
    Duplicate Detection retrieve against."""
    defects = db.query(HistoricalDefect).all()
    return [
        {
            "defect_id": d.defect_id,
            "title": d.title,
            "component": d.component,
            "root_cause": d.root_cause,
            "resolution_summary": d.resolution_summary,
            "tags": d.tags,
        }
        for d in defects
    ]


@app.delete("/api/bugs/{bug_id}")
def delete_bug(bug_id: str, db: Session = Depends(get_db)):
    bug = (
        db.query(BugSubmission)
        .filter((BugSubmission.id == bug_id) | (BugSubmission.bug_id == bug_id))
        .first()
    )
    if not bug:
        raise HTTPException(status_code=404, detail="Bug submission not found.")
    db.delete(bug)
    db.commit()
    return {"message": "Deleted successfully."}


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import uvicorn

    port = int(os.getenv("PORT", "8000"))
    print(f"Starting AI Smart Bug Analyzer & Fix Advisor on http://localhost:{port}")
    uvicorn.run(app, host="0.0.0.0", port=port)