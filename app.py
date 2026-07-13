"""
AI Smart Bug Analyzer & Fix Advisor
Bug Submission Module — single-file version

Everything in one file, per company request: FastAPI backend, PostgreSQL
(via SQLAlchemy, with automatic local SQLite fallback), and the full
HTML/CSS/JS frontend served inline from the same process.

Tech stack covered here (per project tech stack doc):
  - Backend:  Python, FastAPI
  - Database: PostgreSQL (SQLAlchemy ORM; falls back to SQLite for local/dev
              runs if DATABASE_URL is not set)
  - Frontend: HTML, CSS, JavaScript (served inline by FastAPI, no build step)

Run with:
    pip install fastapi uvicorn sqlalchemy psycopg2-binary python-multipart pydantic pypdf python-docx
    uvicorn app:app --reload --port 8000

Then open http://127.0.0.1:8000 in your browser.

Configure PostgreSQL by setting the DATABASE_URL environment variable, e.g.:
    export DATABASE_URL=postgresql://username:password@localhost:5432/bug_analyzer
If unset, a local SQLite file (bug_analyzer.db) is used automatically so this
runs immediately without any database setup.
"""

import io
import os
import enum
import uuid
from datetime import datetime
from typing import List, Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from sqlalchemy import Column, String, Text, Integer, DateTime, Enum, create_engine, desc, func as sa_func
from sqlalchemy.orm import declarative_base, sessionmaker, Session
from sqlalchemy.sql import func
from pydantic import BaseModel

# Optional third-party parsers for PDF / DOCX text extraction.
# The app still runs without them, but .pdf / .docx uploads will be
# rejected with a clear error until the packages are installed.
try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None

try:
    import docx as docx_lib  # python-docx
except ImportError:  # pragma: no cover
    docx_lib = None


# ---------------------------------------------------------------------------
# Database (PostgreSQL via SQLAlchemy, SQLite fallback for zero-setup local runs)
# ---------------------------------------------------------------------------

DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./bug_analyzer.db")
connect_args = {"check_same_thread": False} if DATABASE_URL.startswith("sqlite") else {}

engine = create_engine(DATABASE_URL, connect_args=connect_args)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------

class SourceType(str, enum.Enum):
    PASTE = "paste"
    UPLOAD = "upload"


class Severity(str, enum.Enum):
    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"
    CRITICAL = "critical"


class BugSubmission(Base):
    __tablename__ = "bug_submissions"

    # Internal primary key (UUID) — used for API lookups / routing.
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))

    # Sequential numeric counter, used to build the human-readable bug_id.
    bug_number = Column(Integer, nullable=False)

    # Human-readable, sequential bug identifier, e.g. "BUG-00001".
    # Matches the `bug_id` field from the BUG_RECORD design.
    bug_id = Column(String(20), unique=True, index=True, nullable=False)

    title = Column(String(255), nullable=False)
    severity = Column(Enum(Severity), nullable=False, default=Severity.MEDIUM)
    source_type = Column(Enum(SourceType), nullable=False)
    original_filename = Column(String(255), nullable=True)
    file_type = Column(String(50), nullable=True)
    content = Column(Text, nullable=False)
    content_length = Column(String(20), nullable=True)
    created_at = Column(DateTime(timezone=True), server_default=func.now())


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------

class BugSubmissionResponse(BaseModel):
    id: str
    bug_id: str
    title: str
    severity: str
    source_type: str
    original_filename: Optional[str]
    file_type: Optional[str]
    content_length: Optional[str]
    created_at: datetime

    class Config:
        from_attributes = True


class BugSubmissionDetail(BugSubmissionResponse):
    content: str


class BugSubmissionCreatedResponse(BaseModel):
    id: str
    bug_id: str
    message: str


# ---------------------------------------------------------------------------
# App config
# ---------------------------------------------------------------------------

# Plain-text / log formats, read directly as UTF-8 text.
TEXT_EXTENSIONS = {".txt", ".log", ".json", ".xml", ".csv", ".stacktrace", ".out"}
# Rich-document formats, read via a dedicated parser to extract plain text.
DOCUMENT_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_EXTENSIONS = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB
MAX_FILES_PER_UPLOAD = 10
UPLOAD_DIR = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

BUG_ID_PREFIX = "BUG-"
BUG_ID_PAD = 5  # BUG-00001

Base.metadata.create_all(bind=engine)

app = FastAPI(
    title="AI Smart Bug Analyzer & Fix Advisor — Bug Submission Module",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _validate_severity(severity: str) -> Severity:
    try:
        return Severity(severity.lower())
    except ValueError:
        allowed = "low, medium, high, critical"
        raise HTTPException(
            status_code=400,
            detail=f"Invalid severity '{severity}'. Must be one of: {allowed}.",
        )


def _validate_extension(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed types: {allowed}",
        )
    return ext


def _extract_pdf_text(filename: str, raw: bytes) -> str:
    if PdfReader is None:
        raise HTTPException(
            status_code=500,
            detail="PDF support is not installed on the server. Run: pip install pypdf",
        )
    try:
        reader = PdfReader(io.BytesIO(raw))
        pages_text = []
        for page in reader.pages:
            pages_text.append(page.extract_text() or "")
        text = "\n".join(pages_text).strip()
    except Exception:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' could not be read as a valid PDF file.",
        )
    if not text:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' contains no extractable text (it may be a scanned/image-only PDF).",
        )
    return text


def _extract_docx_text(filename: str, raw: bytes) -> str:
    if docx_lib is None:
        raise HTTPException(
            status_code=500,
            detail="DOCX support is not installed on the server. Run: pip install python-docx",
        )
    try:
        document = docx_lib.Document(io.BytesIO(raw))
        paragraphs = [p.text for p in document.paragraphs]
        # Also pull text out of any tables in the document.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        paragraphs.append(cell.text)
        text = "\n".join(paragraphs).strip()
    except Exception:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' could not be read as a valid DOCX file.",
        )
    if not text:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' contains no extractable text.",
        )
    return text


def _extract_text_content(filename: str, ext: str, raw: bytes) -> str:
    """
    Returns the plain-text content to store for a given uploaded file,
    dispatching to the right reader based on its extension.
    """
    if ext == ".pdf":
        return _extract_pdf_text(filename, raw)
    if ext == ".docx":
        return _extract_docx_text(filename, raw)
    # Plain-text / log formats.
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        raise HTTPException(
            status_code=400, detail=f"'{filename}' must be UTF-8 encoded plain text."
        )


def _next_bug_number(db: Session) -> int:
    """
    Returns the next sequential bug number, based on the current max
    in the table. Safe enough for a single-process dev/demo app; for
    high-concurrency production use, swap this for a DB sequence or
    SELECT ... FOR UPDATE.
    """
    current_max = db.query(sa_func.max(BugSubmission.bug_number)).scalar()
    return (current_max or 0) + 1


def _format_bug_id(bug_number: int) -> str:
    return f"{BUG_ID_PREFIX}{bug_number:0{BUG_ID_PAD}d}"


def _serialize(bug: BugSubmission) -> dict:
    return {
        "id": bug.id,
        "bug_id": bug.bug_id,
        "title": bug.title,
        "severity": bug.severity.value if hasattr(bug.severity, "value") else bug.severity,
        "source_type": bug.source_type.value if hasattr(bug.source_type, "value") else bug.source_type,
        "original_filename": bug.original_filename,
        "file_type": bug.file_type,
        "content_length": bug.content_length,
        "created_at": bug.created_at,
    }


# ---------------------------------------------------------------------------
# Frontend (served inline — no separate static files/build step)
# ---------------------------------------------------------------------------

FRONTEND_HTML = '<!DOCTYPE html>\n<html lang="en">\n<head>\n<meta charset="UTF-8" />\n<meta name="viewport" content="width=device-width, initial-scale=1.0" />\n<title>Bug Submission — AI Smart Bug Analyzer &amp; Fix Advisor</title>\n<style>\n  @import url(\'https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;500;600&family=IBM+Plex+Sans:wght@400;500;600&display=swap\');\n\n  :root{\n    --ink:#000000;\n    --panel:#181C25;\n    --panel-2:#1F2430;\n    --panel-3:#242A38;\n    --line:#2B303C;\n    --line-soft:#232833;\n    --text:#E7EAF0;\n    --muted:#8992A6;\n    --muted-dim:#5C6478;\n    --accent:#57C7D4;\n    --accent-dim:rgba(87,199,212,0.12);\n    --accent-line:rgba(87,199,212,0.35);\n    --danger:#E2574C;\n    --danger-dim:rgba(226,87,76,0.12);\n    --warning:#E0A438;\n    --warning-dim:rgba(224,164,56,0.12);\n    --success:#48B884;\n    --success-dim:rgba(72,184,132,0.12);\n  }\n\n  *{box-sizing:border-box;}\n  html,body{margin:0;padding:0;}\n  body{\n    background:var(--ink);\n    color:var(--text);\n    font-family:\'IBM Plex Sans\', sans-serif;\n    font-size:15px;\n    line-height:1.6;\n    -webkit-font-smoothing:antialiased;\n  }\n  .mono{font-family:\'IBM Plex Mono\', monospace;}\n  .wrap{ max-width:1040px; margin:0 auto; padding:40px 24px 64px; }\n\n  .eyebrow{\n    font-family:\'IBM Plex Mono\', monospace; font-size:11px; letter-spacing:0.14em;\n    text-transform:uppercase; color:var(--muted-dim); margin:0 0 10px;\n  }\n  .header-row{ display:flex; align-items:flex-end; justify-content:space-between; gap:16px; margin-bottom:6px; flex-wrap:wrap; }\n  h1{ font-family:\'IBM Plex Mono\', monospace; font-size:26px; font-weight:600; letter-spacing:-0.01em; margin:0; }\n  .subtitle{ color:var(--muted); font-size:14.5px; margin:10px 0 32px; max-width:560px; }\n  .btn-clear{\n    background:transparent; border:1px solid var(--line); color:var(--muted);\n    font-family:\'IBM Plex Mono\', monospace; font-size:12px; letter-spacing:0.04em;\n    padding:8px 14px; border-radius:6px; cursor:pointer; transition:border-color .15s ease, color .15s ease;\n  }\n  .btn-clear:hover{ border-color:var(--muted-dim); color:var(--text); }\n\n  .grid{ display:grid; grid-template-columns:1.55fr 1fr; gap:20px; align-items:start; }\n  @media (max-width:820px){ .grid{ grid-template-columns:1fr; } }\n\n  .panel{ background:var(--panel); border:1px solid var(--line); border-radius:10px; padding:22px; }\n  .label{\n    font-family:\'IBM Plex Mono\', monospace; font-size:11px; letter-spacing:0.1em;\n    text-transform:uppercase; color:var(--muted-dim); display:block; margin-bottom:8px;\n  }\n\n  input[type="text"], select, textarea{\n    width:100%; background:var(--panel-2); border:1px solid var(--line); color:var(--text);\n    border-radius:7px; padding:11px 13px; font-family:\'IBM Plex Sans\', sans-serif; font-size:14px;\n    outline:none; transition:border-color .15s ease, box-shadow .15s ease;\n  }\n  input[type="text"]:focus, select:focus, textarea:focus{\n    border-color:var(--accent-line); box-shadow:0 0 0 3px var(--accent-dim);\n  }\n  textarea{\n    font-family:\'IBM Plex Mono\', monospace; font-size:13px; line-height:1.65;\n    min-height:190px; resize:vertical; white-space:pre;\n  }\n  textarea::placeholder, input::placeholder{ color:var(--muted-dim); }\n\n  .field{ margin-bottom:20px; }\n  .field-row{ display:flex; gap:14px; }\n  .field-row .field{ flex:1; margin-bottom:0; }\n\n  .severity{ display:flex; gap:8px; }\n  .sev-opt{\n    flex:1; text-align:center; padding:9px 8px; border-radius:7px; border:1px solid var(--line);\n    background:var(--panel-2); color:var(--muted); font-family:\'IBM Plex Mono\', monospace; font-size:12px;\n    letter-spacing:0.03em; cursor:pointer; user-select:none; display:flex; align-items:center;\n    justify-content:center; gap:6px; transition:all .15s ease;\n  }\n  .sev-dot{ width:6px; height:6px; border-radius:50%; background:var(--muted-dim); flex-shrink:0; }\n  .sev-opt[data-sev="low"] .sev-dot{ background:var(--success); }\n  .sev-opt[data-sev="medium"] .sev-dot{ background:var(--warning); }\n  .sev-opt[data-sev="high"] .sev-dot{ background:#E8873F; }\n  .sev-opt[data-sev="critical"] .sev-dot{ background:var(--danger); }\n  .sev-opt.active{ border-color:var(--accent-line); background:var(--accent-dim); color:var(--text); }\n\n  hr.divider{ border:none; border-top:1px solid var(--line-soft); margin:22px 0; }\n\n  .dropzone{\n    border:1.5px dashed var(--line); border-radius:9px; padding:26px 20px; text-align:center;\n    cursor:pointer; transition:border-color .15s ease, background .15s ease;\n  }\n  .dropzone.drag{ border-color:var(--accent); background:var(--accent-dim); }\n  .dropzone .dz-icon{ font-family:\'IBM Plex Mono\', monospace; font-size:20px; color:var(--muted-dim); margin-bottom:8px; }\n  .dropzone .dz-title{ font-size:14px; color:var(--text); margin-bottom:4px; }\n  .dropzone .dz-sub{ font-size:12.5px; color:var(--muted-dim); font-family:\'IBM Plex Mono\', monospace; }\n  input[type="file"]{ display:none; }\n\n  .file-list{ margin-top:14px; display:flex; flex-direction:column; gap:8px; }\n  .file-chip{\n    display:flex; align-items:center; gap:10px; background:var(--panel-2); border:1px solid var(--line);\n    border-radius:7px; padding:9px 11px; font-size:13px;\n  }\n  .file-chip.file-error{ border-color:rgba(226,87,76,0.4); }\n  .file-chip .fc-ext{\n    font-family:\'IBM Plex Mono\', monospace; font-size:10px; letter-spacing:0.04em;\n    background:var(--panel-3); color:var(--muted); border-radius:4px; padding:3px 6px; flex-shrink:0;\n  }\n  .file-chip .fc-name{ flex:1; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }\n  .file-chip .fc-meta{ color:var(--muted-dim); font-size:11.5px; font-family:\'IBM Plex Mono\', monospace; flex-shrink:0; }\n  .file-chip .fc-error-text{ color:var(--danger); font-size:11.5px; font-family:\'IBM Plex Mono\', monospace; flex-shrink:0; }\n  .file-chip .fc-remove{\n    background:none; border:none; color:var(--muted-dim); cursor:pointer;\n    font-size:15px; line-height:1; padding:2px 4px; border-radius:4px;\n  }\n  .file-chip .fc-remove:hover{ color:var(--danger); background:var(--danger-dim); }\n\n  .submit-btn{\n    width:100%; margin-top:22px; background:var(--accent); color:#0A2C31; border:none; border-radius:8px;\n    padding:13px; font-family:\'IBM Plex Mono\', monospace; font-size:13.5px; font-weight:600;\n    letter-spacing:0.02em; cursor:pointer; transition:filter .15s ease, opacity .15s ease;\n  }\n  .submit-btn:hover{ filter:brightness(1.08); }\n  .submit-btn:disabled{ opacity:0.45; cursor:not-allowed; filter:none; }\n\n  .hint{ font-size:12px; color:var(--muted-dim); margin-top:9px; text-align:center; }\n  .hint.error-hint{ color:var(--danger); }\n\n  .preview-panel{ position:sticky; top:20px; }\n  .preview-head{ display:flex; align-items:center; justify-content:space-between; margin-bottom:16px; }\n  .status-pill{\n    display:inline-flex; align-items:center; gap:7px; font-family:\'IBM Plex Mono\', monospace;\n    font-size:11px; letter-spacing:0.05em; text-transform:uppercase; color:var(--muted-dim);\n  }\n  .status-dot{ width:7px; height:7px; border-radius:50%; background:var(--muted-dim); }\n  .status-dot.live{ background:var(--accent); animation:pulse 1.6s ease-in-out infinite; }\n  @keyframes pulse{ 0%,100%{ box-shadow:0 0 0 0 var(--accent-dim); } 50%{ box-shadow:0 0 0 5px transparent; } }\n\n  .empty-state{ padding:30px 6px; text-align:center; color:var(--muted-dim); font-size:13px; }\n  .empty-state .ee-icon{ font-family:\'IBM Plex Mono\', monospace; font-size:22px; margin-bottom:10px; opacity:0.6; }\n\n  .badge{\n    display:inline-block; font-family:\'IBM Plex Mono\', monospace; font-size:11.5px; letter-spacing:0.03em;\n    padding:5px 10px; border-radius:5px; background:var(--accent-dim); color:var(--accent);\n    border:1px solid var(--accent-line); margin-bottom:16px;\n  }\n\n  .kv{ margin-bottom:14px; }\n  .kv-label{\n    font-family:\'IBM Plex Mono\', monospace; font-size:10.5px; text-transform:uppercase;\n    letter-spacing:0.08em; color:var(--muted-dim); margin-bottom:4px;\n  }\n  .kv-value{ font-family:\'IBM Plex Mono\', monospace; font-size:13px; color:var(--text); word-break:break-word; line-height:1.5; }\n  .kv-value.danger-text{ color:var(--danger); }\n\n  .preview-fade{ animation:fadeIn .25s ease; }\n  @keyframes fadeIn{ from{ opacity:0; transform:translateY(3px); } to{ opacity:1; transform:translateY(0); } }\n\n  .attached-mini{ margin-top:6px; display:flex; flex-direction:column; gap:6px; }\n  .attached-mini-item{\n    display:flex; justify-content:space-between; font-family:\'IBM Plex Mono\', monospace; font-size:11.5px;\n    color:var(--muted); border-top:1px solid var(--line-soft); padding-top:6px;\n  }\n\n  .confirm-banner{\n    margin-top:20px; background:var(--success-dim); border:1px solid rgba(72,184,132,0.35);\n    border-radius:9px; padding:15px 16px; display:none;\n  }\n  .confirm-banner.show{ display:block; animation:fadeIn .25s ease; }\n  .confirm-banner.error-banner{ background:var(--danger-dim); border-color:rgba(226,87,76,0.35); }\n  .confirm-title{ font-family:\'IBM Plex Mono\', monospace; font-size:13px; color:var(--success); margin-bottom:4px; }\n  .confirm-banner.error-banner .confirm-title{ color:var(--danger); }\n  .confirm-sub{ font-size:12.5px; color:var(--muted); }\n\n  .history{ margin-top:36px; }\n  .history-head{ display:flex; align-items:center; justify-content:space-between; margin-bottom:14px; }\n  .history-title{\n    font-family:\'IBM Plex Mono\', monospace; font-size:11px; letter-spacing:0.1em;\n    text-transform:uppercase; color:var(--muted-dim);\n  }\n  .btn-refresh{\n    background:transparent; border:1px solid var(--line); color:var(--muted);\n    font-family:\'IBM Plex Mono\', monospace; font-size:11px; letter-spacing:0.04em;\n    padding:5px 10px; border-radius:5px; cursor:pointer;\n  }\n  .btn-refresh:hover{ border-color:var(--muted-dim); color:var(--text); }\n  .history-empty{ color:var(--muted-dim); font-size:13px; padding:16px 0; border-top:1px solid var(--line-soft); }\n  .history-item{\n    display:flex; align-items:center; gap:14px; padding:12px 0; border-top:1px solid var(--line-soft); cursor:pointer;\n  }\n  .history-item:hover .history-title-txt{ color:var(--accent); }\n  .history-id{\n    font-family:\'IBM Plex Mono\', monospace; font-size:12px; color:var(--accent); flex-shrink:0;\n    width:150px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap;\n  }\n  .history-title-txt{ flex:1; font-size:13.5px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; transition:color .15s ease; }\n  .history-sev{\n    font-family:\'IBM Plex Mono\', monospace; font-size:10.5px; text-transform:uppercase;\n    letter-spacing:0.04em; padding:3px 8px; border-radius:5px; flex-shrink:0;\n  }\n  .history-time{ font-family:\'IBM Plex Mono\', monospace; font-size:11.5px; color:var(--muted-dim); flex-shrink:0; }\n\n  .modal{ position:fixed; inset:0; background:rgba(0,0,0,0.6); display:flex; align-items:center; justify-content:center; padding:20px; z-index:50; }\n  .modal.hidden{ display:none; }\n  .modal-content{\n    background:var(--panel); border:1px solid var(--line); border-radius:10px; max-width:700px; width:100%;\n    max-height:80vh; overflow-y:auto; padding:24px; position:relative;\n  }\n  .modal-close{ position:absolute; top:12px; right:16px; background:none; border:none; color:var(--muted-dim); font-size:1.4rem; cursor:pointer; }\n  .modal-meta{ color:var(--muted); font-size:12.5px; font-family:\'IBM Plex Mono\', monospace; }\n  .modal-body{\n    white-space:pre-wrap; word-break:break-word; background:var(--panel-2); border:1px solid var(--line);\n    border-radius:8px; padding:14px; font-family:\'IBM Plex Mono\', monospace; font-size:12.5px; margin-top:12px;\n  }\n</style>\n</head>\n<body>\n<div class="wrap">\n\n  <p class="eyebrow">AI Smart Bug Analyzer &amp; Fix Advisor</p>\n  <div class="header-row">\n    <h1>Bug submission</h1>\n    <button class="btn-clear" id="clearBtn" type="button">Clear form</button>\n  </div>\n  <p class="subtitle">Paste an error, or drop a log file. Submissions are saved to the analyzer\'s database.</p>\n\n  <div class="grid">\n\n    <div>\n      <div class="panel">\n\n        <div class="field-row">\n          <div class="field">\n            <label class="label" for="title">Title</label>\n            <input type="text" id="title" placeholder="Checkout fails with 500 on submit" />\n          </div>\n        </div>\n\n        <div class="field">\n          <label class="label">Severity</label>\n          <div class="severity" id="severity">\n            <div class="sev-opt" data-sev="low"><span class="sev-dot"></span>Low</div>\n            <div class="sev-opt active" data-sev="medium"><span class="sev-dot"></span>Medium</div>\n            <div class="sev-opt" data-sev="high"><span class="sev-dot"></span>High</div>\n            <div class="sev-opt" data-sev="critical"><span class="sev-dot"></span>Critical</div>\n          </div>\n        </div>\n\n        <hr class="divider" />\n\n        <div class="field">\n          <label class="label" for="pasteArea">Paste report, stack trace, or log</label>\n          <textarea id="pasteArea" placeholder="Traceback (most recent call last):&#10;  File &quot;app/views.py&quot;, line 88, in checkout&#10;    charge = stripe.Charge.create(...)&#10;stripe.error.CardError: Your card was declined."></textarea>\n        </div>\n\n        <div class="field">\n          <label class="label">Attach files</label>\n          <div class="dropzone" id="dropzone">\n            <div class="dz-icon">&#8595;</div>\n            <div class="dz-title">Drop files here, or click to browse</div>\n            <div class="dz-sub">.log .txt .json .xml .csv .stacktrace .out .pdf .docx — up to 10 files, 5MB each</div>\n          </div>\n          <input type="file" id="fileInput" multiple accept=".txt,.log,.json,.xml,.csv,.stacktrace,.out,.pdf,.docx" />\n          <div class="file-list" id="fileList"></div>\n        </div>\n\n        <button class="submit-btn" id="submitBtn" disabled>Submit for analysis</button>\n        <div class="hint" id="submitHint">Add pasted text or at least one file to enable submission.</div>\n\n        <div class="confirm-banner" id="confirmBanner">\n          <div class="confirm-title" id="confirmTitle"></div>\n          <div class="confirm-sub" id="confirmSub"></div>\n        </div>\n\n      </div>\n\n      <div class="history">\n        <div class="history-head">\n          <div class="history-title">Submission history</div>\n          <button class="btn-refresh" id="refreshHistoryBtn" type="button">Refresh</button>\n        </div>\n        <div id="historyList">\n          <div class="history-empty">Loading submissions...</div>\n        </div>\n      </div>\n    </div>\n\n    <div class="preview-panel">\n      <div class="panel">\n        <div class="preview-head">\n          <span class="label" style="margin-bottom:0;">Diagnostic preview</span>\n          <span class="status-pill"><span class="status-dot" id="statusDot"></span><span id="statusText">Idle</span></span>\n        </div>\n        <div id="previewBody">\n          <div class="empty-state">\n            <div class="ee-icon">&#10022;</div>\n            Paste an error or drop a file to begin analysis.\n          </div>\n        </div>\n      </div>\n    </div>\n\n  </div>\n</div>\n\n<div id="detailModal" class="modal hidden">\n  <div class="modal-content">\n    <button id="closeModal" class="modal-close" type="button">&times;</button>\n    <h3 id="modalTitle" class="mono" style="margin:0 0 6px;"></h3>\n    <p id="modalMeta" class="modal-meta"></p>\n    <pre id="modalContent" class="modal-body"></pre>\n  </div>\n</div>\n\n<script>\n(function(){\n  const API_BASE = "";\n\n  const pasteArea = document.getElementById(\'pasteArea\');\n  const dropzone = document.getElementById(\'dropzone\');\n  const fileInput = document.getElementById(\'fileInput\');\n  const fileList = document.getElementById(\'fileList\');\n  const submitBtn = document.getElementById(\'submitBtn\');\n  const submitHint = document.getElementById(\'submitHint\');\n  const previewBody = document.getElementById(\'previewBody\');\n  const statusDot = document.getElementById(\'statusDot\');\n  const statusText = document.getElementById(\'statusText\');\n  const confirmBanner = document.getElementById(\'confirmBanner\');\n  const confirmTitle = document.getElementById(\'confirmTitle\');\n  const confirmSub = document.getElementById(\'confirmSub\');\n  const historyList = document.getElementById(\'historyList\');\n  const titleInput = document.getElementById(\'title\');\n  const severityWrap = document.getElementById(\'severity\');\n  const clearBtn = document.getElementById(\'clearBtn\');\n  const refreshHistoryBtn = document.getElementById(\'refreshHistoryBtn\');\n\n  const MAX_FILES = 10;\n  const MAX_FILE_SIZE = 5 * 1024 * 1024;\n  const ALLOWED_EXT = [\'txt\',\'log\',\'json\',\'xml\',\'csv\',\'stacktrace\',\'out\',\'pdf\',\'docx\'];\n  const BINARY_EXT = [\'pdf\',\'docx\'];\n\n  let attachedFiles = [];\n  let currentSeverity = \'medium\';\n  let isSubmitting = false;\n\n  severityWrap.querySelectorAll(\'.sev-opt\').forEach(function(opt){\n    opt.addEventListener(\'click\', function(){\n      severityWrap.querySelectorAll(\'.sev-opt\').forEach(function(o){ o.classList.remove(\'active\'); });\n      opt.classList.add(\'active\');\n      currentSeverity = opt.dataset.sev;\n    });\n  });\n\n  function analyzeContent(text){\n    const result = {\n      type: null, exception: null, message: null, file: null, line: null,\n      language: null, lines: text.split(\'\\n\').length, chars: text.length\n    };\n    if (!text || !text.trim()) return null;\n\n    if (/Traceback \\(most recent call last\\)/.test(text)) {\n      result.type = \'Python traceback\';\n      result.language = \'Python\';\n      const fileMatch = text.match(/File "([^"]+)", line (\\d+)/g);\n      if (fileMatch && fileMatch.length) {\n        const m = fileMatch[fileMatch.length - 1].match(/File "([^"]+)", line (\\d+)/);\n        if (m) { result.file = m[1]; result.line = m[2]; }\n      }\n      const excMatch = text.match(/^([A-Za-z_.]+(?:Error|Exception|Warning))\\s*:\\s*(.*)$/m);\n      if (excMatch) { result.exception = excMatch[1]; result.message = excMatch[2]; }\n    } else if (/at [\\w$.<>]+\\s*\\(?.*\\.java:\\d+\\)?/.test(text) || /Exception in thread/.test(text)) {\n      result.type = \'Java stack trace\';\n      result.language = \'Java\';\n      const m = text.match(/([\\w.$]+):(\\d+)\\)/) || text.match(/\\(([\\w.$]+\\.java):(\\d+)\\)/);\n      if (m) { result.file = m[1]; result.line = m[2]; }\n      const excMatch = text.match(/([\\w.]+(?:Exception|Error))(?::\\s*(.*))?/);\n      if (excMatch) { result.exception = excMatch[1]; result.message = excMatch[2] || null; }\n    } else if (/at .*\\(.*:\\d+:\\d+\\)/.test(text) || /Uncaught|TypeError|ReferenceError/.test(text)) {\n      result.type = \'JS/Node runtime error\';\n      result.language = \'JavaScript\';\n      const m = text.match(/at .*\\(?([\\w./\\\\-]+\\.(?:js|ts|jsx|tsx)):(\\d+):(\\d+)\\)?/);\n      if (m) { result.file = m[1]; result.line = m[2]; }\n      const excMatch = text.match(/^([A-Za-z]+Error)\\s*:\\s*(.*)$/m);\n      if (excMatch) { result.exception = excMatch[1]; result.message = excMatch[2]; }\n    } else if (/^\\d{4}-\\d{2}-\\d{2}[T ]\\d{2}:\\d{2}:\\d{2}/m.test(text)) {\n      result.type = \'Structured log\';\n      const lvl = text.match(/\\b(ERROR|WARN|FATAL|CRITICAL)\\b/);\n      if (lvl) result.exception = lvl[1] + \' entries detected\';\n    } else {\n      result.type = \'Unstructured report\';\n    }\n\n    if (!result.file) {\n      const generic = text.match(/([\\w./\\\\-]+\\.(?:py|js|ts|java|go|rb|cpp|c|cs|php)):(\\d+)/);\n      if (generic) { result.file = generic[1]; result.line = generic[2]; }\n    }\n    return result;\n  }\n\n  function extForFile(name){\n    const parts = name.split(\'.\');\n    return parts.length > 1 ? parts[parts.length-1].toUpperCase() : \'FILE\';\n  }\n  function extLower(name){\n    const parts = name.split(\'.\');\n    return parts.length > 1 ? parts[parts.length-1].toLowerCase() : \'\';\n  }\n  function formatSize(bytes){\n    if (bytes < 1024) return bytes + \' B\';\n    if (bytes < 1024*1024) return (bytes/1024).toFixed(1) + \' KB\';\n    return (bytes/(1024*1024)).toFixed(1) + \' MB\';\n  }\n  function escapeHtml(str){\n    return String(str).replace(/[&<>"\']/g, function(c){\n      return { \'&\':\'&amp;\', \'<\':\'&lt;\', \'>\':\'&gt;\', \'"\':\'&quot;\', "\'":\'&#39;\' }[c];\n    });\n  }\n\n  function renderFileList(){\n    fileList.innerHTML = \'\';\n    attachedFiles.forEach(function(f, idx){\n      const chip = document.createElement(\'div\');\n      chip.className = \'file-chip\' + (f.error ? \' file-error\' : \'\');\n      chip.innerHTML =\n        \'<span class="fc-ext">\' + extForFile(f.name) + \'</span>\' +\n        \'<span class="fc-name">\' + escapeHtml(f.name) + \'</span>\' +\n        (f.error\n          ? \'<span class="fc-error-text">\' + escapeHtml(f.error) + \'</span>\'\n          : \'<span class="fc-meta">\' + formatSize(f.size) + \'</span>\') +\n        \'<button class="fc-remove" type="button" aria-label="Remove file">&#10005;</button>\';\n      chip.querySelector(\'.fc-remove\').addEventListener(\'click\', function(){\n        attachedFiles.splice(idx, 1);\n        renderFileList();\n        updatePreview();\n      });\n      fileList.appendChild(chip);\n    });\n  }\n\n  function updatePreview(){\n    const text = pasteArea.value;\n    const pasteAnalysis = analyzeContent(text);\n    const hasFiles = attachedFiles.length > 0;\n    const hasValidFiles = attachedFiles.some(function(f){ return !f.error; });\n    const hasContent = (text && text.trim().length > 0) || hasValidFiles;\n\n    submitBtn.disabled = !hasContent || isSubmitting;\n    if (!isSubmitting) {\n      submitHint.textContent = hasContent ? \'Ready to submit.\' : \'Add pasted text or at least one file to enable submission.\';\n      submitHint.classList.remove(\'error-hint\');\n    }\n\n    if (!hasContent) {\n      statusDot.classList.remove(\'live\');\n      statusText.textContent = \'Idle\';\n      previewBody.innerHTML = \'<div class="empty-state"><div class="ee-icon">&#10022;</div>Paste an error or drop a file to begin analysis.</div>\';\n      return;\n    }\n\n    statusDot.classList.add(\'live\');\n    statusText.textContent = \'Reading\';\n\n    let html = \'\';\n    if (pasteAnalysis) {\n      html += \'<div class="badge preview-fade">\' + pasteAnalysis.type + \'</div>\';\n      if (pasteAnalysis.exception) {\n        html += \'<div class="kv preview-fade"><div class="kv-label">Exception</div><div class="kv-value danger-text">\' + escapeHtml(pasteAnalysis.exception) + \'</div></div>\';\n      }\n      if (pasteAnalysis.message) {\n        html += \'<div class="kv preview-fade"><div class="kv-label">Message</div><div class="kv-value">\' + escapeHtml(pasteAnalysis.message) + \'</div></div>\';\n      }\n      if (pasteAnalysis.file) {\n        html += \'<div class="kv preview-fade"><div class="kv-label">Location</div><div class="kv-value">\' + escapeHtml(pasteAnalysis.file) + (pasteAnalysis.line ? \':\' + pasteAnalysis.line : \'\') + \'</div></div>\';\n      }\n      if (pasteAnalysis.language) {\n        html += \'<div class="kv preview-fade"><div class="kv-label">Language</div><div class="kv-value">\' + pasteAnalysis.language + \'</div></div>\';\n      }\n      html += \'<div class="kv preview-fade"><div class="kv-label">Pasted content</div><div class="kv-value">\' + pasteAnalysis.lines + \' lines &middot; \' + pasteAnalysis.chars + \' chars</div></div>\';\n    } else if (!hasFiles) {\n      html += \'<div class="empty-state"><div class="ee-icon">&#10022;</div>Paste an error or drop a file to begin analysis.</div>\';\n    }\n\n    if (hasFiles) {\n      html += \'<div class="kv preview-fade"><div class="kv-label">Attached files (\' + attachedFiles.length + \')</div><div class="attached-mini">\';\n      attachedFiles.forEach(function(f){\n        let label;\n        if (f.error) {\n          label = f.error;\n        } else if (f.binary) {\n          label = \'Document file — parsed on server (preview not shown client-side)\';\n        } else if (f.content) {\n          label = (analyzeContent(f.content) || {type:\'Unreadable\'}).type;\n        } else {\n          label = \'Reading...\';\n        }\n        html += \'<div class="attached-mini-item"><span>\' + escapeHtml(f.name) + \'</span><span>\' + escapeHtml(label) + \'</span></div>\';\n      });\n      html += \'</div></div>\';\n    }\n\n    previewBody.innerHTML = html;\n  }\n\n  pasteArea.addEventListener(\'input\', updatePreview);\n\n  dropzone.addEventListener(\'click\', function(){ fileInput.click(); });\n  dropzone.addEventListener(\'dragover\', function(e){ e.preventDefault(); dropzone.classList.add(\'drag\'); });\n  dropzone.addEventListener(\'dragleave\', function(){ dropzone.classList.remove(\'drag\'); });\n  dropzone.addEventListener(\'drop\', function(e){\n    e.preventDefault();\n    dropzone.classList.remove(\'drag\');\n    handleFiles(e.dataTransfer.files);\n  });\n  fileInput.addEventListener(\'change\', function(e){\n    handleFiles(e.target.files);\n    fileInput.value = \'\';\n  });\n\n  function handleFiles(fileListObj){\n    const incoming = Array.from(fileListObj);\n\n    if (attachedFiles.length + incoming.length > MAX_FILES) {\n      submitHint.textContent = \'Max \' + MAX_FILES + \' files per submission. Remove some files first.\';\n      submitHint.classList.add(\'error-hint\');\n      return;\n    }\n\n    incoming.forEach(function(file){\n      const ext = extLower(file.name);\n      const isBinary = BINARY_EXT.includes(ext);\n      const entry = { file: file, name: file.name, size: file.size, content: null, error: null, binary: isBinary };\n\n      if (!ALLOWED_EXT.includes(ext)) {\n        entry.error = \'Unsupported type (.\' + ext + \')\';\n      } else if (file.size === 0) {\n        entry.error = \'File is empty\';\n      } else if (file.size > MAX_FILE_SIZE) {\n        entry.error = \'Too large (max 5MB)\';\n      }\n\n      attachedFiles.push(entry);\n\n      // PDF/DOCX are binary formats — text is extracted server-side, so we\n      // skip the client-side text preview for those and just show a badge.\n      if (!entry.error && !isBinary) {\n        const reader = new FileReader();\n        reader.onload = function(ev){\n          entry.content = ev.target.result;\n          renderFileList();\n          updatePreview();\n        };\n        reader.readAsText(file);\n      }\n    });\n    renderFileList();\n    updatePreview();\n  }\n\n  submitBtn.addEventListener(\'click\', async function(){\n    if (submitBtn.disabled || isSubmitting) return;\n\n    const title = titleInput.value.trim();\n    const text = pasteArea.value.trim();\n    const validFiles = attachedFiles.filter(function(f){ return !f.error; });\n\n    if (!title) {\n      submitHint.textContent = \'Title is required.\';\n      submitHint.classList.add(\'error-hint\');\n      titleInput.focus();\n      return;\n    }\n    if (attachedFiles.some(function(f){ return f.error; })) {\n      submitHint.textContent = \'Remove or fix invalid files before submitting.\';\n      submitHint.classList.add(\'error-hint\');\n      return;\n    }\n    if (!text && validFiles.length === 0) {\n      submitHint.textContent = \'Add pasted text or at least one file to enable submission.\';\n      submitHint.classList.add(\'error-hint\');\n      return;\n    }\n\n    isSubmitting = true;\n    submitBtn.disabled = true;\n    submitBtn.textContent = \'Submitting...\';\n    statusText.textContent = \'Submitting\';\n    hideConfirm();\n\n    try {\n      const results = [];\n\n      if (text) {\n        const fd = new FormData();\n        fd.append(\'title\', title);\n        fd.append(\'content\', text);\n        fd.append(\'severity\', currentSeverity);\n        const res = await fetch(API_BASE + \'/api/bugs/paste\', { method: \'POST\', body: fd });\n        const data = await res.json();\n        if (!res.ok) throw new Error(data.detail || \'Paste submission failed.\');\n        results.push(data);\n      }\n\n      if (validFiles.length > 0) {\n        const fd = new FormData();\n        fd.append(\'title\', title);\n        fd.append(\'severity\', currentSeverity);\n        validFiles.forEach(function(f){ fd.append(\'files\', f.file); });\n        const res = await fetch(API_BASE + \'/api/bugs/upload\', { method: \'POST\', body: fd });\n        const data = await res.json();\n        if (!res.ok) throw new Error(data.detail || \'File upload failed.\');\n        results.push(data);\n      }\n\n      showConfirm(false, results.length + \' submission(s) queued\', results.map(function(r){ return r.bug_id; }).join(\', \'));\n      resetForm();\n      await loadHistory();\n\n    } catch (err) {\n      showConfirm(true, \'Submission failed\', err.message);\n    } finally {\n      isSubmitting = false;\n      submitBtn.textContent = \'Submit for analysis\';\n      updatePreview();\n    }\n  });\n\n  function showConfirm(isError, title, sub){\n    confirmBanner.classList.toggle(\'error-banner\', isError);\n    confirmTitle.textContent = title;\n    confirmSub.textContent = sub;\n    confirmBanner.classList.add(\'show\');\n    setTimeout(hideConfirm, 5000);\n  }\n  function hideConfirm(){\n    confirmBanner.classList.remove(\'show\');\n  }\n\n  function resetForm(){\n    pasteArea.value = \'\';\n    attachedFiles = [];\n    titleInput.value = \'\';\n    severityWrap.querySelectorAll(\'.sev-opt\').forEach(function(o){ o.classList.remove(\'active\'); });\n    severityWrap.querySelector(\'[data-sev="medium"]\').classList.add(\'active\');\n    currentSeverity = \'medium\';\n    renderFileList();\n    updatePreview();\n  }\n\n  clearBtn.addEventListener(\'click\', function(){\n    resetForm();\n    submitHint.textContent = \'Add pasted text or at least one file to enable submission.\';\n    submitHint.classList.remove(\'error-hint\');\n  });\n\n  const sevColors = {\n    low: { bg: \'var(--success-dim)\', text: \'var(--success)\' },\n    medium: { bg: \'var(--warning-dim)\', text: \'var(--warning)\' },\n    high: { bg: \'rgba(232,135,63,0.12)\', text: \'#E8873F\' },\n    critical: { bg: \'var(--danger-dim)\', text: \'var(--danger)\' }\n  };\n\n  async function loadHistory(){\n    try {\n      const res = await fetch(API_BASE + \'/api/bugs\');\n      const bugs = await res.json();\n      if (!res.ok) throw new Error(bugs.detail || \'Failed to load history.\');\n\n      if (!bugs.length) {\n        historyList.innerHTML = \'<div class="history-empty">Nothing submitted yet. Your queued analyses will appear here.</div>\';\n        return;\n      }\n\n      historyList.innerHTML = \'\';\n      bugs.forEach(function(bug){\n        const c = sevColors[bug.severity] || sevColors.medium;\n        const time = new Date(bug.created_at).toLocaleString();\n        const item = document.createElement(\'div\');\n        item.className = \'history-item\';\n        item.innerHTML =\n          \'<span class="history-id">\' + escapeHtml(bug.bug_id) + \'</span>\' +\n          \'<span class="history-title-txt">\' + escapeHtml(bug.title) + \'</span>\' +\n          \'<span class="history-sev" style="background:\' + c.bg + \';color:\' + c.text + \'">\' + bug.severity + \'</span>\' +\n          \'<span class="history-time">\' + time + \'</span>\';\n        item.addEventListener(\'click\', function(){ openDetail(bug.id); });\n        historyList.appendChild(item);\n      });\n    } catch (err) {\n      historyList.innerHTML = \'<div class="history-empty">Failed to load history: \' + escapeHtml(err.message) + \'</div>\';\n    }\n  }\n\n  refreshHistoryBtn.addEventListener(\'click\', loadHistory);\n\n  const modal = document.getElementById(\'detailModal\');\n  document.getElementById(\'closeModal\').addEventListener(\'click\', function(){ modal.classList.add(\'hidden\'); });\n  modal.addEventListener(\'click\', function(e){ if (e.target === modal) modal.classList.add(\'hidden\'); });\n\n  async function openDetail(id){\n    try {\n      const res = await fetch(API_BASE + \'/api/bugs/\' + id);\n      const bug = await res.json();\n      if (!res.ok) throw new Error(bug.detail || \'Failed to load detail.\');\n      document.getElementById(\'modalTitle\').textContent = bug.bug_id + \' — \' + bug.title;\n      document.getElementById(\'modalMeta\').textContent =\n        (bug.source_type === \'upload\' ? \'File: \' + bug.original_filename : \'Source: Pasted text\') +\n        \' · \' + bug.severity + \' · \' + new Date(bug.created_at).toLocaleString() + \' · \' + bug.content_length + \' chars\';\n      document.getElementById(\'modalContent\').textContent = bug.content;\n      modal.classList.remove(\'hidden\');\n    } catch (err) {\n      showConfirm(true, \'Could not load submission\', err.message);\n    }\n  }\n\n  updatePreview();\n  loadHistory();\n})();\n</script>\n</body>\n</html>\n'


@app.get("/", response_class=HTMLResponse)
def serve_frontend():
    return FRONTEND_HTML


# ---------------------------------------------------------------------------
# API routes
# ---------------------------------------------------------------------------

@app.get("/api/health")
def health_check():
    return {"status": "ok", "time": datetime.utcnow().isoformat()}


@app.post("/api/bugs/paste", response_model=BugSubmissionCreatedResponse)
def submit_pasted_bug(
    title: str = Form(...),
    content: str = Form(...),
    severity: str = Form("medium"),
    db: Session = Depends(get_db),
):
    content = content.strip()
    if not content:
        raise HTTPException(status_code=400, detail="Pasted content cannot be empty.")
    if not title.strip():
        raise HTTPException(status_code=400, detail="Title cannot be empty.")
    sev = _validate_severity(severity)

    bug_number = _next_bug_number(db)
    bug = BugSubmission(
        bug_number=bug_number,
        bug_id=_format_bug_id(bug_number),
        title=title.strip(),
        severity=sev,
        source_type=SourceType.PASTE,
        original_filename=None,
        file_type="text/plain",
        content=content,
        content_length=str(len(content)),
    )
    db.add(bug)
    db.commit()
    db.refresh(bug)

    return {"id": bug.id, "bug_id": bug.bug_id, "message": "Bug report submitted successfully."}


@app.post("/api/bugs/upload", response_model=BugSubmissionCreatedResponse)
async def submit_uploaded_bug(
    title: str = Form(...),
    severity: str = Form("medium"),
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
):
    if not title.strip():
        raise HTTPException(status_code=400, detail="Title cannot be empty.")
    if not files:
        raise HTTPException(status_code=400, detail="At least one file is required.")
    if len(files) > MAX_FILES_PER_UPLOAD:
        raise HTTPException(
            status_code=400,
            detail=f"Too many files. Max {MAX_FILES_PER_UPLOAD} files per submission.",
        )
    sev = _validate_severity(severity)

    created_ids = []
    first_bug_id = None
    for file in files:
        ext = _validate_extension(file.filename)

        raw = await file.read()
        if len(raw) == 0:
            raise HTTPException(status_code=400, detail=f"'{file.filename}' is empty.")
        if len(raw) > MAX_FILE_SIZE_BYTES:
            max_mb = MAX_FILE_SIZE_BYTES // (1024 * 1024)
            raise HTTPException(
                status_code=400,
                detail=f"'{file.filename}' is too large. Max size is {max_mb} MB.",
            )

        # Dispatches to plain-text decoding, PDF text extraction, or DOCX
        # text extraction, depending on the file's extension.
        text_content = _extract_text_content(file.filename, ext, raw)

        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S%f")
        safe_name = f"{timestamp}_{os.path.basename(file.filename)}"
        disk_path = os.path.join(UPLOAD_DIR, safe_name)
        with open(disk_path, "wb") as f:
            f.write(raw)

        bug_title = title.strip() if len(files) == 1 else f"{title.strip()} — {file.filename}"
        bug_number = _next_bug_number(db)
        bug = BugSubmission(
            bug_number=bug_number,
            bug_id=_format_bug_id(bug_number),
            title=bug_title,
            severity=sev,
            source_type=SourceType.UPLOAD,
            original_filename=file.filename,
            file_type=ext,
            content=text_content,
            content_length=str(len(text_content)),
        )
        db.add(bug)
        db.commit()
        db.refresh(bug)
        created_ids.append(bug.id)
        if first_bug_id is None:
            first_bug_id = bug.bug_id

    return {
        "id": created_ids[0],
        "bug_id": first_bug_id,
        "message": f"{len(created_ids)} file(s) uploaded and submitted successfully.",
    }


@app.get("/api/bugs", response_model=List[BugSubmissionResponse])
def list_bugs(db: Session = Depends(get_db)):
    bugs = db.query(BugSubmission).order_by(desc(BugSubmission.created_at)).all()
    return [_serialize(b) for b in bugs]


@app.get("/api/bugs/{bug_id}", response_model=BugSubmissionDetail)
def get_bug(bug_id: str, db: Session = Depends(get_db)):
    """
    Looks up a bug by either its internal UUID (id) or its human-readable
    bug_id (e.g. "BUG-00001"), so both identifiers work as a lookup key.
    """
    bug = (
        db.query(BugSubmission)
        .filter((BugSubmission.id == bug_id) | (BugSubmission.bug_id == bug_id))
        .first()
    )
    if not bug:
        raise HTTPException(status_code=404, detail="Bug submission not found.")
    data = _serialize(bug)
    data["content"] = bug.content
    return data


@app.delete("/api/bugs/{bug_id}")
def delete_bug(bug_id: str, db: Session = Depends(get_db)):
    bug = (
        db.query(BugSubmission)
        .filter((BugSubmission.id == bug_id) | (BugSubmission.bug_id == bug_id))
        .first()
    )
    if not bug:
        raise HTTPException(status_code=404, detail="Bug submission not found.")
    db.delete(bug)
    db.commit()
    return {"message": "Deleted successfully."}