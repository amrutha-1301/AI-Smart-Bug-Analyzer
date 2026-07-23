"""
AI Smart Bug Analyzer & Fix Advisor
Milestone 2 — Triage Agent + Log Analysis Agent + Multi-Agent Orchestration
Single-file version (per company convention established in Milestone 1)

Builds directly on top of the Milestone 1 Bug Submission Module. Everything
still runs from one process: FastAPI backend, PostgreSQL (via SQLAlchemy,
with automatic local SQLite fallback), and the full HTML/CSS/JS frontend
served inline.

Tech stack (per project tech stack doc):
  - Backend:  Python, FastAPI
  - Database: PostgreSQL (SQLAlchemy ORM; falls back to SQLite for local/dev
              runs if DATABASE_URL is not set)
  - Frontend: HTML, CSS, JavaScript (served inline by FastAPI, no build step)

What's new in Milestone 2 (see design doc "AI Agent Layer" / "Orchestration
Flow"):
  - TriageAgent        -> severity, priority, component, confidence, reasoning
  - LogAnalysisAgent    -> exception type, message, failure point, affected
                           code path(s), language, error signature, confidence,
                           reasoning
  - Orchestrator        -> runs both agents automatically on every submission
                           (paste or upload) and persists a combined,
                           structured record keyed by bug_id so Milestone 3
                           (Duplicate Detection + Root Cause) and Milestone 4
                           (Remediation) can consume it directly, with no
                           re-parsing required.
  - validate_agents.py  -> separate script that runs both agents against a
                           seeded seeded seeded dataset of varied bug report
                           styles/error types and reports accuracy.

Run with:
    pip install fastapi uvicorn sqlalchemy psycopg2-binary python-multipart pydantic pypdf python-docx
    python3 app.py
    (or: uvicorn app:app --reload --port 3000)

Then open http://localhost:3000 in Chrome (or any browser).

Configure PostgreSQL by setting the DATABASE_URL environment variable, e.g.:
    export DATABASE_URL=postgresql://username:password@localhost:5432/bug_analyzer
If unset, a local SQLite file (bug_analyzer.db) is used automatically so this
runs immediately without any database setup.
"""

import io
import os
import re
import json
import enum
import uuid
import hashlib
from datetime import datetime
from typing import List, Optional, Dict, Any, Tuple

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from sqlalchemy import Column, String, Text, Integer, Float, DateTime, Enum, ForeignKey, create_engine, desc, func as sa_func
from sqlalchemy.orm import declarative_base, sessionmaker, Session, relationship
from sqlalchemy.sql import func
from pydantic import BaseModel

# Optional third-party parsers for PDF / DOCX text extraction.
# The app still runs without them, but .pdf / .docx uploads will be
# rejected with a clear error until the packages are installed.
try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None

try:
    import docx as docx_lib  # python-docx
except ImportError:  # pragma: no cover
    docx_lib = None


# ---------------------------------------------------------------------------
# Database (PostgreSQL via SQLAlchemy, SQLite fallback for zero-setup local runs)
# ---------------------------------------------------------------------------

DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./bug_analyzer.db")
connect_args = {"check_same_thread": False} if DATABASE_URL.startswith("sqlite") else {}

engine = create_engine(DATABASE_URL, connect_args=connect_args)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------

class SourceType(str, enum.Enum):
    PASTE = "paste"
    UPLOAD = "upload"


class Severity(str, enum.Enum):
    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"
    CRITICAL = "critical"


class BugSubmission(Base):
    __tablename__ = "bug_submissions"

    # Internal primary key (UUID) — used for API lookups / routing.
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))

    # Sequential numeric counter, used to build the human-readable bug_id.
    bug_number = Column(Integer, nullable=False)

    # Human-readable, sequential bug identifier, e.g. "BUG-00001".
    # Matches the `bug_id` field from the BUG_RECORD design.
    bug_id = Column(String(20), unique=True, index=True, nullable=False)

    title = Column(String(255), nullable=False)
    severity = Column(Enum(Severity), nullable=False, default=Severity.MEDIUM)
    source_type = Column(Enum(SourceType), nullable=False)
    original_filename = Column(String(255), nullable=True)
    file_type = Column(String(50), nullable=True)
    content = Column(Text, nullable=False)
    content_length = Column(String(20), nullable=True)
    created_at = Column(DateTime(timezone=True), server_default=func.now())

    analysis = relationship(
        "AgentAnalysis", back_populates="bug", uselist=False, cascade="all, delete-orphan"
    )


class AgentAnalysis(Base):
    """
    Milestone 2 output: the combined, structured result of the Triage Agent
    and the Log Analysis Agent for a single bug submission.

    This mirrors the BUG_RECORD fields from the Knowledge Base data model
    (severity, type, component, root_cause placeholder for Milestone 3) and
    is intentionally denormalized/JSON-friendly so downstream agents
    (Duplicate Detection, Root Cause Analysis, Remediation in Milestones 3-4)
    can consume `combined_json` directly without re-parsing raw text.
    """
    __tablename__ = "agent_analyses"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    bug_id = Column(String, ForeignKey("bug_submissions.id"), unique=True, nullable=False, index=True)

    # --- Triage Agent output --------------------------------------------
    triage_severity = Column(String(20), nullable=False)
    triage_priority = Column(String(10), nullable=False)          # P0-P3
    triage_type = Column(String(30), nullable=False)               # Bug / Enhancement / Others
    triage_component = Column(String(100), nullable=False)
    triage_confidence = Column(Float, nullable=False)
    triage_reasoning = Column(Text, nullable=False)

    # --- Log Analysis Agent output ---------------------------------------
    log_exception_type = Column(String(150), nullable=True)
    log_message = Column(Text, nullable=True)
    log_language = Column(String(30), nullable=True)
    log_failure_file = Column(String(255), nullable=True)
    log_failure_line = Column(String(20), nullable=True)
    log_affected_paths = Column(Text, nullable=True)               # JSON list
    log_error_signature = Column(String(64), nullable=True)        # hash, for future dup-detection
    log_confidence = Column(Float, nullable=False)
    log_reasoning = Column(Text, nullable=False)

    # Full structured payload, exactly what Milestone 3/4 agents consume.
    combined_json = Column(Text, nullable=False)

    created_at = Column(DateTime(timezone=True), server_default=func.now())

    bug = relationship("BugSubmission", back_populates="analysis")


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------

class BugSubmissionResponse(BaseModel):
    id: str
    bug_id: str
    title: str
    severity: str
    source_type: str
    original_filename: Optional[str]
    file_type: Optional[str]
    content_length: Optional[str]
    created_at: datetime

    class Config:
        from_attributes = True


class BugSubmissionDetail(BugSubmissionResponse):
    content: str


class BugSubmissionCreatedResponse(BaseModel):
    id: str
    bug_id: str
    message: str
    analysis: Optional[Dict[str, Any]] = None


class AgentAnalysisResponse(BaseModel):
    bug_id: str
    triage: Dict[str, Any]
    log_analysis: Dict[str, Any]
    analyzed_at: datetime

    class Config:
        from_attributes = True


# ---------------------------------------------------------------------------
# App config
# ---------------------------------------------------------------------------

# Plain-text / log formats, read directly as UTF-8 text.
TEXT_EXTENSIONS = {".txt", ".log", ".json", ".xml", ".csv", ".stacktrace", ".out"}
# Rich-document formats, read via a dedicated parser to extract plain text.
DOCUMENT_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_EXTENSIONS = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB
MAX_FILES_PER_UPLOAD = 10
UPLOAD_DIR = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

BUG_ID_PREFIX = "BUG-"
BUG_ID_PAD = 5  # BUG-00001

Base.metadata.create_all(bind=engine)

app = FastAPI(
    title="AI Smart Bug Analyzer & Fix Advisor — Bug Submission + Agent Layer",
    version="2.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _validate_severity(severity: str) -> Severity:
    try:
        return Severity(severity.lower())
    except ValueError:
        allowed = "low, medium, high, critical"
        raise HTTPException(
            status_code=400,
            detail=f"Invalid severity '{severity}'. Must be one of: {allowed}.",
        )


def _validate_extension(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed types: {allowed}",
        )
    return ext


def _extract_pdf_text(filename: str, raw: bytes) -> str:
    if PdfReader is None:
        raise HTTPException(
            status_code=500,
            detail="PDF support is not installed on the server. Run: pip install pypdf",
        )
    try:
        reader = PdfReader(io.BytesIO(raw))
        pages_text = []
        for page in reader.pages:
            pages_text.append(page.extract_text() or "")
        text = "\n".join(pages_text).strip()
    except Exception:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' could not be read as a valid PDF file.",
        )
    if not text:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' contains no extractable text (it may be a scanned/image-only PDF).",
        )
    return text


def _extract_docx_text(filename: str, raw: bytes) -> str:
    if docx_lib is None:
        raise HTTPException(
            status_code=500,
            detail="DOCX support is not installed on the server. Run: pip install python-docx",
        )
    try:
        document = docx_lib.Document(io.BytesIO(raw))
        paragraphs = [p.text for p in document.paragraphs]
        # Also pull text out of any tables in the document.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        paragraphs.append(cell.text)
        text = "\n".join(paragraphs).strip()
    except Exception:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' could not be read as a valid DOCX file.",
        )
    if not text:
        raise HTTPException(
            status_code=400,
            detail=f"'{filename}' contains no extractable text.",
        )
    return text


def _extract_text_content(filename: str, ext: str, raw: bytes) -> str:
    """
    Returns the plain-text content to store for a given uploaded file,
    dispatching to the right reader based on its extension.
    """
    if ext == ".pdf":
        return _extract_pdf_text(filename, raw)
    if ext == ".docx":
        return _extract_docx_text(filename, raw)
    # Plain-text / log formats.
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        raise HTTPException(
            status_code=400, detail=f"'{filename}' must be UTF-8 encoded plain text."
        )


def _next_bug_number(db: Session) -> int:
    """
    Returns the next sequential bug number, based on the current max
    in the table. Safe enough for a single-process dev/demo app; for
    high-concurrency production use, swap this for a DB sequence or
    SELECT ... FOR UPDATE.
    """
    current_max = db.query(sa_func.max(BugSubmission.bug_number)).scalar()
    return (current_max or 0) + 1


def _format_bug_id(bug_number: int) -> str:
    return f"{BUG_ID_PREFIX}{bug_number:0{BUG_ID_PAD}d}"


def _serialize(bug: BugSubmission) -> dict:
    return {
        "id": bug.id,
        "bug_id": bug.bug_id,
        "title": bug.title,
        "severity": bug.severity.value if hasattr(bug.severity, "value") else bug.severity,
        "source_type": bug.source_type.value if hasattr(bug.source_type, "value") else bug.source_type,
        "original_filename": bug.original_filename,
        "file_type": bug.file_type,
        "content_length": bug.content_length,
        "created_at": bug.created_at,
    }


# ---------------------------------------------------------------------------
# AI Agent Layer — Milestone 2
#
# Both agents are implemented as deterministic, rule/heuristic-based
# classifiers (regex + keyword scoring) rather than calling out to an
# external LLM. This keeps Milestone 2 runnable with zero API keys / zero
# network dependency, matches the "everything in one file, zero setup"
# convention from Milestone 1, and gives fully explainable, reproducible
# output (required for the accuracy validation step below). The class
# boundary (`TriageAgent.classify`, `LogAnalysisAgent.analyze`) matches the
# "Agent Responsibilities" table in the design doc 1:1, so swapping the
# internals for a real LLM Service / RAG-backed call later (Milestone 3+)
# does not require touching the orchestrator or the API layer.
# ---------------------------------------------------------------------------

class TriageAgent:
    """
    Classifies a bug report's severity, priority, and affected component.

    Input:  raw bug report / stack trace text (str)
    Output: { severity, priority, type, component, confidence, reasoning }
    """

    SEVERITY_KEYWORDS = {
        Severity.CRITICAL: [
            "data loss", "corrupt", "security", "exploit", "vulnerability",
            "breach", "outage", "down", "crash", "crashed", "crashing",
            "deadlock", "unresponsive", "cannot start", "won't start",
            "production is down", "database is down", "payment failure",
            "segfault", "segmentation fault", "fatal", "kernel panic",
        ],
        Severity.HIGH: [
            "exception", "error", "fail", "failed", "failure", "broken",
            "blocked", "blocker", "500", "internal server error",
            "null pointer", "nullpointerexception", "stack trace",
            "traceback", "cannot login", "can't login", "not working",
            "regression", "timeout", "connection refused",
        ],
        Severity.MEDIUM: [
            "warning", "warn", "intermittent", "sometimes", "occasionally",
            "slow", "performance", "deprecated", "inconsistent",
            "unexpected behavior", "edge case", "workaround",
        ],
        Severity.LOW: [
            "typo", "cosmetic", "misaligned", "misalignment", "ui glitch",
            "minor", "spacing", "color", "colour", "tooltip", "wording",
            "nice to have", "enhancement",
        ],
    }

    SEVERITY_WEIGHT = {
        Severity.CRITICAL: 4,
        Severity.HIGH: 3,
        Severity.MEDIUM: 2,
        Severity.LOW: 1,
    }

    # Structural signals independent of keywords.
    STRUCTURAL_CRITICAL_PATTERNS = [
        r"Traceback \(most recent call last\)",           # unhandled Python exception
        r"Exception in thread",                             # unhandled Java exception
        r"Segmentation fault",
        r"PANIC|panic:",
    ]

    COMPONENT_KEYWORDS = {
        "Authentication": ["login", "auth", "token", "session", "password", "oauth", "jwt"],
        "Payments": ["payment", "checkout", "stripe", "billing", "invoice", "charge", "refund"],
        "Database": ["database", "sql", "query", "postgres", "sqlite", "orm", "connection pool", "deadlock"],
        "API": ["api", "endpoint", "request", "response", "500", "404", "rest", "http"],
        "Frontend/UI": ["ui", "button", "css", "layout", "render", "dom", "browser", "misaligned", "tooltip"],
        "File Upload": ["upload", "file", "attachment", "pdf", "docx", "multipart"],
        "Networking": ["timeout", "connection refused", "network", "socket", "dns"],
        "Infrastructure": ["deploy", "docker", "kubernetes", "server", "outage", "down", "crash"],
    }

    TYPE_KEYWORDS = {
        "Enhancement": ["enhancement", "feature request", "would be nice", "nice to have", "improve"],
        "Others": ["question", "how do i", "documentation", "docs"],
    }

    def classify(self, content: str) -> Dict[str, Any]:
        text = content or ""
        lower = text.lower()
        reasoning_parts: List[str] = []

        # --- severity scoring -------------------------------------------------
        scores: Dict[Severity, int] = {s: 0 for s in Severity}
        matched_keywords: Dict[Severity, List[str]] = {s: [] for s in Severity}

        for sev, keywords in self.SEVERITY_KEYWORDS.items():
            for kw in keywords:
                if kw in lower:
                    scores[sev] += 1
                    matched_keywords[sev].append(kw)

        structural_hit = False
        for pattern in self.STRUCTURAL_CRITICAL_PATTERNS:
            if re.search(pattern, text):
                scores[Severity.HIGH] += 2
                structural_hit = True
                reasoning_parts.append(f"structural signal matched pattern '{pattern}'")

        # Weighted pick: prefer the severity with the highest (score * weight).
        # Tie-break on raw match count (more specific evidence wins), then on
        # severity rank itself so ties still favor the more severe bucket.
        # A single incidental keyword hit (e.g. "fail" in a cosmetic report)
        # is treated as weak signal and should NOT override a Medium default.
        best_severity = Severity.MEDIUM
        best_weighted = -1
        best_matches = -1
        total_matches = 0

        severity_rank = {
            Severity.LOW: 0,
            Severity.MEDIUM: 1,
            Severity.HIGH: 2,
            Severity.CRITICAL: 3,
        }

        for sev, score in scores.items():
            total_matches += score
            weighted = score * self.SEVERITY_WEIGHT[sev]
            if score == 0:
                continue
            if (
                weighted > best_weighted
                or (weighted == best_weighted and score > best_matches)
                or (weighted == best_weighted and score == best_matches
                    and severity_rank[sev] > severity_rank[best_severity])
            ):
                best_weighted = weighted
                best_matches = score
                best_severity = sev

        # Require a minimum strength of evidence before trusting anything
        # above Medium. One incidental keyword match (score=1, e.g. the
        # word "fail" showing up once) is too weak to justify High/Critical
        # on its own — fall back to Medium unless there's a structural hit
        # or multiple corroborating keyword matches.
        MIN_SCORE_FOR_ELEVATED = 2
        if total_matches == 0:
            best_severity = Severity.MEDIUM
            reasoning_parts.append("no strong severity keywords found; defaulted to Medium")
        elif (
            best_severity in (Severity.HIGH, Severity.CRITICAL)
            and best_matches < MIN_SCORE_FOR_ELEVATED
            and not structural_hit
        ):
            reasoning_parts.append(
                f"only a single weak '{best_severity.value}' signal found "
                f"({matched_keywords[best_severity]}); downgraded to Medium "
                f"since evidence was insufficient"
            )
            best_severity = Severity.MEDIUM
        else:
            kws = matched_keywords[best_severity]
            if kws:
                reasoning_parts.append(
                    f"matched {len(kws)} {best_severity.value}-severity keyword(s): "
                    + ", ".join(sorted(set(kws))[:5])
                )

        # --- confidence ---------------------------------------------------
        # Ratio of the winning severity's matches to total matches, nudged
        # up for a direct structural hit (e.g. a raw traceback is a strong,
        # unambiguous signal), bounded to [0.35, 0.97] so the agent never
        # claims false certainty on sparse input.
        if total_matches > 0:
            confidence = matched_keywords[best_severity].__len__() / max(total_matches, 1)
            confidence = 0.5 + confidence * 0.4
        else:
            confidence = 0.4
        if structural_hit and best_severity in (Severity.HIGH, Severity.CRITICAL):
            confidence = min(confidence + 0.15, 0.97)
        confidence = round(min(max(confidence, 0.35), 0.97), 2)

        # --- priority (derived from severity, informed by urgency words) --
        urgency_words = ["urgent", "asap", "immediately", "blocking release", "blocker"]
        urgent_hit = any(w in lower for w in urgency_words)
        priority_map = {
            Severity.CRITICAL: "P0",
            Severity.HIGH: "P1",
            Severity.MEDIUM: "P2",
            Severity.LOW: "P3",
        }
        priority = priority_map[best_severity]
        if urgent_hit and priority in ("P1", "P2"):
            # bump priority by one level, but never override severity itself
            priority = "P0" if priority == "P1" else "P1"
            reasoning_parts.append("urgency language detected; priority bumped one level")

        # --- component ------------------------------------------------------
        component_scores: Dict[str, int] = {c: 0 for c in self.COMPONENT_KEYWORDS}
        for comp, keywords in self.COMPONENT_KEYWORDS.items():
            for kw in keywords:
                if kw in lower:
                    component_scores[comp] += 1
        best_component = max(component_scores, key=component_scores.get)
        if component_scores[best_component] == 0:
            best_component = "Unclassified"
            reasoning_parts.append("no component keywords matched; marked Unclassified")
        else:
            reasoning_parts.append(
                f"component inferred from keyword(s) associated with '{best_component}'"
            )

        # --- type -------------------------------------------------------
        bug_type = "Bug"
        for t, keywords in self.TYPE_KEYWORDS.items():
            if any(kw in lower for kw in keywords):
                bug_type = t
                break

        return {
            "severity": best_severity.value,
            "priority": priority,
            "type": bug_type,
            "component": best_component,
            "confidence": confidence,
            "reasoning": "; ".join(reasoning_parts) if reasoning_parts else "insufficient signal in report",
        }


class LogAnalysisAgent:
    """
    Parses stack traces / error logs and extracts structured error
    signatures.

    Input:  raw bug report / stack trace text (str)
    Output: { exception_type, message, failure_point, affected_code_paths,
              language, error_signature, confidence, reasoning }
    """

    CODE_FILE_PATTERN = re.compile(
        r"([\w./\\-]+\.(?:py|js|ts|jsx|tsx|java|go|rb|cpp|c|cs|php|kt|rs)):(\d+)"
    )

    def analyze(self, content: str) -> Dict[str, Any]:
        text = content or ""
        reasoning_parts: List[str] = []

        language = None
        exception_type = None
        message = None
        failure_file = None
        failure_line = None

        if re.search(r"Traceback \(most recent call last\)", text):
            language = "Python"
            reasoning_parts.append("detected Python traceback header")
            file_matches = re.findall(r'File "([^"]+)", line (\d+)', text)
            if file_matches:
                failure_file, failure_line = file_matches[-1]
                reasoning_parts.append("failure point taken as last frame in traceback")
            exc_match = re.search(r"^([A-Za-z_.]+(?:Error|Exception|Warning))\s*:\s*(.*)$", text, re.MULTILINE)
            if exc_match:
                exception_type, message = exc_match.group(1), exc_match.group(2).strip()

        elif re.search(r"Exception in thread|at [\w$.<>]+\(.*\.java:\d+\)", text):
            language = "Java"
            reasoning_parts.append("detected Java stack trace signature")
            m = re.search(r"\(([\w.$]+\.java):(\d+)\)", text)
            if m:
                failure_file, failure_line = m.group(1), m.group(2)
            exc_match = re.search(r"([\w.]+(?:Exception|Error))(?::\s*(.*))?", text)
            if exc_match:
                exception_type = exc_match.group(1)
                message = (exc_match.group(2) or "").strip() or None

        elif re.search(r"at .*\(.*:\d+:\d+\)|Uncaught|TypeError|ReferenceError", text):
            language = "JavaScript"
            reasoning_parts.append("detected JS/Node runtime error signature")
            # Prefer the file:line:col captured inside "(...)" so a greedy
            # leading ".*" in the caller-name portion can't eat into the
            # filename itself (e.g. "app.js" -> false match "p.js").
            m = re.search(r"\(([^()]+\.(?:js|ts|jsx|tsx)):(\d+):(\d+)\)", text)
            if not m:
                m = re.search(r"([\w./\\-]+\.(?:js|ts|jsx|tsx)):(\d+):(\d+)", text)
            if m:
                failure_file, failure_line = m.group(1), m.group(2)
            # "Uncaught TypeError: msg" / "TypeError: msg" anywhere in the
            # text, not just at line-start (browser console output often
            # prefixes with "Uncaught ").
            exc_match = re.search(r"(?:Uncaught\s+)?([A-Za-z]+(?:Error|Exception))\s*:\s*(.*)", text)
            if exc_match:
                exception_type = exc_match.group(1)
                message = exc_match.group(2).strip().split("\n")[0] or None

        elif re.search(r"^\d{4}-\d{2}-\d{2}[T ]\d{2}:\d{2}:\d{2}", text, re.MULTILINE):
            language = "log"
            reasoning_parts.append("detected structured timestamped log format")
            lvl = re.search(r"\b(ERROR|WARN|FATAL|CRITICAL)\b", text)
            if lvl:
                exception_type = f"{lvl.group(1)} log entry"

        else:
            reasoning_parts.append("no recognized stack trace signature; treated as unstructured report")

        # Generic fallback for failure_file if language-specific patterns missed it.
        if not failure_file:
            m = self.CODE_FILE_PATTERN.search(text)
            if m:
                failure_file, failure_line = m.group(1), m.group(2)
                reasoning_parts.append("failure point recovered via generic file:line pattern")

        # Affected code path(s): every distinct source file mentioned, in
        # order of appearance, capped to keep the payload small.
        affected_paths = []
        for m in self.CODE_FILE_PATTERN.finditer(text):
            path = m.group(1)
            if path not in affected_paths:
                affected_paths.append(path)
        # Python tracebacks use `File "path", line N` rather than `path:N`,
        # so the generic pattern above won't catch them — fold in every
        # `File "..."` reference too, plus the resolved failure_file.
        for m in re.finditer(r'File "([^"]+)"', text):
            path = m.group(1)
            if path not in affected_paths:
                affected_paths.append(path)
        if failure_file and failure_file not in affected_paths:
            affected_paths.insert(0, failure_file)
        affected_paths = affected_paths[:10]

        # Error signature: stable hash of exception type + normalized
        # message, for Milestone 3's Duplicate Detection Agent to key off
        # of before semantic/embedding search is wired in.
        sig_source = f"{exception_type or 'unknown'}::{(message or '').strip().lower()[:200]}"
        error_signature = hashlib.sha256(sig_source.encode("utf-8")).hexdigest()[:16]

        confidence = 0.4
        if language and language != "log":
            confidence += 0.25
        if exception_type:
            confidence += 0.2
        if failure_file:
            confidence += 0.1
        confidence = round(min(confidence, 0.97), 2)

        if not text.strip():
            confidence = 0.0
            reasoning_parts = ["empty content — nothing to analyze"]

        return {
            "exception_type": exception_type,
            "message": message,
            "language": language,
            "failure_file": failure_file,
            "failure_line": failure_line,
            "affected_paths": affected_paths,
            "error_signature": error_signature,
            "confidence": confidence,
            "reasoning": "; ".join(reasoning_parts),
        }


triage_agent = TriageAgent()
log_analysis_agent = LogAnalysisAgent()


def run_orchestration(content: str) -> Dict[str, Any]:
    """
    Multi-Agent Orchestrator (Milestone 2 scope): runs the Triage Agent and
    Log Analysis Agent on submission and combines their outputs into one
    structured payload. Both agents currently run independently off the raw
    content; Milestone 2's orchestration flow (design doc, section 3) has
    Log Analysis feed error signatures onward to Duplicate Detection /
    Root Cause in Milestone 3 — this function is the seam where that
    downstream wiring plugs in later without changing the API layer.
    """
    triage_result = triage_agent.classify(content)
    log_result = log_analysis_agent.analyze(content)
    return {
        "triage": triage_result,
        "log_analysis": log_result,
        "analyzed_at": datetime.utcnow().isoformat(),
    }


def _persist_analysis(db: Session, bug: BugSubmission, combined: Dict[str, Any]) -> AgentAnalysis:
    triage = combined["triage"]
    log_res = combined["log_analysis"]

    record = AgentAnalysis(
        bug_id=bug.id,
        triage_severity=triage["severity"],
        triage_priority=triage["priority"],
        triage_type=triage["type"],
        triage_component=triage["component"],
        triage_confidence=triage["confidence"],
        triage_reasoning=triage["reasoning"],
        log_exception_type=log_res["exception_type"],
        log_message=log_res["message"],
        log_language=log_res["language"],
        log_failure_file=log_res["failure_file"],
        log_failure_line=log_res["failure_line"],
        log_affected_paths=json.dumps(log_res["affected_paths"]),
        log_error_signature=log_res["error_signature"],
        log_confidence=log_res["confidence"],
        log_reasoning=log_res["reasoning"],
        combined_json=json.dumps({"bug_id": bug.bug_id, **combined}),
    )
    db.add(record)
    db.commit()
    db.refresh(record)
    return record


def _serialize_analysis(record: AgentAnalysis, bug_id: str) -> Dict[str, Any]:
    return {
        "bug_id": bug_id,
        "triage": {
            "severity": record.triage_severity,
            "priority": record.triage_priority,
            "type": record.triage_type,
            "component": record.triage_component,
            "confidence": record.triage_confidence,
            "reasoning": record.triage_reasoning,
        },
        "log_analysis": {
            "exception_type": record.log_exception_type,
            "message": record.log_message,
            "language": record.log_language,
            "failure_file": record.log_failure_file,
            "failure_line": record.log_failure_line,
            "affected_paths": json.loads(record.log_affected_paths or "[]"),
            "error_signature": record.log_error_signature,
            "confidence": record.log_confidence,
            "reasoning": record.log_reasoning,
        },
        "analyzed_at": record.created_at,
    }


# ---------------------------------------------------------------------------
# Frontend (served inline — no separate static files/build step)
# ---------------------------------------------------------------------------

FRONTEND_HTML = '<!DOCTYPE html>\n<html lang="en">\n<head>\n<meta charset="UTF-8" />\n<meta name="viewport" content="width=device-width, initial-scale=1.0" />\n<title>Bug Submission — AI Smart Bug Analyzer &amp; Fix Advisor</title>\n<style>\n  @import url(\'https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;500;600&family=IBM+Plex+Sans:wght@400;500;600&display=swap\');\n\n  :root{\n    --ink:#000000;\n    --panel:#181C25;\n    --panel-2:#1F2430;\n    --panel-3:#242A38;\n    --line:#2B303C;\n    --line-soft:#232833;\n    --text:#E7EAF0;\n    --muted:#8992A6;\n    --muted-dim:#5C6478;\n    --accent:#57C7D4;\n    --accent-dim:rgba(87,199,212,0.12);\n    --accent-line:rgba(87,199,212,0.35);\n    --danger:#E2574C;\n    --danger-dim:rgba(226,87,76,0.12);\n    --warning:#E0A438;\n    --warning-dim:rgba(224,164,56,0.12);\n    --success:#48B884;\n    --success-dim:rgba(72,184,132,0.12);\n  }\n\n  *{box-sizing:border-box;}\n  html,body{margin:0;padding:0;}\n  body{\n    background:var(--ink);\n    color:var(--text);\n    font-family:\'IBM Plex Sans\', sans-serif;\n    font-size:15px;\n    line-height:1.6;\n    -webkit-font-smoothing:antialiased;\n  }\n  .mono{font-family:\'IBM Plex Mono\', monospace;}\n  .wrap{ max-width:1040px; margin:0 auto; padding:40px 24px 64px; }\n\n  .eyebrow{\n    font-family:\'IBM Plex Mono\', monospace; font-size:11px; letter-spacing:0.14em;\n    text-transform:uppercase; color:var(--muted-dim); margin:0 0 10px;\n  }\n  .header-row{ display:flex; align-items:flex-end; justify-content:space-between; gap:16px; margin-bottom:6px; flex-wrap:wrap; }\n  h1{ font-family:\'IBM Plex Mono\', monospace; font-size:26px; font-weight:600; letter-spacing:-0.01em; margin:0; }\n  .subtitle{ color:var(--muted); font-size:14.5px; margin:10px 0 32px; max-width:560px; }\n  .btn-clear{\n    background:transparent; border:1px solid var(--line); color:var(--muted);\n    font-family:\'IBM Plex Mono\', monospace; font-size:12px; letter-spacing:0.04em;\n    padding:8px 14px; border-radius:6px; cursor:pointer; transition:border-color .15s ease, color .15s ease;\n  }\n  .btn-clear:hover{ border-color:var(--muted-dim); color:var(--text); }\n\n  .grid{ display:grid; grid-template-columns:1.55fr 1fr; gap:20px; align-items:start; }\n  @media (max-width:820px){ .grid{ grid-template-columns:1fr; } }\n\n  .panel{ background:var(--panel); border:1px solid var(--line); border-radius:10px; padding:22px; }\n  .label{\n    font-family:\'IBM Plex Mono\', monospace; font-size:11px; letter-spacing:0.1em;\n    text-transform:uppercase; color:var(--muted-dim); display:block; margin-bottom:8px;\n  }\n\n  input[type="text"], select, textarea{\n    width:100%; background:var(--panel-2); border:1px solid var(--line); color:var(--text);\n    border-radius:7px; padding:11px 13px; font-family:\'IBM Plex Sans\', sans-serif; font-size:14px;\n    outline:none; transition:border-color .15s ease, box-shadow .15s ease;\n  }\n  input[type="text"]:focus, select:focus, textarea:focus{\n    border-color:var(--accent-line); box-shadow:0 0 0 3px var(--accent-dim);\n  }\n  textarea{\n    font-family:\'IBM Plex Mono\', monospace; font-size:13px; line-height:1.65;\n    min-height:190px; resize:vertical; white-space:pre;\n  }\n  textarea::placeholder, input::placeholder{ color:var(--muted-dim); }\n\n  .field{ margin-bottom:20px; }\n  .field-row{ display:flex; gap:14px; }\n  .field-row .field{ flex:1; margin-bottom:0; }\n\n\n  hr.divider{ border:none; border-top:1px solid var(--line-soft); margin:22px 0; }\n\n  .dropzone{\n    border:1.5px dashed var(--line); border-radius:9px; padding:26px 20px; text-align:center;\n    cursor:pointer; transition:border-color .15s ease, background .15s ease;\n  }\n  .dropzone.drag{ border-color:var(--accent); background:var(--accent-dim); }\n  .dropzone .dz-icon{ font-family:\'IBM Plex Mono\', monospace; font-size:20px; color:var(--muted-dim); margin-bottom:8px; }\n  .dropzone .dz-title{ font-size:14px; color:var(--text); margin-bottom:4px; }\n  .dropzone .dz-sub{ font-size:12.5px; color:var(--muted-dim); font-family:\'IBM Plex Mono\', monospace; }\n  input[type="file"]{ display:none; }\n\n  .file-list{ margin-top:14px; display:flex; flex-direction:column; gap:8px; }\n  .file-chip{\n    display:flex; align-items:center; gap:10px; background:var(--panel-2); border:1px solid var(--line);\n    border-radius:7px; padding:9px 11px; font-size:13px;\n  }\n  .file-chip.file-error{ border-color:rgba(226,87,76,0.4); }\n  .file-chip .fc-ext{\n    font-family:\'IBM Plex Mono\', monospace; font-size:10px; letter-spacing:0.04em;\n    background:var(--panel-3); color:var(--muted); border-radius:4px; padding:3px 6px; flex-shrink:0;\n  }\n  .file-chip .fc-name{ flex:1; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }\n  .file-chip .fc-meta{ color:var(--muted-dim); font-size:11.5px; font-family:\'IBM Plex Mono\', monospace; flex-shrink:0; }\n  .file-chip .fc-error-text{ color:var(--danger); font-size:11.5px; font-family:\'IBM Plex Mono\', monospace; flex-shrink:0; }\n  .file-chip .fc-remove{\n    background:none; border:none; color:var(--muted-dim); cursor:pointer;\n    font-size:15px; line-height:1; padding:2px 4px; border-radius:4px;\n  }\n  .file-chip .fc-remove:hover{ color:var(--danger); background:var(--danger-dim); }\n\n  .submit-btn{\n    width:100%; margin-top:22px; background:var(--accent); color:#0A2C31; border:none; border-radius:8px;\n    padding:13px; font-family:\'IBM Plex Mono\', monospace; font-size:13.5px; font-weight:600;\n    letter-spacing:0.02em; cursor:pointer; transition:filter .15s ease, opacity .15s ease;\n  }\n  .submit-btn:hover{ filter:brightness(1.08); }\n  .submit-btn:disabled{ opacity:0.45; cursor:not-allowed; filter:none; }\n\n  .hint{ font-size:12px; color:var(--muted-dim); margin-top:9px; text-align:center; }\n  .hint.error-hint{ color:var(--danger); }\n\n  .preview-panel{ position:sticky; top:20px; }\n  .preview-head{ display:flex; align-items:center; justify-content:space-between; margin-bottom:16px; }\n  .status-pill{\n    display:inline-flex; align-items:center; gap:7px; font-family:\'IBM Plex Mono\', monospace;\n    font-size:11px; letter-spacing:0.05em; text-transform:uppercase; color:var(--muted-dim);\n  }\n  .status-dot{ width:7px; height:7px; border-radius:50%; background:var(--muted-dim); }\n  .status-dot.live{ background:var(--accent); animation:pulse 1.6s ease-in-out infinite; }\n  @keyframes pulse{ 0%,100%{ box-shadow:0 0 0 0 var(--accent-dim); } 50%{ box-shadow:0 0 0 5px transparent; } }\n\n  .empty-state{ padding:30px 6px; text-align:center; color:var(--muted-dim); font-size:13px; }\n  .empty-state .ee-icon{ font-family:\'IBM Plex Mono\', monospace; font-size:22px; margin-bottom:10px; opacity:0.6; }\n\n  .badge{\n    display:inline-block; font-family:\'IBM Plex Mono\', monospace; font-size:11.5px; letter-spacing:0.03em;\n    padding:5px 10px; border-radius:5px; background:var(--accent-dim); color:var(--accent);\n    border:1px solid var(--accent-line); margin-bottom:16px;\n  }\n\n  .kv{ margin-bottom:14px; }\n  .kv-label{\n    font-family:\'IBM Plex Mono\', monospace; font-size:10.5px; text-transform:uppercase;\n    letter-spacing:0.08em; color:var(--muted-dim); margin-bottom:4px;\n  }\n  .kv-value{ font-family:\'IBM Plex Mono\', monospace; font-size:13px; color:var(--text); word-break:break-word; line-height:1.5; }\n  .kv-value.danger-text{ color:var(--danger); }\n\n  .preview-fade{ animation:fadeIn .25s ease; }\n  @keyframes fadeIn{ from{ opacity:0; transform:translateY(3px); } to{ opacity:1; transform:translateY(0); } }\n\n  .attached-mini{ margin-top:6px; display:flex; flex-direction:column; gap:6px; }\n  .attached-mini-item{\n    display:flex; justify-content:space-between; font-family:\'IBM Plex Mono\', monospace; font-size:11.5px;\n    color:var(--muted); border-top:1px solid var(--line-soft); padding-top:6px;\n  }\n\n  .confirm-banner{\n    margin-top:20px; background:var(--success-dim); border:1px solid rgba(72,184,132,0.35);\n    border-radius:9px; padding:15px 16px; display:none;\n  }\n  .confirm-banner.show{ display:block; animation:fadeIn .25s ease; }\n  .confirm-banner.error-banner{ background:var(--danger-dim); border-color:rgba(226,87,76,0.35); }\n  .confirm-title{ font-family:\'IBM Plex Mono\', monospace; font-size:13px; color:var(--success); margin-bottom:4px; }\n  .confirm-banner.error-banner .confirm-title{ color:var(--danger); }\n  .confirm-sub{ font-size:12.5px; color:var(--muted); }\n\n  .history{ margin-top:36px; }\n  .history-head{ display:flex; align-items:center; justify-content:space-between; margin-bottom:14px; }\n  .history-title{\n    font-family:\'IBM Plex Mono\', monospace; font-size:11px; letter-spacing:0.1em;\n    text-transform:uppercase; color:var(--muted-dim);\n  }\n  .btn-refresh{\n    background:transparent; border:1px solid var(--line); color:var(--muted);\n    font-family:\'IBM Plex Mono\', monospace; font-size:11px; letter-spacing:0.04em;\n    padding:5px 10px; border-radius:5px; cursor:pointer;\n  }\n  .btn-refresh:hover{ border-color:var(--muted-dim); color:var(--text); }\n  .history-empty{ color:var(--muted-dim); font-size:13px; padding:16px 0; border-top:1px solid var(--line-soft); }\n  .history-item{\n    display:flex; align-items:center; gap:14px; padding:12px 0; border-top:1px solid var(--line-soft); cursor:pointer;\n  }\n  .history-item:hover .history-title-txt{ color:var(--accent); }\n  .history-id{\n    font-family:\'IBM Plex Mono\', monospace; font-size:12px; color:var(--accent); flex-shrink:0;\n    width:150px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap;\n  }\n  .history-title-txt{ flex:1; font-size:13.5px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; transition:color .15s ease; }\n  .history-sev{\n    font-family:\'IBM Plex Mono\', monospace; font-size:10.5px; text-transform:uppercase;\n    letter-spacing:0.04em; padding:3px 8px; border-radius:5px; flex-shrink:0;\n  }\n  .history-time{ font-family:\'IBM Plex Mono\', monospace; font-size:11.5px; color:var(--muted-dim); flex-shrink:0; }\n\n  .modal{ position:fixed; inset:0; background:rgba(0,0,0,0.6); display:flex; align-items:center; justify-content:center; padding:20px; z-index:50; }\n  .modal.hidden{ display:none; }\n  .modal-content{\n    background:var(--panel); border:1px solid var(--line); border-radius:10px; max-width:700px; width:100%;\n    max-height:80vh; overflow-y:auto; padding:24px; position:relative;\n  }\n  .modal-close{ position:absolute; top:12px; right:16px; background:none; border:none; color:var(--muted-dim); font-size:1.4rem; cursor:pointer; }\n  .modal-meta{ color:var(--muted); font-size:12.5px; font-family:\'IBM Plex Mono\', monospace; }\n  .modal-body{\n    white-space:pre-wrap; word-break:break-word; background:var(--panel-2); border:1px solid var(--line);\n    border-radius:8px; padding:14px; font-family:\'IBM Plex Mono\', monospace; font-size:12.5px; margin-top:12px;\n  }\n  .agent-section{ margin-top:18px; border-top:1px solid var(--line-soft); padding-top:16px; }\n  .agent-section-title{\n    font-family:\'IBM Plex Mono\', monospace; font-size:11px; letter-spacing:0.1em; text-transform:uppercase;\n    color:var(--accent); margin-bottom:10px;\n  }\n  .agent-grid{ display:grid; grid-template-columns:1fr 1fr; gap:10px 18px; margin-bottom:8px; }\n  .agent-kv-label{ font-family:\'IBM Plex Mono\', monospace; font-size:10.5px; text-transform:uppercase; letter-spacing:0.06em; color:var(--muted-dim); }\n  .agent-kv-value{ font-family:\'IBM Plex Mono\', monospace; font-size:12.5px; color:var(--text); word-break:break-word; }\n  .agent-reasoning{ font-size:12.5px; color:var(--muted); margin-top:8px; line-height:1.5; }\n  .confidence-bar{ height:5px; border-radius:3px; background:var(--panel-3); margin-top:4px; overflow:hidden; }\n  .confidence-fill{ height:100%; background:var(--accent); }\n</style>\n</head>\n<body>\n<div class="wrap">\n\n  <p class="eyebrow">AI Smart Bug Analyzer &amp; Fix Advisor</p>\n  <div class="header-row">\n    <h1>Bug submission</h1>\n    <button class="btn-clear" id="clearBtn" type="button">Clear form</button>\n  </div>\n  <p class="subtitle">Paste an error, or drop a log file. The Triage Agent predicts severity/priority/component and the Log Analysis Agent parses the error automatically on submit.</p>\n\n  <div class="grid">\n\n    <div>\n      <div class="panel">\n\n        <div class="field-row">\n          <div class="field">\n            <label class="label" for="title">Title</label>\n            <input type="text" id="title" placeholder="Checkout fails with 500 on submit" />\n          </div>\n        </div>\n\n        <hr class="divider" />\n\n        <div class="field">\n          <label class="label" for="pasteArea">Paste report, stack trace, or log</label>\n          <textarea id="pasteArea" placeholder="Traceback (most recent call last):&#10;  File &quot;app/views.py&quot;, line 88, in checkout&#10;    charge = stripe.Charge.create(...)&#10;stripe.error.CardError: Your card was declined."></textarea>\n        </div>\n\n        <div class="field">\n          <label class="label">Attach files</label>\n          <div class="dropzone" id="dropzone">\n            <div class="dz-icon">&#8595;</div>\n            <div class="dz-title">Drop files here, or click to browse</div>\n            <div class="dz-sub">.log .txt .json .xml .csv .stacktrace .out .pdf .docx — up to 10 files, 5MB each</div>\n          </div>\n          <input type="file" id="fileInput" multiple accept=".txt,.log,.json,.xml,.csv,.stacktrace,.out,.pdf,.docx" />\n          <div class="file-list" id="fileList"></div>\n        </div>\n\n        <button class="submit-btn" id="submitBtn" disabled>Submit for analysis</button>\n        <div class="hint" id="submitHint">Add pasted text or at least one file to enable submission.</div>\n\n        <div class="confirm-banner" id="confirmBanner">\n          <div class="confirm-title" id="confirmTitle"></div>\n          <div class="confirm-sub" id="confirmSub"></div>\n        </div>\n\n      </div>\n\n      <div class="history">\n        <div class="history-head">\n          <div class="history-title">Submission history</div>\n          <button class="btn-refresh" id="refreshHistoryBtn" type="button">Refresh</button>\n        </div>\n        <div id="historyList">\n          <div class="history-empty">Loading submissions...</div>\n        </div>\n      </div>\n    </div>\n\n    <div class="preview-panel">\n      <div class="panel">\n        <div class="preview-head">\n          <span class="label" style="margin-bottom:0;">Diagnostic preview</span>\n          <span class="status-pill"><span class="status-dot" id="statusDot"></span><span id="statusText">Idle</span></span>\n        </div>\n        <div id="previewBody">\n          <div class="empty-state">\n            <div class="ee-icon">&#10022;</div>\n            Paste an error or drop a file to begin analysis.\n          </div>\n        </div>\n      </div>\n    </div>\n\n  </div>\n</div>\n\n<div id="detailModal" class="modal hidden">\n  <div class="modal-content">\n    <button id="closeModal" class="modal-close" type="button">&times;</button>\n    <h3 id="modalTitle" class="mono" style="margin:0 0 6px;"></h3>\n    <p id="modalMeta" class="modal-meta"></p>\n    <pre id="modalContent" class="modal-body"></pre>\n    <div id="modalAgentSection"></div>\n  </div>\n</div>\n\n<script>\n(function(){\n  const API_BASE = "";\n\n  const pasteArea = document.getElementById(\'pasteArea\');\n  const dropzone = document.getElementById(\'dropzone\');\n  const fileInput = document.getElementById(\'fileInput\');\n  const fileList = document.getElementById(\'fileList\');\n  const submitBtn = document.getElementById(\'submitBtn\');\n  const submitHint = document.getElementById(\'submitHint\');\n  const previewBody = document.getElementById(\'previewBody\');\n  const statusDot = document.getElementById(\'statusDot\');\n  const statusText = document.getElementById(\'statusText\');\n  const confirmBanner = document.getElementById(\'confirmBanner\');\n  const confirmTitle = document.getElementById(\'confirmTitle\');\n  const confirmSub = document.getElementById(\'confirmSub\');\n  const historyList = document.getElementById(\'historyList\');\n  const titleInput = document.getElementById(\'title\');\n  const clearBtn = document.getElementById(\'clearBtn\');\n  const refreshHistoryBtn = document.getElementById(\'refreshHistoryBtn\');\n\n  const MAX_FILES = 10;\n  const MAX_FILE_SIZE = 5 * 1024 * 1024;\n  const ALLOWED_EXT = [\'txt\',\'log\',\'json\',\'xml\',\'csv\',\'stacktrace\',\'out\',\'pdf\',\'docx\'];\n  const BINARY_EXT = [\'pdf\',\'docx\'];\n\n  let attachedFiles = [];\n  let isSubmitting = false;\n\n  function analyzeContent(text){\n    const result = {\n      type: null, exception: null, message: null, file: null, line: null,\n      language: null, lines: text.split(\'\\n\').length, chars: text.length\n    };\n    if (!text || !text.trim()) return null;\n\n    if (/Traceback \\(most recent call last\\)/.test(text)) {\n      result.type = \'Python traceback\';\n      result.language = \'Python\';\n      const fileMatch = text.match(/File "([^"]+)", line (\\d+)/g);\n      if (fileMatch && fileMatch.length) {\n        const m = fileMatch[fileMatch.length - 1].match(/File "([^"]+)", line (\\d+)/);\n        if (m) { result.file = m[1]; result.line = m[2]; }\n      }\n      const excMatch = text.match(/^([A-Za-z_.]+(?:Error|Exception|Warning))\\s*:\\s*(.*)$/m);\n      if (excMatch) { result.exception = excMatch[1]; result.message = excMatch[2]; }\n    } else if (/at [\\w$.<>]+\\s*\\(?.*\\.java:\\d+\\)?/.test(text) || /Exception in thread/.test(text)) {\n      result.type = \'Java stack trace\';\n      result.language = \'Java\';\n      const m = text.match(/([\\w.$]+):(\\d+)\\)/) || text.match(/\\(([\\w.$]+\\.java):(\\d+)\\)/);\n      if (m) { result.file = m[1]; result.line = m[2]; }\n      const excMatch = text.match(/([\\w.]+(?:Exception|Error))(?::\\s*(.*))?/);\n      if (excMatch) { result.exception = excMatch[1]; result.message = excMatch[2] || null; }\n    } else if (/at .*\\(.*:\\d+:\\d+\\)/.test(text) || /Uncaught|TypeError|ReferenceError/.test(text)) {\n      result.type = \'JS/Node runtime error\';\n      result.language = \'JavaScript\';\n      const m = text.match(/at .*\\(?([\\w./\\\\-]+\\.(?:js|ts|jsx|tsx)):(\\d+):(\\d+)\\)?/);\n      if (m) { result.file = m[1]; result.line = m[2]; }\n      const excMatch = text.match(/^([A-Za-z]+Error)\\s*:\\s*(.*)$/m);\n      if (excMatch) { result.exception = excMatch[1]; result.message = excMatch[2]; }\n    } else if (/^\\d{4}-\\d{2}-\\d{2}[T ]\\d{2}:\\d{2}:\\d{2}/m.test(text)) {\n      result.type = \'Structured log\';\n      const lvl = text.match(/\\b(ERROR|WARN|FATAL|CRITICAL)\\b/);\n      if (lvl) result.exception = lvl[1] + \' entries detected\';\n    } else {\n      result.type = \'Unstructured report\';\n    }\n\n    if (!result.file) {\n      const generic = text.match(/([\\w./\\\\-]+\\.(?:py|js|ts|java|go|rb|cpp|c|cs|php)):(\\d+)/);\n      if (generic) { result.file = generic[1]; result.line = generic[2]; }\n    }\n    return result;\n  }\n\n  function extForFile(name){\n    const parts = name.split(\'.\');\n    return parts.length > 1 ? parts[parts.length-1].toUpperCase() : \'FILE\';\n  }\n  function extLower(name){\n    const parts = name.split(\'.\');\n    return parts.length > 1 ? parts[parts.length-1].toLowerCase() : \'\';\n  }\n  function formatSize(bytes){\n    if (bytes < 1024) return bytes + \' B\';\n    if (bytes < 1024*1024) return (bytes/1024).toFixed(1) + \' KB\';\n    return (bytes/(1024*1024)).toFixed(1) + \' MB\';\n  }\n  function escapeHtml(str){\n    return String(str).replace(/[&<>"\']/g, function(c){\n      return { \'&\':\'&amp;\', \'<\':\'&lt;\', \'>\':\'&gt;\', \'"\':\'&quot;\', "\'":\'&#39;\' }[c];\n    });\n  }\n\n  function renderFileList(){\n    fileList.innerHTML = \'\';\n    attachedFiles.forEach(function(f, idx){\n      const chip = document.createElement(\'div\');\n      chip.className = \'file-chip\' + (f.error ? \' file-error\' : \'\');\n      chip.innerHTML =\n        \'<span class="fc-ext">\' + extForFile(f.name) + \'</span>\' +\n        \'<span class="fc-name">\' + escapeHtml(f.name) + \'</span>\' +\n        (f.error\n          ? \'<span class="fc-error-text">\' + escapeHtml(f.error) + \'</span>\'\n          : \'<span class="fc-meta">\' + formatSize(f.size) + \'</span>\') +\n        \'<button class="fc-remove" type="button" aria-label="Remove file">&#10005;</button>\';\n      chip.querySelector(\'.fc-remove\').addEventListener(\'click\', function(){\n        attachedFiles.splice(idx, 1);\n        renderFileList();\n        updatePreview();\n      });\n      fileList.appendChild(chip);\n    });\n  }\n\n  function updatePreview(){\n    const text = pasteArea.value;\n    const pasteAnalysis = analyzeContent(text);\n    const hasFiles = attachedFiles.length > 0;\n    const hasValidFiles = attachedFiles.some(function(f){ return !f.error; });\n    const hasContent = (text && text.trim().length > 0) || hasValidFiles;\n\n    submitBtn.disabled = !hasContent || isSubmitting;\n    if (!isSubmitting) {\n      submitHint.textContent = hasContent ? \'Ready to submit.\' : \'Add pasted text or at least one file to enable submission.\';\n      submitHint.classList.remove(\'error-hint\');\n    }\n\n    if (!hasContent) {\n      statusDot.classList.remove(\'live\');\n      statusText.textContent = \'Idle\';\n      previewBody.innerHTML = \'<div class="empty-state"><div class="ee-icon">&#10022;</div>Paste an error or drop a file to begin analysis.</div>\';\n      return;\n    }\n\n    statusDot.classList.add(\'live\');\n    statusText.textContent = \'Reading\';\n\n    let html = \'\';\n    if (pasteAnalysis) {\n      html += \'<div class="badge preview-fade">\' + pasteAnalysis.type + \'</div>\';\n      if (pasteAnalysis.exception) {\n        html += \'<div class="kv preview-fade"><div class="kv-label">Exception</div><div class="kv-value danger-text">\' + escapeHtml(pasteAnalysis.exception) + \'</div></div>\';\n      }\n      if (pasteAnalysis.message) {\n        html += \'<div class="kv preview-fade"><div class="kv-label">Message</div><div class="kv-value">\' + escapeHtml(pasteAnalysis.message) + \'</div></div>\';\n      }\n      if (pasteAnalysis.file) {\n        html += \'<div class="kv preview-fade"><div class="kv-label">Location</div><div class="kv-value">\' + escapeHtml(pasteAnalysis.file) + (pasteAnalysis.line ? \':\' + pasteAnalysis.line : \'\') + \'</div></div>\';\n      }\n      if (pasteAnalysis.language) {\n        html += \'<div class="kv preview-fade"><div class="kv-label">Language</div><div class="kv-value">\' + pasteAnalysis.language + \'</div></div>\';\n      }\n      html += \'<div class="kv preview-fade"><div class="kv-label">Pasted content</div><div class="kv-value">\' + pasteAnalysis.lines + \' lines &middot; \' + pasteAnalysis.chars + \' chars</div></div>\';\n    } else if (!hasFiles) {\n      html += \'<div class="empty-state"><div class="ee-icon">&#10022;</div>Paste an error or drop a file to begin analysis.</div>\';\n    }\n\n    if (hasFiles) {\n      html += \'<div class="kv preview-fade"><div class="kv-label">Attached files (\' + attachedFiles.length + \')</div><div class="attached-mini">\';\n      attachedFiles.forEach(function(f){\n        let label;\n        if (f.error) {\n          label = f.error;\n        } else if (f.binary) {\n          label = \'Document file — parsed on server (preview not shown client-side)\';\n        } else if (f.content) {\n          label = (analyzeContent(f.content) || {type:\'Unreadable\'}).type;\n        } else {\n          label = \'Reading...\';\n        }\n        html += \'<div class="attached-mini-item"><span>\' + escapeHtml(f.name) + \'</span><span>\' + escapeHtml(label) + \'</span></div>\';\n      });\n      html += \'</div></div>\';\n    }\n\n    previewBody.innerHTML = html;\n  }\n\n  pasteArea.addEventListener(\'input\', updatePreview);\n\n  dropzone.addEventListener(\'click\', function(){ fileInput.click(); });\n  dropzone.addEventListener(\'dragover\', function(e){ e.preventDefault(); dropzone.classList.add(\'drag\'); });\n  dropzone.addEventListener(\'dragleave\', function(){ dropzone.classList.remove(\'drag\'); });\n  dropzone.addEventListener(\'drop\', function(e){\n    e.preventDefault();\n    dropzone.classList.remove(\'drag\');\n    handleFiles(e.dataTransfer.files);\n  });\n  fileInput.addEventListener(\'change\', function(e){\n    handleFiles(e.target.files);\n    fileInput.value = \'\';\n  });\n\n  function handleFiles(fileListObj){\n    const incoming = Array.from(fileListObj);\n\n    if (attachedFiles.length + incoming.length > MAX_FILES) {\n      submitHint.textContent = \'Max \' + MAX_FILES + \' files per submission. Remove some files first.\';\n      submitHint.classList.add(\'error-hint\');\n      return;\n    }\n\n    incoming.forEach(function(file){\n      const ext = extLower(file.name);\n      const isBinary = BINARY_EXT.includes(ext);\n      const entry = { file: file, name: file.name, size: file.size, content: null, error: null, binary: isBinary };\n\n      if (!ALLOWED_EXT.includes(ext)) {\n        entry.error = \'Unsupported type (.\' + ext + \')\';\n      } else if (file.size === 0) {\n        entry.error = \'File is empty\';\n      } else if (file.size > MAX_FILE_SIZE) {\n        entry.error = \'Too large (max 5MB)\';\n      }\n\n      attachedFiles.push(entry);\n\n      if (!entry.error && !isBinary) {\n        const reader = new FileReader();\n        reader.onload = function(ev){\n          entry.content = ev.target.result;\n          renderFileList();\n          updatePreview();\n        };\n        reader.readAsText(file);\n      }\n    });\n    renderFileList();\n    updatePreview();\n  }\n\n  submitBtn.addEventListener(\'click\', async function(){\n    if (submitBtn.disabled || isSubmitting) return;\n\n    const title = titleInput.value.trim();\n    const text = pasteArea.value.trim();\n    const validFiles = attachedFiles.filter(function(f){ return !f.error; });\n\n    if (!title) {\n      submitHint.textContent = \'Title is required.\';\n      submitHint.classList.add(\'error-hint\');\n      titleInput.focus();\n      return;\n    }\n    if (attachedFiles.some(function(f){ return f.error; })) {\n      submitHint.textContent = \'Remove or fix invalid files before submitting.\';\n      submitHint.classList.add(\'error-hint\');\n      return;\n    }\n    if (!text && validFiles.length === 0) {\n      submitHint.textContent = \'Add pasted text or at least one file to enable submission.\';\n      submitHint.classList.add(\'error-hint\');\n      return;\n    }\n\n    isSubmitting = true;\n    submitBtn.disabled = true;\n    submitBtn.textContent = \'Submitting...\';\n    statusText.textContent = \'Analyzing\';\n    hideConfirm();\n\n    try {\n      const results = [];\n\n      if (text) {\n        const fd = new FormData();\n        fd.append(\'title\', title);\n        fd.append(\'content\', text);\n        const res = await fetch(API_BASE + \'/api/bugs/paste\', { method: \'POST\', body: fd });\n        const data = await res.json();\n        if (!res.ok) throw new Error(data.detail || \'Paste submission failed.\');\n        results.push(data);\n      }\n\n      if (validFiles.length > 0) {\n        const fd = new FormData();\n        fd.append(\'title\', title);\n        validFiles.forEach(function(f){ fd.append(\'files\', f.file); });\n        const res = await fetch(API_BASE + \'/api/bugs/upload\', { method: \'POST\', body: fd });\n        const data = await res.json();\n        if (!res.ok) throw new Error(data.detail || \'File upload failed.\');\n        results.push(data);\n      }\n\n      const firstAnalysis = results.find(function(r){ return r.analysis; });\n      const bugIds = results.map(function(r){ return r.bug_id; }).join(\', \');\n\n      if (firstAnalysis && firstAnalysis.analysis) {\n        showConfirm(\n          false,\n          bugIds + \' analyzed\',\n          \'Predicted severity: \' + firstAnalysis.analysis.triage.severity + \' · \' + firstAnalysis.analysis.triage.component\n        );\n        statusDot.classList.add(\'live\');\n        statusText.textContent = \'Analyzed\';\n        previewBody.innerHTML =\n          \'<div class="badge preview-fade">\' + bugIds + \' — analysis complete</div>\' +\n          renderAgentSection(firstAnalysis.analysis);\n        resetForm(true);\n      } else {\n        showConfirm(false, results.length + \' submission(s) analyzed\', bugIds);\n        resetForm(false);\n      }\n\n      await loadHistory();\n\n    } catch (err) {\n      showConfirm(true, \'Submission failed\', err.message);\n    } finally {\n      isSubmitting = false;\n      submitBtn.textContent = \'Submit for analysis\';\n      updatePreview();\n    }\n  });\n\n  function showConfirm(isError, title, sub){\n    confirmBanner.classList.toggle(\'error-banner\', isError);\n    confirmTitle.textContent = title;\n    confirmSub.textContent = sub;\n    confirmBanner.classList.add(\'show\');\n    setTimeout(hideConfirm, 5000);\n  }\n  function hideConfirm(){\n    confirmBanner.classList.remove(\'show\');\n  }\n\n  function resetForm(keepPreview){\n    pasteArea.value = \'\';\n    attachedFiles = [];\n    titleInput.value = \'\';\n    renderFileList();\n    if (!keepPreview) {\n      updatePreview();\n    } else {\n      // Form fields are cleared, but leave previewBody showing the agent\n      // analysis result from the submission that just completed.\n      submitBtn.disabled = true;\n      submitHint.textContent = \'Add pasted text or at least one file to enable submission.\';\n      submitHint.classList.remove(\'error-hint\');\n    }\n  }\n\n  clearBtn.addEventListener(\'click\', function(){\n    resetForm();\n    submitHint.textContent = \'Add pasted text or at least one file to enable submission.\';\n    submitHint.classList.remove(\'error-hint\');\n  });\n\n  const sevColors = {\n    low: { bg: \'var(--success-dim)\', text: \'var(--success)\' },\n    medium: { bg: \'var(--warning-dim)\', text: \'var(--warning)\' },\n    high: { bg: \'rgba(232,135,63,0.12)\', text: \'#E8873F\' },\n    critical: { bg: \'var(--danger-dim)\', text: \'var(--danger)\' }\n  };\n\n  async function loadHistory(){\n    try {\n      const res = await fetch(API_BASE + \'/api/bugs\');\n      const bugs = await res.json();\n      if (!res.ok) throw new Error(bugs.detail || \'Failed to load history.\');\n\n      if (!bugs.length) {\n        historyList.innerHTML = \'<div class="history-empty">Nothing submitted yet. Your queued analyses will appear here.</div>\';\n        return;\n      }\n\n      historyList.innerHTML = \'\';\n      bugs.forEach(function(bug){\n        const c = sevColors[bug.severity] || sevColors.medium;\n        const time = new Date(bug.created_at).toLocaleString();\n        const item = document.createElement(\'div\');\n        item.className = \'history-item\';\n        item.innerHTML =\n          \'<span class="history-id">\' + escapeHtml(bug.bug_id) + \'</span>\' +\n          \'<span class="history-title-txt">\' + escapeHtml(bug.title) + \'</span>\' +\n          \'<span class="history-sev" style="background:\' + c.bg + \';color:\' + c.text + \'">\' + bug.severity + \'</span>\' +\n          \'<span class="history-time">\' + time + \'</span>\';\n        item.addEventListener(\'click\', function(){ openDetail(bug.id); });\n        historyList.appendChild(item);\n      });\n    } catch (err) {\n      historyList.innerHTML = \'<div class="history-empty">Failed to load history: \' + escapeHtml(err.message) + \'</div>\';\n    }\n  }\n\n  refreshHistoryBtn.addEventListener(\'click\', loadHistory);\n\n  const modal = document.getElementById(\'detailModal\');\n  document.getElementById(\'closeModal\').addEventListener(\'click\', function(){ modal.classList.add(\'hidden\'); });\n  modal.addEventListener(\'click\', function(e){ if (e.target === modal) modal.classList.add(\'hidden\'); });\n\n  function renderAgentSection(analysis){\n    if (!analysis) {\n      return \'<div class="agent-section"><div class="agent-section-title">Agent Analysis</div><div style="color:var(--muted-dim);font-size:12.5px;">Not available for this submission.</div></div>\';\n    }\n    const t = analysis.triage;\n    const l = analysis.log_analysis;\n    let html = \'<div class="agent-section"><div class="agent-section-title">Triage Agent</div>\';\n    html += \'<div class="agent-grid">\';\n    html += \'<div><div class="agent-kv-label">Severity</div><div class="agent-kv-value">\' + escapeHtml(t.severity) + \'</div></div>\';\n    html += \'<div><div class="agent-kv-label">Priority</div><div class="agent-kv-value">\' + escapeHtml(t.priority) + \'</div></div>\';\n    html += \'<div><div class="agent-kv-label">Component</div><div class="agent-kv-value">\' + escapeHtml(t.component) + \'</div></div>\';\n    html += \'<div><div class="agent-kv-label">Type</div><div class="agent-kv-value">\' + escapeHtml(t.type) + \'</div></div>\';\n    html += \'</div>\';\n    html += \'<div class="agent-kv-label">Confidence — \' + Math.round(t.confidence*100) + \'%</div>\';\n    html += \'<div class="confidence-bar"><div class="confidence-fill" style="width:\' + Math.round(t.confidence*100) + \'%"></div></div>\';\n    html += \'<div class="agent-reasoning">\' + escapeHtml(t.reasoning) + \'</div>\';\n    html += \'</div>\';\n\n    html += \'<div class="agent-section"><div class="agent-section-title">Log Analysis Agent</div>\';\n    html += \'<div class="agent-grid">\';\n    html += \'<div><div class="agent-kv-label">Exception</div><div class="agent-kv-value">\' + escapeHtml(l.exception_type || \'—\') + \'</div></div>\';\n    html += \'<div><div class="agent-kv-label">Language</div><div class="agent-kv-value">\' + escapeHtml(l.language || \'—\') + \'</div></div>\';\n    html += \'<div><div class="agent-kv-label">Failure point</div><div class="agent-kv-value">\' + escapeHtml((l.failure_file || \'—\') + (l.failure_line ? (\':\' + l.failure_line) : \'\')) + \'</div></div>\';\n    html += \'<div><div class="agent-kv-label">Error signature</div><div class="agent-kv-value">\' + escapeHtml(l.error_signature || \'—\') + \'</div></div>\';\n    html += \'</div>\';\n    if (l.message) {\n      html += \'<div class="agent-kv-label">Message</div><div class="agent-kv-value" style="margin-bottom:8px;">\' + escapeHtml(l.message) + \'</div>\';\n    }\n    if (l.affected_paths && l.affected_paths.length) {\n      html += \'<div class="agent-kv-label">Affected code path(s)</div><div class="agent-kv-value" style="margin-bottom:8px;">\' + escapeHtml(l.affected_paths.join(\', \')) + \'</div>\';\n    }\n    html += \'<div class="agent-kv-label">Confidence — \' + Math.round(l.confidence*100) + \'%</div>\';\n    html += \'<div class="confidence-bar"><div class="confidence-fill" style="width:\' + Math.round(l.confidence*100) + \'%"></div></div>\';\n    html += \'<div class="agent-reasoning">\' + escapeHtml(l.reasoning) + \'</div>\';\n    html += \'</div>\';\n    return html;\n  }\n\n  async function openDetail(id){\n    try {\n      const res = await fetch(API_BASE + \'/api/bugs/\' + id);\n      const bug = await res.json();\n      if (!res.ok) throw new Error(bug.detail || \'Failed to load detail.\');\n      document.getElementById(\'modalTitle\').textContent = bug.bug_id + \' — \' + bug.title;\n      document.getElementById(\'modalMeta\').textContent =\n        (bug.source_type === \'upload\' ? \'File: \' + bug.original_filename : \'Source: Pasted text\') +\n        \' · \' + bug.severity + \' · \' + new Date(bug.created_at).toLocaleString() + \' · \' + bug.content_length + \' chars\';\n      document.getElementById(\'modalContent\').textContent = bug.content;\n\n      let analysis = null;\n      try {\n        const aRes = await fetch(API_BASE + \'/api/bugs/\' + id + \'/analysis\');\n        if (aRes.ok) analysis = await aRes.json();\n      } catch (e) { /* analysis not available */ }\n      document.getElementById(\'modalAgentSection\').innerHTML = renderAgentSection(analysis);\n\n      modal.classList.remove(\'hidden\');\n    } catch (err) {\n      showConfirm(true, \'Could not load submission\', err.message);\n    }\n  }\n\n  updatePreview();\n  loadHistory();\n})();\n</script>\n</body>\n</html>\n'


@app.get("/", response_class=HTMLResponse)
def serve_frontend():
    return FRONTEND_HTML


# ---------------------------------------------------------------------------
# API routes
# ---------------------------------------------------------------------------

@app.get("/api/health")
def health_check():
    return {"status": "ok", "time": datetime.utcnow().isoformat()}


@app.post("/api/bugs/paste", response_model=BugSubmissionCreatedResponse)
def submit_pasted_bug(
    title: str = Form(...),
    content: str = Form(...),
    db: Session = Depends(get_db),
):
    content = content.strip()
    if not content:
        raise HTTPException(status_code=400, detail="Pasted content cannot be empty.")
    if not title.strip():
        raise HTTPException(status_code=400, detail="Title cannot be empty.")

    # --- Milestone 2: run Triage + Log Analysis BEFORE creating the record,
    # so severity is set by the Triage Agent's prediction rather than a
    # manual user choice. -------------------------------------------------
    combined = run_orchestration(content)
    sev = _validate_severity(combined["triage"]["severity"])

    bug_number = _next_bug_number(db)
    bug = BugSubmission(
        bug_number=bug_number,
        bug_id=_format_bug_id(bug_number),
        title=title.strip(),
        severity=sev,
        source_type=SourceType.PASTE,
        original_filename=None,
        file_type="text/plain",
        content=content,
        content_length=str(len(content)),
    )
    db.add(bug)
    db.commit()
    db.refresh(bug)

    record = _persist_analysis(db, bug, combined)

    return {
        "id": bug.id,
        "bug_id": bug.bug_id,
        "message": "Bug report submitted and analyzed successfully.",
        "analysis": _serialize_analysis(record, bug.bug_id),
    }


@app.post("/api/bugs/upload", response_model=BugSubmissionCreatedResponse)
async def submit_uploaded_bug(
    title: str = Form(...),
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
):
    if not title.strip():
        raise HTTPException(status_code=400, detail="Title cannot be empty.")
    if not files:
        raise HTTPException(status_code=400, detail="At least one file is required.")
    if len(files) > MAX_FILES_PER_UPLOAD:
        raise HTTPException(
            status_code=400,
            detail=f"Too many files. Max {MAX_FILES_PER_UPLOAD} files per submission.",
        )

    created_ids = []
    first_bug_id = None
    first_bug = None
    first_combined = None
    for file in files:
        ext = _validate_extension(file.filename)

        raw = await file.read()
        if len(raw) == 0:
            raise HTTPException(status_code=400, detail=f"'{file.filename}' is empty.")
        if len(raw) > MAX_FILE_SIZE_BYTES:
            max_mb = MAX_FILE_SIZE_BYTES // (1024 * 1024)
            raise HTTPException(
                status_code=400,
                detail=f"'{file.filename}' is too large. Max size is {max_mb} MB.",
            )

        # Dispatches to plain-text decoding, PDF text extraction, or DOCX
        # text extraction, depending on the file's extension.
        text_content = _extract_text_content(file.filename, ext, raw)

        # --- Milestone 2: run Triage + Log Analysis BEFORE creating the
        # record, so severity is set by the Triage Agent's prediction
        # rather than a manual user choice. --------------------------------
        combined = run_orchestration(text_content)
        sev = _validate_severity(combined["triage"]["severity"])

        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S%f")
        safe_name = f"{timestamp}_{os.path.basename(file.filename)}"
        disk_path = os.path.join(UPLOAD_DIR, safe_name)
        with open(disk_path, "wb") as f:
            f.write(raw)

        bug_title = title.strip() if len(files) == 1 else f"{title.strip()} — {file.filename}"
        bug_number = _next_bug_number(db)
        bug = BugSubmission(
            bug_number=bug_number,
            bug_id=_format_bug_id(bug_number),
            title=bug_title,
            severity=sev,
            source_type=SourceType.UPLOAD,
            original_filename=file.filename,
            file_type=ext,
            content=text_content,
            content_length=str(len(text_content)),
        )
        db.add(bug)
        db.commit()
        db.refresh(bug)

        record = _persist_analysis(db, bug, combined)

        created_ids.append(bug.id)
        if first_bug_id is None:
            first_bug_id = bug.bug_id
            first_bug = bug
            first_combined = _serialize_analysis(record, bug.bug_id)

    return {
        "id": created_ids[0],
        "bug_id": first_bug_id,
        "message": f"{len(created_ids)} file(s) uploaded, submitted, and analyzed successfully.",
        "analysis": first_combined,
    }


@app.get("/api/bugs", response_model=List[BugSubmissionResponse])
def list_bugs(db: Session = Depends(get_db)):
    bugs = db.query(BugSubmission).order_by(desc(BugSubmission.created_at)).all()
    return [_serialize(b) for b in bugs]


@app.get("/api/bugs/{bug_id}", response_model=BugSubmissionDetail)
def get_bug(bug_id: str, db: Session = Depends(get_db)):
    """
    Looks up a bug by either its internal UUID (id) or its human-readable
    bug_id (e.g. "BUG-00001"), so both identifiers work as a lookup key.
    """
    bug = (
        db.query(BugSubmission)
        .filter((BugSubmission.id == bug_id) | (BugSubmission.bug_id == bug_id))
        .first()
    )
    if not bug:
        raise HTTPException(status_code=404, detail="Bug submission not found.")
    data = _serialize(bug)
    data["content"] = bug.content
    return data


@app.get("/api/bugs/{bug_id}/analysis")
def get_bug_analysis(bug_id: str, db: Session = Depends(get_db)):
    """
    Returns the combined Triage Agent + Log Analysis Agent output for a bug
    (Milestone 2 deliverable), in the exact structure Milestone 3's
    Duplicate Detection and Root Cause agents are expected to consume.
    """
    bug = (
        db.query(BugSubmission)
        .filter((BugSubmission.id == bug_id) | (BugSubmission.bug_id == bug_id))
        .first()
    )
    if not bug:
        raise HTTPException(status_code=404, detail="Bug submission not found.")
    if not bug.analysis:
        raise HTTPException(status_code=404, detail="No agent analysis found for this bug.")
    return _serialize_analysis(bug.analysis, bug.bug_id)


@app.delete("/api/bugs/{bug_id}")
def delete_bug(bug_id: str, db: Session = Depends(get_db)):
    bug = (
        db.query(BugSubmission)
        .filter((BugSubmission.id == bug_id) | (BugSubmission.bug_id == bug_id))
        .first()
    )
    if not bug:
        raise HTTPException(status_code=404, detail="Bug submission not found.")
    db.delete(bug)
    db.commit()
    return {"message": "Deleted successfully."}


# ---------------------------------------------------------------------------
# Entry point — lets this single file be launched directly with
# `python3 app.py`, in addition to `uvicorn app:app --reload --port 3000`.
# Defaults to port 3000 so it opens at http://localhost:3000 in Chrome.
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import uvicorn

    port = int(os.getenv("PORT", "3000"))
    print(f"Starting AI Smart Bug Analyzer & Fix Advisor on http://localhost:{port}")
    uvicorn.run("app:app", host="0.0.0.0", port=port, reload=True)